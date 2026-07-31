import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt

doc = docx.Document()
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.text = "TEST"

tcPr = cell._tc.get_or_add_tcPr()
tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>'))

p = cell.paragraphs[0]
pPr = p._p.get_or_add_pPr()
pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'))

r = p.runs[0]
r.font.name = "Tahoma"
r.font.size = Pt(10)
r.bold = True

print(cell._tc.xml)

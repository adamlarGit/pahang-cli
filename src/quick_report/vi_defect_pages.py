from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING, Any, Sequence

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

from src.quick_report.cbm_render import _build_jinja_env
from src.quick_report.models import ViDefectPagePlan

if TYPE_CHECKING:
    from src.quick_report.defects import ViDefectRecord

# Constants
DEFECTS_PER_PAGE = 6

EMPTY_DEFECT_CELL_MAP = {
    0: ([0, 1, 2], [0], None),
    1: ([0, 1, 2], [1, 2], None),
    2: ([4, 5, 6], [0], 3),
    3: ([4, 5, 6], [1, 2], 3),
    4: ([8, 9, 10], [0], 7),
    5: ([8, 9, 10], [1, 2], 7),
}


def _clear_cell_text(cell: Any) -> None:
    """Clear paragraph text and run text in cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.text = ""


def _set_cell_no_borders(cell: Any) -> None:
    """Construct/update <w:tcBorders> XML element on cell tcPr with w:val='nil'."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = tcBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tcBorders.append(border)
        border.set(qn('w:val'), 'nil')


def _remove_empty_cell_borders(docx_path: Path, active_count: int) -> None:
    """Clear borders and text for unused VI defect card cells safely using python-docx oxml."""
    if active_count >= DEFECTS_PER_PAGE:
        return

    try:
        doc = Document(docx_path)
        if not doc.tables:
            return

        table = doc.tables[0]
        for idx in range(active_count, DEFECTS_PER_PAGE):
            row_indices, col_indices, spacer_row = EMPTY_DEFECT_CELL_MAP[idx]
            for r in row_indices:
                for c in col_indices:
                    if r < len(table.rows) and c < len(table.rows[r].cells):
                        cell = table.cell(r, c)
                        _clear_cell_text(cell)
                        _set_cell_no_borders(cell)
            if spacer_row is not None and spacer_row < len(table.rows):
                for c in col_indices:
                    if c < len(table.rows[spacer_row].cells):
                        cell = table.cell(spacer_row, c)
                        _clear_cell_text(cell)
                        _set_cell_no_borders(cell)

        doc.save(docx_path)
    except Exception as exc:
        logging.warning(f"Failed to remove empty cell borders for {docx_path}: {exc}")


def build_vi_defect_page_context(pe_info: dict[str, Any], chunk: Sequence[ViDefectRecord]) -> dict[str, Any]:
    """Build context for a single VI defect page."""
    context = pe_info.copy()
    raw_chunk = []
    for idx_slot, item in enumerate(chunk, start=1):
        raw_chunk.append(item.to_dict())
        context[f"equipment{idx_slot}"] = item.equipment
        context[f"description{idx_slot}"] = item.defect_area
        context[f"remark{idx_slot}"] = item.additional_remarks

    # Pad to DEFECTS_PER_PAGE with empty dicts for unused slots
    empty_defect = {"equipment": "", "defect_area": "", "remarks": ""}
    for idx_slot in range(len(chunk) + 1, DEFECTS_PER_PAGE + 1):
        raw_chunk.append(empty_defect.copy())
        context[f"equipment{idx_slot}"] = ""
        context[f"description{idx_slot}"] = ""
        context[f"remark{idx_slot}"] = ""

    context["defects"] = raw_chunk
    return context


class ViDefectPageBuilder:
    """Builder for constructing VI defect page rendering plans."""

    def build(
        self,
        defects: Sequence[ViDefectRecord],
        template_path: Path,
        pe_info: dict[str, Any],
        substation_number: int,
    ) -> list[ViDefectPagePlan]:
        """Build immutable ViDefectPagePlan objects for the given VI defects."""
        template_p = Path(template_path)
        if not defects or not template_p.exists():
            return []

        chunks = [defects[i : i + DEFECTS_PER_PAGE] for i in range(0, len(defects), DEFECTS_PER_PAGE)]
        plans: list[ViDefectPagePlan] = []

        for part_number, chunk in enumerate(chunks, start=1):
            context = build_vi_defect_page_context(pe_info, chunk)
            output_filename = f"{substation_number:03d}_6 VI DEFECT part{part_number}.docx"
            plans.append(
                ViDefectPagePlan(
                    template_path=template_p,
                    output_filename=output_filename,
                    context=context,
                    active_defect_count=len(chunk),
                )
            )

        return plans


def generate_vi_defect_pages(
    defects: Sequence[ViDefectRecord],
    template_path: str | Path,
    output_dir: str | Path,
    substation_number: int,
    pe_info: dict[str, Any],
) -> list[Path]:
    """Generate VI defect pages dynamically with automatic pagination."""
    template_p = Path(template_path)
    if not template_p.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    plans = ViDefectPageBuilder().build(
        defects=defects,
        template_path=template_p,
        pe_info=pe_info,
        substation_number=substation_number,
    )

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    output_paths: list[Path] = []

    for plan in plans:
        output_path = out_dir / plan.output_filename
        doc = DocxTemplate(str(plan.template_path))
        doc.render(plan.context, jinja_env=_build_jinja_env(), autoescape=True)
        doc.save(str(output_path))
        del doc
        _remove_empty_cell_borders(output_path, plan.active_defect_count)
        output_paths.append(output_path)

    return output_paths


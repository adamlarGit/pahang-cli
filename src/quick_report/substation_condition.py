"""Substation condition page generation for Quick Reports (Part 5)."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from docxcompose.composer import Composer as DocxComposer
from docxtpl import DocxTemplate

from src.quick_report.utils import clear_cell_text, set_cell_no_borders
from src.testsheet.models import SubstationEquipmentPackage, SubstationTestsheetPackage


def build_substation_condition_pairs(pkg: SubstationTestsheetPackage | None = None) -> list[tuple[str, str]]:
    """Build active 2-column pairs for the substation condition page based on substation technology and equipment inventory."""
    if not pkg or not getattr(pkg, "data", None):
        return [("SUBSTATION OVERVIEW", "SIGNBOARD")]

    data = pkg.data
    eq = getattr(data, "equipment", None)
    if not isinstance(eq, SubstationEquipmentPackage):
        eq = SubstationEquipmentPackage()

    pairs: list[tuple[str, str]] = [("SUBSTATION OVERVIEW", "SIGNBOARD")]

    # Switchgear
    swg_count = len(eq.switchgears)
    if swg_count == 1:
        pairs.append(("SWITCHGEAR", "SWITCHGEAR NAMEPLATE"))
    elif swg_count >= 2:
        for i in range(1, swg_count + 1):
            pairs.append((f"SWITCHGEAR {i}", f"SWITCHGEAR {i} NAMEPLATE"))

    # Transformer
    tx_count = eq.transformer_count
    if tx_count == 1:
        pairs.append(("TRANSFORMER", "TRANSFORMER NAMEPLATE"))
    elif tx_count >= 2:
        for i in range(1, tx_count + 1):
            pairs.append((f"TRANSFORMER {i}", f"TRANSFORMER {i} NAMEPLATE"))

    # LVDB / Feeder Pillar
    lvdb_specs = eq.lvdb_specs
    if len(lvdb_specs) == 1:
        label = "FEEDER PILLAR" if (lvdb_specs[0].label or "").upper() == "FP" else "LVDB"
        pairs.append((label, f"{label} NAMEPLATE"))
    elif len(lvdb_specs) >= 2:
        for i, spec in enumerate(lvdb_specs, 1):
            label = "FEEDER PILLAR" if (spec.label or "").upper() == "FP" else "LVDB"
            source = spec.source.strip() if (spec.source and spec.source.strip()) else f"TX{i}"
            pairs.append((f"{label} {source}", f"{label} {source} NAMEPLATE"))

    # Battery Charger & RTU
    bc_count = len(eq.battery_banks) if len(eq.battery_banks) > 0 else (1 if eq.has_battery_charger else 0)
    if bc_count == 1:
        pairs.append(("BATTERY CHARGER", "BATTERY CHARGER NAMEPLATE"))
    elif bc_count >= 2:
        for i in range(1, bc_count + 1):
            pairs.append((f"BATTERY CHARGER {i}", f"BATTERY CHARGER {i} NAMEPLATE"))
    if eq.has_rtu:
        pairs.append(("RTU", "RTU NAMEPLATE"))

    # 5. Dual SF6 Indicator
    if len(eq.switchgears) >= 2 and eq.has_sf6:
        pairs.append(("SF6 INDICATOR 1", "SF6 INDICATOR 2"))

    # 6. Single Indicators
    single_items: list[str] = []
    if eq.has_efi:
        single_items.append("EFI")
    if eq.has_sf6 and not (len(eq.switchgears) >= 2 and eq.has_sf6):
        single_items.append("SF6 INDICATOR")

    tx_oil_level_single_handled = False
    if tx_count == 1 and len(single_items) % 2 == 1:
        single_items.append("TRANSFORMER OIL LEVEL INDICATOR")
        tx_oil_level_single_handled = True

    for i in range(0, len(single_items), 2):
        if i + 1 < len(single_items):
            pairs.append((single_items[i], single_items[i + 1]))
        else:
            pairs.append((single_items[i], ""))

    # 7. Fire Extinguisher (Dedicated Rooms)
    b_type = (data.building_type or "").upper()
    s_type = (data.substation_type or "").upper()
    if b_type in ("INDOOR", "ATTACH") and "OUTDOOR" not in b_type and "COMPACT" not in b_type and "CS" not in s_type:
        fe = getattr(eq, "fire_extinguisher", None)
        if fe and fe.has_fire_extinguisher:
            # Switchgear Room
            if swg_count == 1:
                pairs.append(("FIRE EXTINGUISHER\n(SWITCHGEAR ROOM)", "FIRE EXTINGUISHER EXPIRY DATE"))
            elif swg_count >= 2:
                for i in range(1, swg_count + 1):
                    pairs.append((f"FIRE EXTINGUISHER\n(SWITCHGEAR {i} ROOM)", "FIRE EXTINGUISHER EXPIRY DATE"))
            elif swg_count == 0:
                pairs.append(("FIRE EXTINGUISHER\n(SWITCHGEAR ROOM)", "FIRE EXTINGUISHER EXPIRY DATE"))

            # Transformer Room(s)
            if tx_count == 1:
                pairs.append(("FIRE EXTINGUISHER\n(TX ROOM)", "FIRE EXTINGUISHER EXPIRY DATE"))
            elif tx_count >= 2:
                for i in range(1, tx_count + 1):
                    pairs.append((f"FIRE EXTINGUISHER\n(TX{i} ROOM)", "FIRE EXTINGUISHER EXPIRY DATE"))

    # 8. Transformer Oil Level Indicator (if not already packed)
    if tx_count >= 2:
        pairs.append(("TRANSFORMER 1 OIL LEVEL INDICATOR", "TRANSFORMER 2 OIL LEVEL INDICATOR"))
    elif tx_count == 1 and not tx_oil_level_single_handled:
        pairs.append(("TRANSFORMER OIL LEVEL INDICATOR", ""))

    return pairs


def _remove_empty_cell_borders_sub_cond(
    docx_path: Path,
    active_count: Sequence[tuple[str, str]] | int | None = None,
) -> None:
    """Clear borders and text for unused substation condition cells safely using python-docx oxml."""
    if active_count is None:
        return

    if isinstance(active_count, int):
        if active_count >= 3:
            return
        chunk: list[tuple[str, str]] = [("X", "X")] * active_count
    elif isinstance(active_count, (list, tuple)):
        chunk = list(active_count)
        if len(chunk) == 3 and all(p[0] and p[1] for p in chunk):
            return
    else:
        return

    try:
        doc = Document(docx_path)
        if not doc.tables:
            return

        if len(doc.tables) >= 3:
            for slot_idx in range(3):
                if slot_idx >= len(doc.tables):
                    break
                table = doc.tables[slot_idx]
                if slot_idx >= len(chunk):
                    for row in table.rows:
                        for cell in row.cells:
                            clear_cell_text(cell)
                            set_cell_no_borders(cell)
                else:
                    pair = chunk[slot_idx]
                    left_active = bool(pair[0]) if len(pair) > 0 else False
                    right_active = bool(pair[1]) if len(pair) > 1 else False

                    if not left_active and not right_active:
                        for row in table.rows:
                            for cell in row.cells:
                                clear_cell_text(cell)
                                set_cell_no_borders(cell)
                    else:
                        for row in table.rows:
                            if not right_active:
                                right_cell = row.cells[-1]
                                clear_cell_text(right_cell)
                                set_cell_no_borders(right_cell)
                            if not left_active:
                                left_cell = row.cells[0]
                                clear_cell_text(left_cell)
                                set_cell_no_borders(left_cell)

        elif len(doc.tables) == 1:
            table = doc.tables[0]
            num_rows = len(table.rows)

            if num_rows >= 8:
                slot_row_map = {0: (0, 1), 1: (3, 4), 2: (6, 7)}
                spacer_above_map = {1: 2, 2: 5}
                spacer_below_map = {0: 2, 1: 5}
            elif num_rows == 6:
                slot_row_map = {0: (0, 1), 1: (2, 3), 2: (4, 5)}
                spacer_above_map = {}
                spacer_below_map = {}
            else:
                slot_row_map = {i: (i * 2, i * 2 + 1) for i in range(3)}
                spacer_above_map = {}
                spacer_below_map = {}

            for slot_idx, (r_hdr, r_photo) in slot_row_map.items():
                if r_hdr >= num_rows or r_photo >= num_rows:
                    continue

                r_sp_above = spacer_above_map.get(slot_idx)
                r_sp_below = spacer_below_map.get(slot_idx)

                if slot_idx >= len(chunk):
                    for r_idx in (r_hdr, r_photo):
                        for cell in table.rows[r_idx].cells:
                            clear_cell_text(cell)
                            set_cell_no_borders(cell)
                    if r_sp_above is not None and r_sp_above < num_rows:
                        for cell in table.rows[r_sp_above].cells:
                            clear_cell_text(cell)
                            set_cell_no_borders(cell)
                else:
                    pair = chunk[slot_idx]
                    left_active = bool(pair[0]) if len(pair) > 0 else False
                    right_active = bool(pair[1]) if len(pair) > 1 else False

                    if not left_active and not right_active:
                        for r_idx in (r_hdr, r_photo):
                            for cell in table.rows[r_idx].cells:
                                clear_cell_text(cell)
                                set_cell_no_borders(cell)
                        if r_sp_above is not None and r_sp_above < num_rows:
                            for cell in table.rows[r_sp_above].cells:
                                clear_cell_text(cell)
                                set_cell_no_borders(cell)
                    else:
                        num_cols = len(table.rows[r_hdr].cells)
                        if not right_active:
                            for r_idx in (r_hdr, r_photo):
                                if num_cols >= 3:
                                    clear_cell_text(table.rows[r_idx].cells[2])
                                    set_cell_no_borders(table.rows[r_idx].cells[2])
                                    set_cell_no_borders(table.rows[r_idx].cells[1])
                                elif num_cols >= 2:
                                    clear_cell_text(table.rows[r_idx].cells[1])
                                    set_cell_no_borders(table.rows[r_idx].cells[1])
                            if r_sp_above is not None and r_sp_above < num_rows and num_cols >= 3:
                                set_cell_no_borders(table.rows[r_sp_above].cells[2])
                                set_cell_no_borders(table.rows[r_sp_above].cells[1])
                            if r_sp_below is not None and r_sp_below < num_rows and num_cols >= 3:
                                set_cell_no_borders(table.rows[r_sp_below].cells[2])
                                set_cell_no_borders(table.rows[r_sp_below].cells[1])

                        if not left_active:
                            for r_idx in (r_hdr, r_photo):
                                clear_cell_text(table.rows[r_idx].cells[0])
                                set_cell_no_borders(table.rows[r_idx].cells[0])
                                if num_cols >= 3:
                                    set_cell_no_borders(table.rows[r_idx].cells[1])
                            if r_sp_above is not None and r_sp_above < num_rows:
                                set_cell_no_borders(table.rows[r_sp_above].cells[0])
                            if r_sp_below is not None and r_sp_below < num_rows:
                                set_cell_no_borders(table.rows[r_sp_below].cells[0])

        doc.save(docx_path)
    except Exception as exc:
        logging.warning(f"Failed to remove empty cell borders for {docx_path}: {exc}")


def generate_substation_condition_pages(
    pe_info: dict,
    condition_pairs_or_pkg: Sequence[tuple[str, str]] | SubstationTestsheetPackage | None,
    template_path: Path | str,
    output_dir: Path | str,
    substation_number: int,
) -> list[Path]:
    """Generate substation condition pages."""
    t_path = Path(template_path)
    if not t_path.exists():
        raise FileNotFoundError(f"Template path does not exist: {t_path}")

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    if isinstance(condition_pairs_or_pkg, (list, tuple)):
        pairs = list(condition_pairs_or_pkg)
    else:
        pairs = build_substation_condition_pairs(condition_pairs_or_pkg)
    chunks = [pairs[i:i + 3] for i in range(0, len(pairs), 3)]
    if not chunks:
        chunks = [[]]

    parts: list[Path] = []

    for idx, chunk in enumerate(chunks, start=1):
        cond_out = out_dir / f"{substation_number:03d}_5 SUBSTATION CONDITION part{idx}.docx"

        padded_chunk = list(chunk)
        while len(padded_chunk) < 3:
            padded_chunk.append(("", ""))

        context = {
            "pairs": [
                {
                    "header_left": p[0],
                    "header_right": p[1],
                    "photo_left": "",
                    "photo_right": "",
                }
                for p in padded_chunk
            ]
        }
        context.update(pe_info)

        doc = DocxTemplate(str(t_path))
        doc.render(context)

        # Shrink trailing paragraph to near-zero height to prevent overflow.
        # Word requires a mandatory <w:p> after every table; if the 3-pair
        # table fills the page, this paragraph spills onto page 2 and
        # causes blank pages during COM assembly. Setting line_spacing_rule
        # to EXACTLY is critical — without it Word ignores the small size.
        if doc.paragraphs:
            last_para = doc.paragraphs[-1]
            pf = last_para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(0.5)
            for run in last_para.runs:
                run.font.size = Pt(0.5)
            if not last_para.runs:
                run = last_para.add_run()
                run.font.size = Pt(0.5)

        doc.save(str(cond_out))

        _remove_empty_cell_borders_sub_cond(cond_out, chunk)

        parts.append(cond_out)

    # Merge all condition part files into a single docx so the COM
    # assembly stage pastes one file (no inter-part page breaks that
    # cause blank pages between full-page condition tables).
    if len(parts) > 1:
        merged_path = out_dir / f"{substation_number:03d}_5 SUBSTATION CONDITION.docx"
        base_doc = Document(str(parts[0]))
        composer = DocxComposer(base_doc)
        for part_path in parts[1:]:
            composer.append(Document(str(part_path)))
        composer.save(str(merged_path))

        # Post-merge: remove trailing section break and re-shrink the last paragraph.
        # docxcompose adds a trailing nextPage section break to the final paragraph
        # and may reset formatting during the merge process.
        merged_doc = Document(str(merged_path))
        if merged_doc.paragraphs:
            last_para = merged_doc.paragraphs[-1]

            # Remove trailing section break if present
            pPr = last_para._p.get_or_add_pPr()
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None:
                pPr.remove(sectPr)

            # Remove trailing <w:br> page breaks injected by docxcompose
            for run in last_para.runs:
                r_elem = run._r
                for br in r_elem.findall(qn("w:br")):
                    r_elem.remove(br)
            for br in last_para._p.findall(qn("w:br")):
                last_para._p.remove(br)

            pf = last_para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(0.5)
            for run in last_para.runs:
                run.font.size = Pt(0.5)
            if not last_para.runs:
                run = last_para.add_run()
                run.font.size = Pt(0.5)
        merged_doc.save(str(merged_path))

        # Clean up individual part files.
        for p in parts:
            try:
                p.unlink(missing_ok=True)
            except Exception:
                pass
        return [merged_path]

    return parts

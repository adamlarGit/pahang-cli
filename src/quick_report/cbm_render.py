"""CBM quick-report rendering engine and context builders."""

from __future__ import annotations

import gc
from pathlib import Path
import re
from typing import Any

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from jinja2 import Environment, Undefined

from src.core.normalizers import (
    format_db_int,
    format_temperature_float,
    normalize_us_characteristic,
)
from src.quick_report.cbm_family import QuickReportFamilySpec
from src.quick_report.defects import CbmDefectRecord
from src.quick_report.prpd import (
    discover_ultratev_survey_dir,
    generate_prpd_graphs_for_swg_panel,
    generate_prpd_graphs_for_transformer,
)
from src.quick_report.utils import clear_cell_text, set_cell_shading
from src.testsheet.models import (
    BatteryBankSpec,
    LVDBSpec,
    SubstationEquipmentPackage,
    SwitchgearPanelSpec,
    TransformerSpec,
)


class PreservingUndefined(Undefined):
    """Render undefined quick-report Jinja placeholders as '-'."""

    def __str__(self) -> str:
        return "-"

    def __getattr__(self, name: str):
        if name.startswith("_"):
            raise AttributeError(name)
        return self

    def __getitem__(self, name):
        return self

    def __iter__(self):
        return iter(())

    def __bool__(self):
        return False

    def __len__(self):
        return 0


class QuickReportContext(dict):
    """Dict wrapper that renders missing/empty values as clean '-'."""

    def __init__(self, value: dict | None = None, path: str = ""):
        super().__init__()
        self._path = path
        if value:
            for key, item in value.items():
                super().__setitem__(key, self._wrap_value(item, self._child_path(key)))

    def _child_path(self, key: str) -> str:
        return f"{self._path}.{key}" if self._path else str(key)

    def _wrap_value(self, value, path: str):
        if isinstance(value, dict):
            return QuickReportContext(value, path)
        if isinstance(value, list):
            return [self._wrap_value(item, path) for item in value]
        if isinstance(value, InlineImage):
            return value
        if path.endswith(".prpd") or path.endswith("_image") or path.endswith(".image") or path == "prpd":
            if isinstance(value, (InlineImage, Path)):
                return value
            if value is None or (isinstance(value, str) and not value.strip()) or value == "-":
                return ""
        if value is None:
            return "-"
        if isinstance(value, str) and not value.strip():
            return "-"
        return value

    def __missing__(self, key):
        return "-"

    def __getattr__(self, name):
        if name.startswith("_"):
            raise AttributeError(name)
        try:
            return self[name]
        except KeyError:
            return "-"


def _build_jinja_env() -> Environment:
    """Create a Jinja environment that renders undefined placeholders as '-'."""
    return Environment(undefined=PreservingUndefined, autoescape=True)


def _preserve_blank_render_values(value):
    """Wrap quick-report render data so missing keys render as clean '-'."""
    if isinstance(value, dict):
        return QuickReportContext(value)
    if isinstance(value, list):
        return [QuickReportContext(item) if isinstance(item, dict) else item for item in value]
    return value


def _process_inline_images(doc: DocxTemplate, context: dict) -> None:
    """Recursively convert file path images (like PRPD graphs) into InlineImage instances."""
    def _convert(obj: Any) -> None:
        if isinstance(obj, dict):
            for k, v in list(obj.items()):
                k_str = str(k)
                if k_str == "prpd" or k_str.endswith("_image") or k_str.endswith(".image"):
                    if isinstance(v, (str, Path)) and str(v).strip() and str(v) != "-":
                        v_path = Path(v)
                        if v_path.is_file():
                            obj[k] = InlineImage(doc, str(v_path), width=Mm(80))
                        else:
                            obj[k] = ""
                    elif isinstance(v, InlineImage):
                        pass
                    else:
                        obj[k] = ""
                else:
                    _convert(v)
        elif isinstance(obj, list):
            for item in obj:
                _convert(item)

    _convert(context)


def _render_docx_template(
    template_path: str | Path,
    output_path: Path,
    context: dict,
    *,
    defective_technologies: set[str] | list[str] | tuple[str, ...] | str | None = None,
    overview: bool | None = None,
) -> Path:
    """Render a DocxTemplate with quick-report placeholder semantics and severity cell shading."""
    doc = DocxTemplate(str(template_path))
    rendered_context = dict(context)
    _process_inline_images(doc, rendered_context)
    doc.render(_preserve_blank_render_values(rendered_context), jinja_env=_build_jinja_env(), autoescape=True)

    is_overview = overview if overview is not None else bool(context.get("__is_overview__", False))
    def_techs: set[str] = set()
    if defective_technologies is not None:
        if isinstance(defective_technologies, str):
            def_techs = {t.strip().upper() for t in defective_technologies.replace(",", " ").replace("+", " ").split() if t.strip()}
        else:
            def_techs = {str(t).strip().upper() for t in defective_technologies if str(t).strip()}
    elif "__defective_technologies__" in context:
        raw_dt = context["__defective_technologies__"]
        if isinstance(raw_dt, str):
            def_techs = {t.strip().upper() for t in raw_dt.replace(",", " ").replace("+", " ").split() if t.strip()}
        elif isinstance(raw_dt, (set, list, tuple)):
            def_techs = {str(t).strip().upper() for t in raw_dt if str(t).strip()}

    # Post-process table cells for severity shading
    for table in doc.docx.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if "__SEVERITY_IR__" in text or "{{ ir.severity }}" in text:
                    clear_cell_text(cell)
                    if not is_overview:
                        set_cell_shading(cell, "EE0000" if "IR" in def_techs else "00B050")
                    else:
                        cell.paragraphs[0].text = "-"
                elif "__SEVERITY_US__" in text or "{{ us.severity }}" in text:
                    clear_cell_text(cell)
                    if not is_overview:
                        set_cell_shading(cell, "EE0000" if "US" in def_techs else "00B050")
                    else:
                        cell.paragraphs[0].text = "-"
                elif "__SEVERITY_TEV__" in text or "{{ tev.severity }}" in text:
                    clear_cell_text(cell)
                    if not is_overview:
                        set_cell_shading(cell, "EE0000" if "TEV" in def_techs else "00B050")
                    else:
                        cell.paragraphs[0].text = "-"

    doc.save(output_path)
    del doc
    gc.collect()
    return output_path


def _text_or_empty(value: Any) -> str:
    if value is None:
        return ""
    try:
        if value != value:
            return ""
    except Exception:
        pass
    return str(value).strip()


def _fallback_dash(value: Any) -> str:
    s = _text_or_empty(value)
    return s if s else "-"


def _format_detail_area(defect_area: Any, additional_remarks: Any, *, overview: bool = False) -> str:
    if overview:
        return "OVERVIEW"
    defect_area_text = _text_or_empty(defect_area)
    remarks_text = _text_or_empty(additional_remarks)
    if remarks_text:
        return f"{defect_area_text}/ {remarks_text}" if defect_area_text else remarks_text
    return defect_area_text if defect_area_text else "-"


def _extract_equipment_package(pe_info: dict[str, Any] | None) -> SubstationEquipmentPackage | None:
    if not pe_info:
        return None
    for key in ("equipment_specs", "equipment_package", "equipment"):
        val = pe_info.get(key)
        if isinstance(val, SubstationEquipmentPackage):
            return val
    return None


def _find_matching_switchgear_panel(
    panels: list[SwitchgearPanelSpec] | tuple[SwitchgearPanelSpec, ...],
    target_id: str,
) -> SwitchgearPanelSpec | None:
    if not panels or not target_id:
        return None
    t = target_id.strip().upper()
    if not t:
        return None

    # 1. Exact match on panel_feeder_no or name
    for p in panels:
        if p.panel_feeder_no and p.panel_feeder_no.strip().upper() == t:
            return p
        if p.name and p.name.strip().upper() == t:
            return p

    # 2. Canonical panel naming match (e.g. PANEL 1, P1)
    for p in panels:
        p_no_str = str(p.panel_no)
        p_feeder = p.panel_feeder_no.strip().upper()
        if t in (f"PANEL {p_no_str}", f"P{p_no_str}", f"PANEL {p_feeder}", f"P{p_feeder}"):
            return p

    # 3. Numeric digit extraction matching
    t_digits = re.findall(r"\d+", t)
    if t_digits:
        t_num = t_digits[0]
        for p in panels:
            if str(p.panel_no) == t_num or p.panel_feeder_no.strip() == t_num:
                return p

    # 4. Substring containment if name has >= 2 chars
    for p in panels:
        p_name = p.name.strip().upper()
        if len(p_name) >= 2 and (p_name in t or t in p_name):
            return p

    return None


def _find_matching_transformer(
    transformers: tuple[TransformerSpec, ...],
    target_id: str,
) -> TransformerSpec | None:
    if not transformers:
        return None
    t = target_id.strip().upper()
    if t:
        for tx in transformers:
            if tx.tx_id and tx.tx_id.strip().upper() == t:
                return tx
        t_digits = re.findall(r"\d+", t)
        if t_digits:
            t_num = t_digits[0]
            for tx in transformers:
                tx_digits = re.findall(r"\d+", tx.tx_id)
                if tx_digits and tx_digits[0] == t_num:
                    return tx
    if len(transformers) == 1:
        return transformers[0]
    return None


def _find_matching_lvdb(
    lvdb_specs: tuple[LVDBSpec, ...],
    target_id: str,
) -> LVDBSpec | None:
    if not lvdb_specs:
        return None
    t = target_id.strip().upper()
    if t:
        for lv in lvdb_specs:
            if lv.name and lv.name.strip().upper() == t:
                return lv
            if lv.source and lv.source.strip().upper() in t:
                return lv
            combined = f"{lv.label} {lv.source}".strip().upper()
            if combined and combined in t:
                return lv
        t_digits = re.findall(r"\d+", t)
        if t_digits:
            t_num = t_digits[0]
            for lv in lvdb_specs:
                lv_digits = re.findall(r"\d+", lv.name)
                if lv_digits and lv_digits[0] == t_num:
                    return lv
    if len(lvdb_specs) == 1:
        return lvdb_specs[0]
    return None


def _find_matching_battery(
    battery_banks: tuple[BatteryBankSpec, ...],
    target_id: str,
) -> BatteryBankSpec | None:
    if not battery_banks:
        return None
    t = target_id.strip().upper()
    if t:
        for b in battery_banks:
            if b.name and (b.name.strip().upper() == t or t in b.name.strip().upper()):
                return b
        t_digits = re.findall(r"\d+", t)
        if t_digits:
            t_num = t_digits[0]
            for b in battery_banks:
                b_digits = re.findall(r"\d+", b.name)
                if b_digits and b_digits[0] == t_num:
                    return b
    if len(battery_banks) == 1:
        return battery_banks[0]
    return None


def _extract_tev_background(pe_info: dict[str, Any] | None) -> str:
    """Extract TEV background dB value from pe_info or testsheet data."""
    if not pe_info:
        return "-"
    ts_data = pe_info.get("testsheet_data")
    if ts_data and getattr(ts_data, "tev_background", None):
        return format_db_int(ts_data.tev_background)
    if isinstance(pe_info.get("substation"), dict):
        bg = pe_info["substation"].get("tev_bg") or pe_info["substation"].get("tev_background")
        if bg and bg != "-":
            return format_db_int(bg)
    if pe_info.get("tev_background"):
        return format_db_int(pe_info["tev_background"])
    if pe_info.get("tev_bg"):
        return format_db_int(pe_info["tev_bg"])
    return "-"


def _build_fp_lvdb_render_context(
    record: CbmDefectRecord,
    *,
    overview: bool,
    item_key: str = "",
    item_suffix: str = "",
    pe_info: dict[str, Any] | None = None,
) -> dict[str, Any]:
    resolved_item_key = _text_or_empty(item_key)
    resolved_item_suffix = _text_or_empty(item_suffix)
    equipment = _text_or_empty(record.equipment).strip()
    area = _format_detail_area(record.defect_area, record.additional_remarks, overview=overview)

    raw_id = record.equipment_id or resolved_item_key or resolved_item_suffix
    dash_split = re.split(r"\s*[-–—]\s*", raw_id, maxsplit=1)
    if len(dash_split) > 1 and dash_split[1].strip():
        labelsource = dash_split[0].strip()
        feederno = dash_split[1].strip()
    else:
        labelsource = raw_id.strip()
        feederno = "-"

    equipment_pkg = _extract_equipment_package(pe_info)
    lvdb_specs = equipment_pkg.lvdb_specs if equipment_pkg else ()
    matched_lv = _find_matching_lvdb(lvdb_specs, labelsource or raw_id)

    fp_mfg = (matched_lv.manufacturer if matched_lv and matched_lv.manufacturer else "") or _text_or_empty(record.brand)
    
    combined_model = f"{equipment} {record.model}".upper()
    if any(k in combined_model for k in ("(J)", "J-SLOTTED", "J SLOTTED", "J-SLOT", "LVDB")):
        fp_model = "J-SLOTTED"
    elif any(k in combined_model for k in ("(D)", " DIN", "DIN ", "-DIN", "/DIN")):
        fp_model = "DIN"
    elif record.model:
        fp_model = _text_or_empty(record.model)
    elif matched_lv and matched_lv.label:
        fp_model = matched_lv.label
    else:
        fp_model = ""

    fp_rating = _text_or_empty(record.rating) or (matched_lv.rating if matched_lv else "")
    fp_serial = matched_lv.serial_no if matched_lv else ""
    
    fp_cable = ""
    if equipment_pkg:
        for swg in equipment_pkg.switchgears:
            for p in swg.panels:
                if p.cable_type:
                    fp_cable = p.cable_type
                    break
            if fp_cable:
                break
    if not fp_cable and "CABLE" in equipment.upper():
        fp_cable = equipment

    tev_bg = _extract_tev_background(pe_info)
    ir_sev = "-" if overview else "__SEVERITY_IR__"
    us_sev = "-" if overview else "__SEVERITY_US__"
    tev_sev = "-" if overview else "__SEVERITY_TEV__"
    def_techs = {record.technology} if record.technology else set()

    return {
        "__is_overview__": overview,
        "__defective_technologies__": def_techs,
        "fp": {
            "labelsource": _fallback_dash(labelsource),
            "feederno": _fallback_dash(feederno),
            "area": area,
            "manufacturer": _fallback_dash(fp_mfg),
            "model": _fallback_dash(fp_model),
            "rating": _fallback_dash(fp_rating),
            "serialnumber": _fallback_dash(fp_serial),
            "cabletype": _fallback_dash(fp_cable),
            "ir": {
                "reading": format_temperature_float(record.ir_reading),
                "severity": ir_sev,
            },
            "us": {
                "reading": format_db_int(record.us_reading),
                "char": normalize_us_characteristic(record.us_char),
                "severity": us_sev,
                "prpd": "",
            },
            "tev": {
                "reading": format_db_int(record.tev_reading),
                "char": _fallback_dash(record.tev_char),
                "bg": format_db_int(tev_bg),
                "severity": tev_sev,
                "prpd": "",
            },
        },
        "ir": {
            "reading": format_temperature_float(record.ir_reading),
            "severity": ir_sev,
        },
        "us": {
            "reading": format_db_int(record.us_reading),
            "char": normalize_us_characteristic(record.us_char),
            "severity": us_sev,
            "prpd": "",
        },
        "tev": {
            "reading": format_db_int(record.tev_reading),
            "char": _fallback_dash(record.tev_char),
            "bg": format_db_int(tev_bg),
            "severity": tev_sev,
            "prpd": "",
        },
    }


def _build_swg_render_context(
    record: CbmDefectRecord,
    *,
    overview: bool,
    item_key: str = "",
    item_suffix: str = "",
    pe_info: dict[str, Any] | None = None,
) -> dict[str, Any]:
    resolved_item_key = _text_or_empty(item_key)
    resolved_item_suffix = _text_or_empty(item_suffix)
    equipment = _text_or_empty(record.equipment).strip()
    area = _format_detail_area(record.defect_area, record.additional_remarks, overview=overview)

    equipment_pkg = _extract_equipment_package(pe_info)
    swg_spec = equipment_pkg.switchgear if equipment_pkg and equipment_pkg.switchgears else None

    # SWG header fields
    swg_manufacturer = (swg_spec.manufacturer if swg_spec and swg_spec.manufacturer else "") or _text_or_empty(record.brand)
    swg_model = _text_or_empty(record.model) or (swg_spec.model if swg_spec else "")
    swg_rating = _text_or_empty(record.rating) or (swg_spec.rating if swg_spec else "")
    swg_serial = swg_spec.serial_no if swg_spec else ""
    swg_type = equipment or (swg_spec.switchgear_type if swg_spec else "")

    # Match panel in testsheet
    all_panels = [p for swg in (equipment_pkg.switchgears if equipment_pkg else ()) for p in swg.panels]
    panel_match_target = record.equipment_id or resolved_item_suffix or resolved_item_key or equipment
    matched_panel = _find_matching_switchgear_panel(all_panels, panel_match_target)

    if matched_panel:
        panel_loadamp = matched_panel.load_amp
        panel_heateramp = matched_panel.heater_amp
        panel_breakerstatus = matched_panel.status
        panel_cabletype = matched_panel.cable_type or (equipment if "CABLE" in equipment.upper() else "")
        panel_serialnumber = matched_panel.serial_no
        panel_us_reading = matched_panel.us_reading or _text_or_empty(record.us_reading)
        panel_us_char = matched_panel.us_char or _text_or_empty(record.us_char)
        panel_tev_reading = matched_panel.tev_reading or _text_or_empty(record.tev_reading)
        panel_tev_ppc = matched_panel.tev_ppc
        panel_tev_char = matched_panel.tev_char or _text_or_empty(record.tev_char)
    else:
        panel_loadamp = ""
        panel_heateramp = ""
        panel_breakerstatus = ""
        panel_cabletype = equipment if "CABLE" in equipment.upper() else ""
        panel_serialnumber = ""
        panel_us_reading = _text_or_empty(record.us_reading)
        panel_us_char = _text_or_empty(record.us_char)
        panel_tev_reading = _text_or_empty(record.tev_reading)
        panel_tev_ppc = ""
        panel_tev_char = _text_or_empty(record.tev_char)

    if " - " in panel_match_target:
        parts = panel_match_target.split(" - ", 1)
        panel_linknumber = parts[0].strip()
        panel_name = parts[1].strip()
    else:
        panel_name = panel_match_target
        panel_linknumber = panel_match_target

    tev_bg = _extract_tev_background(pe_info)
    ir_sev = "-" if overview else "__SEVERITY_IR__"
    us_sev = "-" if overview else "__SEVERITY_US__"
    tev_sev = "-" if overview else "__SEVERITY_TEV__"
    def_techs = {record.technology} if record.technology else set()

    # Discover survey root and retrieve/generate PRPD images if available
    us_prpd: Path | str = ""
    tev_prpd: Path | str = ""
    if pe_info and not overview:
        panel_no = matched_panel.panel_no if matched_panel else 0
        if panel_no == 0:
            digits = re.findall(r"\d+", panel_match_target)
            if digits:
                try:
                    panel_no = int(digits[0])
                except ValueError:
                    panel_no = 0

        prpd_catalog = pe_info.get("prpd_catalog")
        if prpd_catalog and isinstance(prpd_catalog, dict) and "swg" in prpd_catalog:
            swg_catalog = prpd_catalog["swg"]
            if panel_no in swg_catalog:
                entry = swg_catalog[panel_no]
                us_prpd = entry.get("us") or ""
                tev_prpd = entry.get("tev") or ""
            elif len(swg_catalog) == 1 and 0 in swg_catalog:
                entry = swg_catalog[0]
                us_prpd = entry.get("us") or ""
                tev_prpd = entry.get("tev") or ""

        # Fallback to direct on-demand generation if not in catalog
        if not us_prpd and not tev_prpd:
            raw_dir = pe_info.get("raw_data_dir") or pe_info.get("survey_dir") or pe_info.get("raw_dir")
            survey_root = discover_ultratev_survey_dir(raw_dir)
            prpd_out_dir = pe_info.get("prpd_output_dir")

            if survey_root and prpd_out_dir:
                us_png, tev_png = generate_prpd_graphs_for_swg_panel(
                    survey_root=survey_root,
                    panel_no=panel_no,
                    output_dir=Path(prpd_out_dir),
                    feeder_no=panel_match_target,
                    panel_name=panel_name,
                )
                if us_png:
                    us_prpd = us_png
                if tev_png:
                    tev_prpd = tev_png

    return {
        "__is_overview__": overview,
        "__defective_technologies__": def_techs,
        "swg": {
            "area": area,
            "manufacturer": _fallback_dash(swg_manufacturer),
            "model": _fallback_dash(swg_model),
            "rating": _fallback_dash(swg_rating),
            "serialnumber": _fallback_dash(swg_serial),
            "type": _fallback_dash(swg_type),
        },
        "panel": {
            "name": _fallback_dash(panel_name),
            "linknumber": _fallback_dash(panel_linknumber),
            "area": area,
            "breakerstatus": _fallback_dash(panel_breakerstatus),
            "busbarposition": "-",
            "cabletype": _fallback_dash(panel_cabletype),
            "heateramp": _fallback_dash(panel_heateramp),
            "loadamp": _fallback_dash(panel_loadamp),
            "serialnumber": _fallback_dash(panel_serialnumber),
            "ir": {
                "reading": format_temperature_float(record.ir_reading),
                "severity": ir_sev,
            },
            "us": {
                "reading": format_db_int(panel_us_reading),
                "char": normalize_us_characteristic(panel_us_char),
                "severity": us_sev,
                "prpd": us_prpd,
            },
            "tev": {
                "reading": format_db_int(panel_tev_reading),
                "ppc": _fallback_dash(panel_tev_ppc),
                "char": _fallback_dash(panel_tev_char),
                "bg": format_db_int(tev_bg),
                "severity": tev_sev,
                "prpd": tev_prpd,
            },
        },
        "ir": {
            "reading": format_temperature_float(record.ir_reading),
            "severity": ir_sev,
        },
        "us": {
            "reading": format_db_int(panel_us_reading),
            "char": normalize_us_characteristic(panel_us_char),
            "severity": us_sev,
            "prpd": us_prpd,
        },
        "tev": {
            "reading": format_db_int(panel_tev_reading),
            "ppc": _fallback_dash(panel_tev_ppc),
            "char": _fallback_dash(panel_tev_char),
            "bg": format_db_int(tev_bg),
            "severity": tev_sev,
            "prpd": tev_prpd,
        },
    }


def _build_tx_render_context(
    record: CbmDefectRecord,
    *,
    overview: bool,
    item_key: str = "",
    item_suffix: str = "",
    pe_info: dict[str, Any] | None = None,
) -> dict[str, Any]:
    resolved_item_key = _text_or_empty(item_key)
    resolved_item_suffix = _text_or_empty(item_suffix)
    equipment = _text_or_empty(record.equipment).strip()
    area = _format_detail_area(record.defect_area, record.additional_remarks, overview=overview)

    equipment_pkg = _extract_equipment_package(pe_info)
    transformers = equipment_pkg.transformers if equipment_pkg else ()
    tx_match_target = record.equipment_id or resolved_item_key or equipment
    matched_tx = _find_matching_transformer(transformers, tx_match_target)

    # Location from substation building_type (overview) or HV/LV side (detail)
    if overview:
        tx_location = ""
        if pe_info and isinstance(pe_info.get("substation"), dict):
            tx_location = pe_info["substation"].get("building_type", "")
        if not tx_location and resolved_item_suffix:
            tx_location = resolved_item_suffix
    else:
        search_loc = f"{record.defect_area} {record.additional_remarks} {record.equipment_id} {equipment}".upper()
        if any(k in search_loc for k in ("HV", "11KV", "33KV")):
            tx_location = "HV - SIDE"
        elif any(k in search_loc for k in ("LV", "415V")):
            tx_location = "LV - SIDE"
        else:
            tx_location = ""
            if pe_info and isinstance(pe_info.get("substation"), dict):
                tx_location = pe_info["substation"].get("building_type", "")
            if not tx_location and resolved_item_suffix:
                tx_location = resolved_item_suffix

    tx_mfg = (matched_tx.manufacturer if matched_tx and matched_tx.manufacturer else "") or _text_or_empty(record.brand)
    
    raw_model = (matched_tx.type if matched_tx and matched_tx.type else "") or _text_or_empty(record.model)
    if raw_model.upper() in ("H/S", "HERMETICALLY SEALED", "HERMETICALLY SEAL"):
        tx_model = "HERMETICALLY SEAL"
    else:
        tx_model = raw_model

    tx_rating = (matched_tx.rating_kva if matched_tx and matched_tx.rating_kva else "") or _text_or_empty(record.rating)
    tx_serial = matched_tx.serial_no if matched_tx else ""

    tx_cable = ""
    if equipment_pkg:
        for swg in equipment_pkg.switchgears:
            for p in swg.panels:
                if "TX" in p.name.upper() and p.cable_type:
                    tx_cable = p.cable_type
                    break
            if tx_cable:
                break
        if not tx_cable:
            for swg in equipment_pkg.switchgears:
                for p in swg.panels:
                    if p.cable_type:
                        tx_cable = p.cable_type
                        break
                if tx_cable:
                    break
    if not tx_cable and "CABLE" in equipment.upper():
        tx_cable = equipment

    tx_number = record.equipment_id or resolved_item_key
    tx_us_reading = (matched_tx.us_reading if matched_tx and matched_tx.us_reading else "") or _text_or_empty(record.us_reading)
    tx_us_char = (matched_tx.us_char if matched_tx and matched_tx.us_char else "") or _text_or_empty(record.us_char)
    tev_bg = _extract_tev_background(pe_info)

    ir_sev = "-" if overview else "__SEVERITY_IR__"
    us_sev = "-" if overview else "__SEVERITY_US__"
    tev_sev = "-" if overview else "__SEVERITY_TEV__"
    def_techs = {record.technology} if record.technology else set()

    # Discover survey root and retrieve/generate PRPD images if available
    us_prpd: Path | str = ""
    tev_prpd: Path | str = ""
    if pe_info and not overview:
        tx_idx = 1
        if matched_tx and matched_tx.tx_id:
            digits = re.findall(r"\d+", matched_tx.tx_id)
            if digits:
                try:
                    tx_idx = int(digits[0])
                except ValueError:
                    tx_idx = 1
        elif equipment:
            digits = re.findall(r"\d+", equipment)
            if digits:
                try:
                    tx_idx = int(digits[0])
                except ValueError:
                    tx_idx = 1

        prpd_catalog = pe_info.get("prpd_catalog")
        if prpd_catalog and isinstance(prpd_catalog, dict) and "tx" in prpd_catalog:
            tx_catalog = prpd_catalog["tx"]
            if tx_idx in tx_catalog:
                entry = tx_catalog[tx_idx]
                us_prpd = entry.get("us") or ""
                tev_prpd = entry.get("tev") or ""
            elif len(tx_catalog) == 1 and 1 in tx_catalog:
                entry = tx_catalog[1]
                us_prpd = entry.get("us") or ""
                tev_prpd = entry.get("tev") or ""

        # Fallback to direct on-demand generation if not in catalog
        if not us_prpd and not tev_prpd:
            raw_dir = pe_info.get("raw_data_dir") or pe_info.get("survey_dir") or pe_info.get("raw_dir")
            survey_root = discover_ultratev_survey_dir(raw_dir)
            prpd_out_dir = pe_info.get("prpd_output_dir")

            if survey_root and prpd_out_dir:
                us_png, tev_png = generate_prpd_graphs_for_transformer(
                    survey_root=survey_root,
                    tx_idx=tx_idx,
                    output_dir=Path(prpd_out_dir),
                )
                if us_png:
                    us_prpd = us_png
                if tev_png:
                    tev_prpd = tev_png

    return {
        "__is_overview__": overview,
        "__defective_technologies__": def_techs,
        "tx": {
            "number": _fallback_dash(tx_number),
            "location": _fallback_dash(tx_location),
            "area": area,
            "manufacturer": _fallback_dash(tx_mfg),
            "model": _fallback_dash(tx_model),
            "rating": _fallback_dash(tx_rating),
            "serialnumber": _fallback_dash(tx_serial),
            "cabletype": _fallback_dash(tx_cable),
            "ir": {
                "reading": format_temperature_float(record.ir_reading),
                "severity": ir_sev,
            },
            "us": {
                "reading": format_db_int(tx_us_reading),
                "char": normalize_us_characteristic(tx_us_char),
                "severity": us_sev,
                "prpd": us_prpd,
            },
            "tev": {
                "reading": format_db_int(record.tev_reading),
                "char": _fallback_dash(record.tev_char),
                "bg": format_db_int(tev_bg),
                "severity": tev_sev,
                "prpd": tev_prpd,
            },
        },
        "ir": {
            "reading": format_temperature_float(record.ir_reading),
            "severity": ir_sev,
        },
        "us": {
            "reading": format_db_int(tx_us_reading),
            "char": normalize_us_characteristic(tx_us_char),
            "severity": us_sev,
            "prpd": us_prpd,
        },
        "tev": {
            "reading": format_db_int(record.tev_reading),
            "char": _fallback_dash(record.tev_char),
            "bg": format_db_int(tev_bg),
            "severity": tev_sev,
            "prpd": tev_prpd,
        },
    }


def _build_blackbox_render_context(
    record: CbmDefectRecord,
    *,
    overview: bool,
    item_key: str = "",
    item_suffix: str = "",
    pe_info: dict[str, Any] | None = None,
) -> dict[str, Any]:
    resolved_item_key = _text_or_empty(item_key)
    resolved_item_suffix = _text_or_empty(item_suffix)
    area = _format_detail_area(record.defect_area, record.additional_remarks, overview=overview)

    raw_target = record.equipment_id or resolved_item_key or resolved_item_suffix
    digits = re.findall(r"\d+", raw_target)
    if digits:
        bbox_number = digits[0]
    elif raw_target.strip():
        bbox_number = raw_target.strip()
    else:
        bbox_number = "-"

    search_text = f"{record.defect_area} {record.additional_remarks} {record.equipment_id} {resolved_item_suffix} {resolved_item_key}".upper()
    if "LEFT" in search_text:
        bbox_location = "LEFT"
    elif "RIGHT" in search_text:
        bbox_location = "RIGHT"
    elif "FRONT" in search_text:
        bbox_location = "FRONT"
    elif "REAR" in search_text:
        bbox_location = "REAR"
    else:
        bbox_location = "-"

    tev_bg = _extract_tev_background(pe_info)
    ir_sev = "-" if overview else "__SEVERITY_IR__"
    us_sev = "-" if overview else "__SEVERITY_US__"
    tev_sev = "-" if overview else "__SEVERITY_TEV__"
    def_techs = {record.technology} if record.technology else set()

    return {
        "__is_overview__": overview,
        "__defective_technologies__": def_techs,
        "bbox": {
            "number": _fallback_dash(bbox_number),
            "location": _fallback_dash(bbox_location),
            "area": area,
            "ir": {
                "reading": format_temperature_float(record.ir_reading),
                "severity": ir_sev,
            },
            "us": {
                "reading": format_db_int(record.us_reading),
                "char": normalize_us_characteristic(record.us_char),
                "severity": us_sev,
                "prpd": "",
            },
            "tev": {
                "reading": format_db_int(record.tev_reading),
                "char": _fallback_dash(record.tev_char),
                "bg": format_db_int(tev_bg),
                "severity": tev_sev,
                "prpd": "",
            },
        },
        "ir": {
            "reading": format_temperature_float(record.ir_reading),
            "severity": ir_sev,
        },
        "us": {
            "reading": format_db_int(record.us_reading),
            "char": normalize_us_characteristic(record.us_char),
            "severity": us_sev,
            "prpd": "",
        },
        "tev": {
            "reading": format_db_int(record.tev_reading),
            "char": _fallback_dash(record.tev_char),
            "bg": format_db_int(tev_bg),
            "severity": tev_sev,
            "prpd": "",
        },
    }


def _build_battery_render_context(
    record: CbmDefectRecord,
    *,
    overview: bool,
    item_key: str = "",
    item_suffix: str = "",
    pe_info: dict[str, Any] | None = None,
) -> dict[str, Any]:
    resolved_item_key = _text_or_empty(item_key)
    resolved_item_suffix = _text_or_empty(item_suffix)
    area = _format_detail_area(record.defect_area, record.additional_remarks, overview=overview)

    raw_target = record.equipment_id or resolved_item_key or resolved_item_suffix

    equipment_pkg = _extract_equipment_package(pe_info)
    battery_banks = equipment_pkg.battery_banks if equipment_pkg else ()
    matched_batt = _find_matching_battery(battery_banks, raw_target)

    batt_mfg = _text_or_empty(record.brand) or (matched_batt.manufacturer if matched_batt else "")
    batt_model = _text_or_empty(record.model) or (matched_batt.model if matched_batt else "")
    batt_serial = matched_batt.serial_no if matched_batt else ""
    batt_number = raw_target or (matched_batt.name if matched_batt else "")

    tev_bg = _extract_tev_background(pe_info)
    ir_sev = "-" if overview else "__SEVERITY_IR__"
    us_sev = "-" if overview else "__SEVERITY_US__"
    tev_sev = "-" if overview else "__SEVERITY_TEV__"
    def_techs = {record.technology} if record.technology else set()

    return {
        "__is_overview__": overview,
        "__defective_technologies__": def_techs,
        "batt": {
            "number": _fallback_dash(batt_number),
            "manufacturer": _fallback_dash(batt_mfg),
            "model": _fallback_dash(batt_model),
            "serialnumber": _fallback_dash(batt_serial),
            "area": area,
            "ir": {
                "reading": format_temperature_float(record.ir_reading),
                "severity": ir_sev,
            },
            "us": {
                "reading": format_db_int(record.us_reading),
                "char": normalize_us_characteristic(record.us_char),
                "severity": us_sev,
                "prpd": "",
            },
            "tev": {
                "reading": format_db_int(record.tev_reading),
                "char": _fallback_dash(record.tev_char),
                "bg": format_db_int(tev_bg),
                "severity": tev_sev,
                "prpd": "",
            },
        },
        "ir": {
            "reading": format_temperature_float(record.ir_reading),
            "severity": ir_sev,
        },
        "us": {
            "reading": format_db_int(record.us_reading),
            "char": normalize_us_characteristic(record.us_char),
            "severity": us_sev,
            "prpd": "",
        },
        "tev": {
            "reading": format_db_int(record.tev_reading),
            "char": _fallback_dash(record.tev_char),
            "bg": format_db_int(tev_bg),
            "severity": tev_sev,
            "prpd": "",
        },
    }


def _build_family_render_context(
    family_spec: QuickReportFamilySpec,
    record: CbmDefectRecord,
    *,
    overview: bool,
    item_key: str = "",
    item_suffix: str = "",
    pe_info: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if family_spec.id == "fp_lvdb":
        return _build_fp_lvdb_render_context(
            record,
            overview=overview,
            item_key=item_key,
            item_suffix=item_suffix,
            pe_info=pe_info,
        )
    if family_spec.id == "swg":
        return _build_swg_render_context(
            record,
            overview=overview,
            item_key=item_key,
            item_suffix=item_suffix,
            pe_info=pe_info,
        )
    if family_spec.id == "tx":
        return _build_tx_render_context(
            record,
            overview=overview,
            item_key=item_key,
            item_suffix=item_suffix,
            pe_info=pe_info,
        )
    if family_spec.id == "blackbox":
        return _build_blackbox_render_context(
            record,
            overview=overview,
            item_key=item_key,
            item_suffix=item_suffix,
            pe_info=pe_info,
        )
    if family_spec.id == "battery":
        return _build_battery_render_context(
            record,
            overview=overview,
            item_key=item_key,
            item_suffix=item_suffix,
            pe_info=pe_info,
        )
    return {}


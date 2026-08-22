"""Integration and layout tests for Substation Condition DOCX generation (Stage 5 Loader)."""

import logging
from pathlib import Path
import pytest
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

from src.quick_report.substation_condition import (
    _remove_empty_cell_borders_sub_cond,
    generate_substation_condition_pages,
)

TEMPLATE_PATH = Path("templates/QUICK REPORT/SUBSTATION CONFIGURATION/MASTER_SUBSTATION_CONDITION.docx")


def test_generate_substation_condition_pages_single_page(tmp_path: Path):
    """Verify single page (3 pairs) generation and paragraph shrinking."""
    pe_info = {
        "substation": {
            "substation_name_site": "PE TEST 1",
            "name_site": "PE TEST 1",
        }
    }
    pairs = [
        ("SUBSTATION OVERVIEW", "SIGNBOARD"),
        ("SWITCHGEAR", "SWITCHGEAR NAMEPLATE"),
        ("TRANSFORMER", "TRANSFORMER NAMEPLATE"),
    ]

    out_paths = generate_substation_condition_pages(
        pe_info=pe_info,
        condition_pairs_or_pkg=pairs,
        template_path=TEMPLATE_PATH,
        output_dir=tmp_path,
        substation_number=1,
    )

    assert len(out_paths) == 1
    doc_path = out_paths[0]
    assert doc_path.exists()

    doc = Document(doc_path)
    assert len(doc.tables) == 1
    table = doc.tables[0]

    # Verify slot 0, 1, 2 headers
    assert "SUBSTATION OVERVIEW" in table.rows[0].cells[0].text
    assert "SIGNBOARD" in table.rows[0].cells[2].text
    assert "SWITCHGEAR" in table.rows[3].cells[0].text
    assert "SWITCHGEAR NAMEPLATE" in table.rows[3].cells[2].text
    assert "TRANSFORMER" in table.rows[6].cells[0].text
    assert "TRANSFORMER NAMEPLATE" in table.rows[6].cells[2].text

    # Verify trailing paragraph formatting
    assert len(doc.paragraphs) > 0
    last_para = doc.paragraphs[-1]
    assert last_para.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert last_para.paragraph_format.line_spacing == Pt(0.5)
    assert last_para.paragraph_format.space_before == Pt(0)
    assert last_para.paragraph_format.space_after == Pt(0)


def test_generate_substation_condition_pages_multi_page_and_half_pair(tmp_path: Path):
    """Verify 7 pairs -> 3 chunks merged into 1 docx with clean border stripping on incomplete page & half-pair."""
    pe_info = {
        "substation": {
            "substation_name_site": "PE TEST 2",
            "name_site": "PE TEST 2",
        }
    }
    pairs = [
        ("SUBSTATION OVERVIEW", "SIGNBOARD"),
        ("SWITCHGEAR", "SWITCHGEAR NAMEPLATE"),
        ("TRANSFORMER", "TRANSFORMER NAMEPLATE"),
        ("FEEDER PILLAR", "FEEDER PILLAR NAMEPLATE"),
        ("FIRE EXTINGUISHER", "FIRE EXTINGUISHER EXPIRY DATE"),
        ("EFI", "SF6 INDICATOR"),
        ("TRANSFORMER OIL LEVEL INDICATOR", ""),
    ]

    out_paths = generate_substation_condition_pages(
        pe_info=pe_info,
        condition_pairs_or_pkg=pairs,
        template_path=TEMPLATE_PATH,
        output_dir=tmp_path,
        substation_number=2,
    )

    assert len(out_paths) == 1
    merged_path = out_paths[0]
    assert merged_path.exists()
    assert merged_path.name == "002_5 SUBSTATION CONDITION.docx"

    # Verify part files were cleaned up
    part1 = tmp_path / "002_5 SUBSTATION CONDITION part1.docx"
    part2 = tmp_path / "002_5 SUBSTATION CONDITION part2.docx"
    part3 = tmp_path / "002_5 SUBSTATION CONDITION part3.docx"
    assert not part1.exists()
    assert not part2.exists()
    assert not part3.exists()

    doc = Document(merged_path)
    # The merged document contains 3 tables (1 from each part)
    assert len(doc.tables) >= 3

    # Table 2 corresponds to Page 3 (Part 3): 1 half-pair (slot 0), slots 1 & 2 unused
    table3 = doc.tables[2]
    # Slot 0 Left (row 0, cell 0)
    assert "TRANSFORMER OIL LEVEL INDICATOR" in table3.rows[0].cells[0].text
    # Slot 0 Right (row 0 & 1, cell 2) cleared and borders removed
    assert table3.rows[0].cells[2].text == ""
    assert table3.rows[1].cells[2].text == ""
    tcPr_r0_c2 = table3.rows[0].cells[2]._tc.get_or_add_tcPr()
    tcBorders = tcPr_r0_c2.find(qn("w:tcBorders"))
    assert tcBorders is not None
    assert tcBorders.find(qn("w:top")).get(qn("w:val")) == "nil"

    # Slot 1 (rows 3 & 4) cleared and borders removed
    assert table3.rows[3].cells[0].text == ""
    assert table3.rows[3].cells[2].text == ""
    tcPr_s1 = table3.rows[3].cells[0]._tc.get_or_add_tcPr()
    assert tcPr_s1.find(qn("w:tcBorders")) is not None

    # Slot 2 (rows 6 & 7) cleared and borders removed
    assert table3.rows[6].cells[0].text == ""
    assert table3.rows[6].cells[2].text == ""
    tcPr_s2 = table3.rows[6].cells[0]._tc.get_or_add_tcPr()
    assert tcPr_s2.find(qn("w:tcBorders")) is not None


def test_remove_empty_cell_borders_sub_cond_half_pair_direct(tmp_path: Path):
    """Direct test of _remove_empty_cell_borders_sub_cond with chunk containing half-pair in 8-row table."""
    from docxtpl import DocxTemplate

    doc = DocxTemplate(str(TEMPLATE_PATH))
    context = {
        "pairs": [
            {"header_left": "LEFT 1", "header_right": "", "photo_left": "", "photo_right": ""},
            {"header_left": "", "header_right": "", "photo_left": "", "photo_right": ""},
            {"header_left": "", "header_right": "", "photo_left": "", "photo_right": ""},
        ]
    }
    doc.render(context)
    test_docx = tmp_path / "test_half_pair.docx"
    doc.save(test_docx)

    chunk = [("LEFT 1", "")]
    _remove_empty_cell_borders_sub_cond(test_docx, chunk)

    updated = Document(test_docx)
    table = updated.tables[0]

    # Slot 0 left is active
    assert "LEFT 1" in table.rows[0].cells[0].text
    # Slot 0 right has borders stripped
    tcPr_right = table.rows[0].cells[2]._tc.get_or_add_tcPr()
    assert tcPr_right.find(qn("w:tcBorders")) is not None
    # Slot 1 left has borders stripped
    tcPr_s1_left = table.rows[3].cells[0]._tc.get_or_add_tcPr()
    assert tcPr_s1_left.find(qn("w:tcBorders")) is not None


def test_generate_substation_condition_pages_missing_template_raises(tmp_path: Path):
    """Verify FileNotFoundError is raised when template does not exist."""
    with pytest.raises(FileNotFoundError):
        generate_substation_condition_pages(
            pe_info={},
            condition_pairs_or_pkg=[],
            template_path=tmp_path / "non_existent_template.docx",
            output_dir=tmp_path,
            substation_number=1,
        )


def test_remove_empty_cell_borders_sub_cond_error_logged(tmp_path: Path, caplog):
    """Verify logging warning on corrupted or missing path."""
    with caplog.at_level(logging.WARNING):
        _remove_empty_cell_borders_sub_cond(tmp_path / "invalid_path.docx", 1)
    assert "Failed to remove empty cell borders" in caplog.text

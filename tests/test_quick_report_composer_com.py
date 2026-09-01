from pathlib import Path
from unittest.mock import MagicMock, patch
import pytest

from src.quick_report.composer import QuickReportComposer, _paste_with_retry


def test_compile_document_com_recopy_paste(tmp_path: Path):
    """Verify _compile_document calls Add(), Open(), Copy(), Paste(), SaveAs2() when compiled."""
    p1 = tmp_path / "part1.docx"
    p2 = tmp_path / "part2.docx"
    out = tmp_path / "out.docx"

    p1.touch()
    p2.touch()

    composer = QuickReportComposer()

    mock_word = MagicMock()
    mock_main_doc = MagicMock()
    mock_part_doc1 = MagicMock()
    mock_part_doc2 = MagicMock()
    mock_rng = MagicMock()

    mock_word.Documents.Add.return_value = mock_main_doc
    mock_word.Documents.Open.side_effect = [mock_part_doc1, mock_part_doc2]
    mock_main_doc.Content = mock_rng
    mock_main_doc.Tables.Count = 0
    mock_rng.Information.return_value = False

    composer._compile_document([p1, p2], out, word_app=mock_word)

    mock_word.Documents.Add.assert_called_once()
    assert mock_word.Documents.Open.call_count == 2
    mock_part_doc1.Content.Copy.assert_called_once()
    mock_part_doc1.Close.assert_called_once_with(False)
    mock_part_doc2.Content.Copy.assert_called_once()
    mock_part_doc2.Close.assert_called_once_with(False)

    mock_rng.InsertBreak.assert_called_once_with(7)
    assert mock_rng.Paste.call_count == 2
    mock_main_doc.SaveAs2.assert_called_once_with(str(out.resolve()))
    mock_main_doc.Close.assert_called_once_with(False)


def test_compile_document_with_external_word_app(tmp_path: Path):
    """Verify _compile_document reuses provided word_app without quitting it."""
    p1 = tmp_path / "part1.docx"
    out = tmp_path / "out.docx"
    p1.touch()

    composer = QuickReportComposer()

    mock_word = MagicMock()
    mock_main_doc = MagicMock()
    mock_part_doc = MagicMock()
    mock_rng = MagicMock()

    mock_word.Documents.Add.return_value = mock_main_doc
    mock_word.Documents.Open.return_value = mock_part_doc
    mock_main_doc.Content = mock_rng
    mock_main_doc.Tables.Count = 0
    mock_rng.Information.return_value = False

    composer._compile_document([p1], out, word_app=mock_word)

    mock_word.Documents.Add.assert_called_once()
    mock_part_doc.Content.Copy.assert_called_once()
    mock_part_doc.Close.assert_called_once_with(False)
    mock_rng.Paste.assert_called_once()
    mock_main_doc.SaveAs2.assert_called_once_with(str(out.resolve()))
    mock_main_doc.Close.assert_called_once_with(False)


def test_compile_document_raises_when_win32com_missing(tmp_path: Path):
    """Verify _compile_document raises RuntimeError when word_app is None."""
    composer = QuickReportComposer()
    p1 = tmp_path / "part1.docx"
    out = tmp_path / "out.docx"

    with pytest.raises(RuntimeError, match="word_app is required for Quick Report compilation."):
        composer._compile_document([p1], out, word_app=None)


def test_paste_with_retry_success_first_attempt():
    """Verify _paste_with_retry succeeds on first attempt without retrying."""
    mock_rng = MagicMock()
    _paste_with_retry(mock_rng, max_attempts=3, delay=0.01)
    mock_rng.Paste.assert_called_once()


def test_paste_with_retry_success_after_retries():
    """Verify _paste_with_retry retries rng.Paste() until success."""
    mock_rng = MagicMock()
    mock_rng.Paste.side_effect = [Exception("Clipboard locked"), Exception("COM error"), None]
    _paste_with_retry(mock_rng, max_attempts=5, delay=0.01)
    assert mock_rng.Paste.call_count == 3


def test_paste_with_retry_fails_and_reraises():
    """Verify _paste_with_retry re-raises exception when all retry attempts fail."""
    mock_rng = MagicMock()
    mock_rng.Paste.side_effect = Exception("Persistent COM failure")
    with pytest.raises(Exception, match="Persistent COM failure"):
        _paste_with_retry(mock_rng, max_attempts=3, delay=0.01)
    assert mock_rng.Paste.call_count == 3


def test_compile_document_escapes_table_cell(tmp_path: Path):
    """Verify _compile_document escapes table cells when rng is inside a table."""
    p1 = tmp_path / "part1.docx"
    p2 = tmp_path / "part2.docx"
    out = tmp_path / "out.docx"
    p1.touch()
    p2.touch()

    composer = QuickReportComposer()

    mock_word = MagicMock()
    mock_main_doc = MagicMock()
    mock_part_doc = MagicMock()
    mock_rng = MagicMock()
    mock_table = MagicMock()

    mock_word.Documents.Add.return_value = mock_main_doc
    mock_word.Documents.Open.return_value = mock_part_doc
    mock_main_doc.Content = mock_rng
    mock_main_doc.Tables.Count = 1
    mock_main_doc.Tables.return_value = mock_table
    mock_rng.Information.return_value = True  # wdWithInTable = True

    composer._compile_document([p1, p2], out, word_app=mock_word)

    # Verify InsertParagraphAfter was called to escape table cell
    assert mock_table.Range.InsertParagraphAfter.call_count >= 2
    mock_main_doc.Tables.assert_called_with(1)


def _create_mock_template(path: Path, is_table: bool = False, rows: int = 1, cols: int = 1) -> Path:
    from docx import Document

    doc = Document()
    if is_table:
        table = doc.add_table(rows=rows, cols=cols)
        for r in table.rows:
            for c in r.cells:
                c.text = "cell"
    else:
        doc.add_paragraph("Mock Content")
    doc.save(path)
    return path


def _setup_mock_templates(tpl_dir: Path) -> dict[str, Path]:
    tpl_dir.mkdir(parents=True, exist_ok=True)
    return {
        "front_page": _create_mock_template(tpl_dir / "front_page.docx"),
        "cbm_summary": _create_mock_template(tpl_dir / "cbm_summary.docx", is_table=True, rows=2, cols=7),
        "vi_summary": _create_mock_template(tpl_dir / "vi_summary.docx", is_table=True, rows=2, cols=4),
        "swg_overview": _create_mock_template(tpl_dir / "swg_overview.docx"),
        "swg_panel": _create_mock_template(tpl_dir / "swg_panel.docx"),
        "tx_overview": _create_mock_template(tpl_dir / "tx_overview.docx"),
        "tx_detail": _create_mock_template(tpl_dir / "tx_detail.docx"),
        "substation_condition": _create_mock_template(tpl_dir / "substation_condition.docx", is_table=True, rows=8, cols=3),
        "vi_defect": _create_mock_template(tpl_dir / "vi_defect.docx", is_table=True, rows=12, cols=3),
        "sticker": _create_mock_template(tpl_dir / "sticker.docx"),
    }


def test_composer_generate_parts_full_report_all_7_parts(tmp_path: Path):
    """Verify QuickReportComposer generates all 7 parts with strict 2-digit prefixes in exact order for full CBM+VI report."""
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.defects import CbmDefectRecord, ViDefectRecord
    from src.quick_report.models import (
        CbmDefectDetailGroup,
        CbmDefectFamilyPlan,
        CbmDefectGroup,
        QuickReportStationPlan,
    )
    from src.testsheet.models import SubstationTestsheetPackage

    tpls = _setup_mock_templates(tmp_path / "templates")
    out_dir = tmp_path / "output"
    temp_dir = out_dir / "temp_parts"
    temp_dir.mkdir(parents=True, exist_ok=True)

    cbm_defects = (
        CbmDefectRecord(equipment="RMU SF6", defect_area="Cable Box", technology="IR", raw_measurement="60.5"),
        CbmDefectRecord(equipment="RMU SF6", defect_area="Busbar", technology="US", raw_measurement="15.0", us_char="CORONA"),
        CbmDefectRecord(equipment="RMU SF6", defect_area="Cable Box", technology="TEV", raw_measurement="25.0"),
    )
    vi_defects = (
        ViDefectRecord(equipment="SWITCHGEAR", defect_area="Door defect", additional_remarks="Broken handle"),
    )

    spec_swg = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    swg_group = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="",
        defects=cbm_defects,
        overview=cbm_defects[0],
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(cbm_defects[0],)),
        ),
    )
    swg_plan = CbmDefectFamilyPlan(
        spec=spec_swg,
        overview_template=tpls["swg_overview"],
        detail_templates=(("panel_area", tpls["swg_panel"]),),
        groups=(swg_group,),
    )

    pkg = SubstationTestsheetPackage(
        testsheet_path=Path("dummy.xlsx"),
        unsorted_raw_data_dir=Path("dummy_raw"),
        station="TEST STATION",
        month="08. AUGUST",
        date_str="01-09-2026",
        substation_number=1,
        data=None,
    )

    pe_info = {
        "substation": {
            "name_erms": "PE FULL TEST",
            "name_site": "PE FULL TEST",
            "datefrontpage": "01 SEP 2026",
            "date": "01/09/2026",
        }
    }

    plan = QuickReportStationPlan(
        package=pkg,
        pe_info=pe_info,
        cbm_defects=cbm_defects,
        vi_defects=vi_defects,
        suffix=" (IR+US+TEV+VI)",
        suffix_parts=("IR", "US", "TEV", "VI"),
        output_dir=out_dir,
        output_filename="001. PE FULL TEST (IR+US+TEV+VI).docx",
        final_output_path=out_dir / "001. PE FULL TEST (IR+US+TEV+VI).docx",
        condition_pairs=(("SUBSTATION OVERVIEW", "SIGNBOARD"),),
        cond_template_path=tpls["substation_condition"],
        front_page_template=tpls["front_page"],
        cbm_summary_template=tpls["cbm_summary"],
        vi_summary_template=tpls["vi_summary"],
        vi_defect_template=tpls["vi_defect"],
        sticker_template=tpls["sticker"],
        cbm_defect_family_plans=(swg_plan,),
    )

    composer = QuickReportComposer()
    parts = composer._generate_parts(plan, temp_dir)

    expected_filenames = [
        "001_01_front_page.docx",
        "001_02_cbm_summary.docx",
        "001_03_vi_summary.docx",
        "001_04_SWG_OVERVIEW.docx",
        "001_04_SWG_RMU SF6.docx",
        "001_05_substation_condition.docx",
        "001_06_vi_defect_part1.docx",
        "001_07_sticker_page.docx",
    ]

    actual_filenames = [p.name for p in parts]
    assert actual_filenames == expected_filenames

    # Verify all generated parts physically exist on disk
    for part in parts:
        assert part.exists(), f"Part file does not exist: {part}"


def test_composer_generate_parts_pure_vi_report(tmp_path: Path):
    """Verify QuickReportComposer generates parts 1, 3, 5, 6, 7 in exact order for pure VI report."""
    from src.quick_report.defects import ViDefectRecord
    from src.quick_report.models import QuickReportStationPlan
    from src.testsheet.models import SubstationTestsheetPackage

    tpls = _setup_mock_templates(tmp_path / "templates")
    out_dir = tmp_path / "output"
    temp_dir = out_dir / "temp_parts"
    temp_dir.mkdir(parents=True, exist_ok=True)

    vi_defects = (
        ViDefectRecord(equipment="TRANSFORMER", defect_area="Oil leak", additional_remarks="Top up oil"),
    )

    pkg = SubstationTestsheetPackage(
        testsheet_path=Path("dummy.xlsx"),
        unsorted_raw_data_dir=Path("dummy_raw"),
        station="TEST STATION",
        month="08. AUGUST",
        date_str="01-09-2026",
        substation_number=2,
        data=None,
    )

    pe_info = {
        "substation": {
            "name_erms": "PE PURE VI",
            "name_site": "PE PURE VI",
            "datefrontpage": "01 SEP 2026",
            "date": "01/09/2026",
        }
    }

    plan = QuickReportStationPlan(
        package=pkg,
        pe_info=pe_info,
        cbm_defects=(),
        vi_defects=vi_defects,
        suffix=" (VI)",
        suffix_parts=("VI",),
        output_dir=out_dir,
        output_filename="002. PE PURE VI (VI).docx",
        final_output_path=out_dir / "002. PE PURE VI (VI).docx",
        condition_pairs=(("SUBSTATION OVERVIEW", "SIGNBOARD"),),
        cond_template_path=tpls["substation_condition"],
        front_page_template=tpls["front_page"],
        cbm_summary_template=None,
        vi_summary_template=tpls["vi_summary"],
        vi_defect_template=tpls["vi_defect"],
        sticker_template=tpls["sticker"],
        cbm_defect_family_plans=(),
    )

    composer = QuickReportComposer()
    parts = composer._generate_parts(plan, temp_dir)

    expected_filenames = [
        "002_01_front_page.docx",
        "002_03_vi_summary.docx",
        "002_05_substation_condition.docx",
        "002_06_vi_defect_part1.docx",
        "002_07_sticker_page.docx",
    ]

    actual_filenames = [p.name for p in parts]
    assert actual_filenames == expected_filenames

    # Ensure Part 2 and Part 4 are strictly omitted
    assert not any("_02_" in name for name in actual_filenames)
    assert not any("_04_" in name for name in actual_filenames)

    for part in parts:
        assert part.exists()


def test_composer_generate_parts_pure_cbm_report(tmp_path: Path):
    """Verify QuickReportComposer generates parts 1, 2, 4, 5, 7 in exact order for pure CBM report."""
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.defects import CbmDefectRecord
    from src.quick_report.models import (
        CbmDefectDetailGroup,
        CbmDefectFamilyPlan,
        CbmDefectGroup,
        QuickReportStationPlan,
    )
    from src.testsheet.models import SubstationTestsheetPackage

    tpls = _setup_mock_templates(tmp_path / "templates")
    out_dir = tmp_path / "output"
    temp_dir = out_dir / "temp_parts"
    temp_dir.mkdir(parents=True, exist_ok=True)

    cbm_defects = (
        CbmDefectRecord(equipment="TX 1", defect_area="HV Bushing", technology="IR", raw_measurement="55.0"),
    )

    spec_tx = QUICK_REPORT_FAMILY_SPECS_BY_ID["tx"]
    tx_group = CbmDefectGroup(
        item_key="TX 1",
        item_suffix="",
        defects=cbm_defects,
        overview=cbm_defects[0],
        detail_groups=(
            CbmDefectDetailGroup(role_id="tx_hv_side", defects=cbm_defects),
        ),
    )
    tx_plan = CbmDefectFamilyPlan(
        spec=spec_tx,
        overview_template=tpls["tx_overview"],
        detail_templates=(("tx_hv_side", tpls["tx_detail"]),),
        groups=(tx_group,),
    )

    pkg = SubstationTestsheetPackage(
        testsheet_path=Path("dummy.xlsx"),
        unsorted_raw_data_dir=Path("dummy_raw"),
        station="TEST STATION",
        month="08. AUGUST",
        date_str="01-09-2026",
        substation_number=3,
        data=None,
    )

    pe_info = {
        "substation": {
            "name_erms": "PE PURE CBM",
            "name_site": "PE PURE CBM",
            "datefrontpage": "01 SEP 2026",
            "date": "01/09/2026",
        }
    }

    plan = QuickReportStationPlan(
        package=pkg,
        pe_info=pe_info,
        cbm_defects=cbm_defects,
        vi_defects=(),
        suffix=" (IR)",
        suffix_parts=("IR",),
        output_dir=out_dir,
        output_filename="003. PE PURE CBM (IR).docx",
        final_output_path=out_dir / "003. PE PURE CBM (IR).docx",
        condition_pairs=(("SUBSTATION OVERVIEW", "SIGNBOARD"),),
        cond_template_path=tpls["substation_condition"],
        front_page_template=tpls["front_page"],
        cbm_summary_template=tpls["cbm_summary"],
        vi_summary_template=None,
        vi_defect_template=None,
        sticker_template=tpls["sticker"],
        cbm_defect_family_plans=(tx_plan,),
    )

    composer = QuickReportComposer()
    parts = composer._generate_parts(plan, temp_dir)

    expected_filenames = [
        "003_01_front_page.docx",
        "003_02_cbm_summary.docx",
        "003_04_TX_OVERVIEW.docx",
        "003_04_TX_TX 1.docx",
        "003_05_substation_condition.docx",
        "003_07_sticker_page.docx",
    ]

    actual_filenames = [p.name for p in parts]
    assert actual_filenames == expected_filenames

    # Ensure Part 3 and Part 6 are strictly omitted
    assert not any("_03_" in name for name in actual_filenames)
    assert not any("_06_" in name for name in actual_filenames)

    for part in parts:
        assert part.exists()


def test_composer_generate_parts_clean_report_without_defects(tmp_path: Path):
    """Verify QuickReportComposer generates parts 1, 5, 7 for defect-free clean report."""
    from src.quick_report.models import QuickReportStationPlan
    from src.testsheet.models import SubstationTestsheetPackage

    tpls = _setup_mock_templates(tmp_path / "templates")
    out_dir = tmp_path / "output"
    temp_dir = out_dir / "temp_parts"
    temp_dir.mkdir(parents=True, exist_ok=True)

    pkg = SubstationTestsheetPackage(
        testsheet_path=Path("dummy.xlsx"),
        unsorted_raw_data_dir=Path("dummy_raw"),
        station="TEST STATION",
        month="08. AUGUST",
        date_str="01-09-2026",
        substation_number=4,
        data=None,
    )

    pe_info = {
        "substation": {
            "name_erms": "PE CLEAN",
            "name_site": "PE CLEAN",
            "datefrontpage": "01 SEP 2026",
            "date": "01/09/2026",
        }
    }

    plan = QuickReportStationPlan(
        package=pkg,
        pe_info=pe_info,
        cbm_defects=(),
        vi_defects=(),
        suffix="",
        suffix_parts=(),
        output_dir=out_dir,
        output_filename="004. PE CLEAN.docx",
        final_output_path=out_dir / "004. PE CLEAN.docx",
        condition_pairs=(("SUBSTATION OVERVIEW", "SIGNBOARD"),),
        cond_template_path=tpls["substation_condition"],
        front_page_template=tpls["front_page"],
        cbm_summary_template=None,
        vi_summary_template=None,
        vi_defect_template=None,
        sticker_template=tpls["sticker"],
        cbm_defect_family_plans=(),
    )

    composer = QuickReportComposer()
    parts = composer._generate_parts(plan, temp_dir)

    expected_filenames = [
        "004_01_front_page.docx",
        "004_05_substation_condition.docx",
        "004_07_sticker_page.docx",
    ]

    actual_filenames = [p.name for p in parts]
    assert actual_filenames == expected_filenames

    # Ensure Parts 2, 3, 4, 6 are strictly omitted
    assert not any("_02_" in name for name in actual_filenames)
    assert not any("_03_" in name for name in actual_filenames)
    assert not any("_04_" in name for name in actual_filenames)
    assert not any("_06_" in name for name in actual_filenames)

    for part in parts:
        assert part.exists()


def test_composer_load_end_to_end_lifecycle_and_cleanup(tmp_path: Path):
    """Verify QuickReportComposer.load renders all parts, invokes compilation, and cleans up temp_parts."""
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.defects import CbmDefectRecord, ViDefectRecord
    from src.quick_report.models import (
        CbmDefectDetailGroup,
        CbmDefectFamilyPlan,
        CbmDefectGroup,
        QuickReportStationPlan,
    )
    from src.testsheet.models import SubstationTestsheetPackage

    tpls = _setup_mock_templates(tmp_path / "templates")
    out_dir = tmp_path / "output"

    cbm_defects = (
        CbmDefectRecord(equipment="RMU SF6", defect_area="Cable Box", technology="IR", raw_measurement="60.5"),
    )
    vi_defects = (
        ViDefectRecord(equipment="SWITCHGEAR", defect_area="Door defect", additional_remarks="Broken handle"),
    )

    spec_swg = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    swg_group = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="",
        defects=cbm_defects,
        overview=cbm_defects[0],
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(cbm_defects[0],)),
        ),
    )
    swg_plan = CbmDefectFamilyPlan(
        spec=spec_swg,
        overview_template=tpls["swg_overview"],
        detail_templates=(("panel_area", tpls["swg_panel"]),),
        groups=(swg_group,),
    )

    pkg = SubstationTestsheetPackage(
        testsheet_path=Path("dummy.xlsx"),
        unsorted_raw_data_dir=Path("dummy_raw"),
        station="TEST STATION",
        month="08. AUGUST",
        date_str="01-09-2026",
        substation_number=1,
        data=None,
    )

    pe_info = {
        "substation": {
            "name_erms": "PE TEST",
            "name_site": "PE TEST",
            "datefrontpage": "01 SEP 2026",
            "date": "01/09/2026",
        }
    }

    final_output = out_dir / "001. PE TEST (IR+VI).docx"

    plan = QuickReportStationPlan(
        package=pkg,
        pe_info=pe_info,
        cbm_defects=cbm_defects,
        vi_defects=vi_defects,
        suffix=" (IR+VI)",
        suffix_parts=("IR", "VI"),
        output_dir=out_dir,
        output_filename="001. PE TEST (IR+VI).docx",
        final_output_path=final_output,
        condition_pairs=(("SUBSTATION OVERVIEW", "SIGNBOARD"),),
        cond_template_path=tpls["substation_condition"],
        front_page_template=tpls["front_page"],
        cbm_summary_template=tpls["cbm_summary"],
        vi_summary_template=tpls["vi_summary"],
        vi_defect_template=tpls["vi_defect"],
        sticker_template=tpls["sticker"],
        cbm_defect_family_plans=(swg_plan,),
    )

    mock_word = MagicMock()
    mock_main_doc = MagicMock()
    mock_part_doc = MagicMock()
    mock_rng = MagicMock()

    mock_word.Documents.Add.return_value = mock_main_doc
    mock_word.Documents.Open.return_value = mock_part_doc
    mock_main_doc.Content = mock_rng
    mock_main_doc.Tables.Count = 0
    mock_rng.Information.return_value = False

    composer = QuickReportComposer()
    result_path = composer.load(plan, word_app=mock_word)

    assert result_path == final_output
    # 7 parts generated: front, cbm_sum, vi_sum, swg_ov, swg_det, cond, vi_def, sticker (8 files)
    assert mock_word.Documents.Open.call_count == 8
    mock_main_doc.SaveAs2.assert_called_once_with(str(final_output.resolve()))

    # Verify temp_parts was completely cleaned up
    temp_dir = out_dir / "temp_parts"
    assert not temp_dir.exists()


def test_composer_load_cleans_temp_dir_on_compilation_error(tmp_path: Path):
    """Verify QuickReportComposer.load cleans up temp_parts even when an exception occurs during compilation."""
    from src.quick_report.models import QuickReportStationPlan
    from src.testsheet.models import SubstationTestsheetPackage

    tpls = _setup_mock_templates(tmp_path / "templates")
    out_dir = tmp_path / "output_err"

    pkg = SubstationTestsheetPackage(
        testsheet_path=Path("dummy.xlsx"),
        unsorted_raw_data_dir=Path("dummy_raw"),
        station="TEST STATION",
        month="08. AUGUST",
        date_str="01-09-2026",
        substation_number=1,
        data=None,
    )

    plan = QuickReportStationPlan(
        package=pkg,
        pe_info={"substation": {"name_erms": "PE ERR", "name_site": "PE ERR"}},
        cbm_defects=(),
        vi_defects=(),
        suffix="",
        suffix_parts=(),
        output_dir=out_dir,
        output_filename="001. PE ERR.docx",
        final_output_path=out_dir / "001. PE ERR.docx",
        condition_pairs=(),
        cond_template_path=None,
        front_page_template=tpls["front_page"],
        cbm_summary_template=None,
        vi_summary_template=None,
        vi_defect_template=None,
        sticker_template=tpls["sticker"],
        cbm_defect_family_plans=(),
    )

    mock_word = MagicMock()
    mock_word.Documents.Add.side_effect = RuntimeError("COM initialization failure")

    composer = QuickReportComposer()
    with pytest.raises(RuntimeError, match="COM initialization failure"):
        composer.load(plan, word_app=mock_word)

    # Verify temp_parts is still cleaned up
    temp_dir = out_dir / "temp_parts"
    assert not temp_dir.exists()


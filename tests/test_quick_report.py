"""Tests for the Quick Report Generation engine."""

from __future__ import annotations

import logging
from pathlib import Path
from unittest.mock import MagicMock, Mock, patch

import pytest

from src.project.environment import ProjectEnvironment
from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
from src.quick_report.cbm_summary import prepare_tech_summary_rows
from src.quick_report.composer import QuickReportComposer
from src.quick_report.defects import CbmDefectRecord, MasterQr03DefectRepository, ViDefectRecord
from src.quick_report.extractor import QuickReportExtractor
from src.quick_report.models import QuickReportStationPlan
from src.quick_report.transformer import QuickReportTransformer
from src.quick_report.utils import normalize_functional_location_input
from src.workflows.models import QuickReportMode, QuickReportRequest, QuickReportResult
from src.workflows.quick_report import QuickReportWorkflow


def test_fl_normalization():
    """Verify functional location normalization."""
    assert normalize_functional_location_input(" F/L 12345 ") == "12345"
    assert normalize_functional_location_input("f/l 12345") == "12345"
    assert normalize_functional_location_input("12345") == "12345"
    assert normalize_functional_location_input("  F/L   AB-CD  ") == "AB-CD"


def test_quick_report_result_enrichment():
    """Verify QuickReportResult fields."""
    result = QuickReportResult(
        reports_generated=2,
        generated_paths=[Path("a.docx"), Path("b.docx")],
        warnings=["Warn 1"],
        errors=["Err 1"],
    )
    assert result.reports_generated == 2
    assert len(result.generated_paths) == 2
    assert len(result.warnings) == 1
    assert len(result.errors) == 1


def test_family_spec_lookup():
    """Verify all 5 CBM family specs are registered."""
    assert "fp_lvdb" in QUICK_REPORT_FAMILY_SPECS_BY_ID
    assert "swg" in QUICK_REPORT_FAMILY_SPECS_BY_ID
    assert "tx" in QUICK_REPORT_FAMILY_SPECS_BY_ID
    assert "blackbox" in QUICK_REPORT_FAMILY_SPECS_BY_ID
    assert "battery" in QUICK_REPORT_FAMILY_SPECS_BY_ID


def test_cbm_tech_summary_pairing():
    """Verify CBM tech summary 1-to-1 IR/US/TEV pairing behavior."""
    defects = [
        CbmDefectRecord(
            equipment="RMU",
            defect_area="Area 1",
            additional_remarks="Remark 1",
            technology="IR",
            raw_measurement="54.2",
            ir_reading="54.2",
        ),
        CbmDefectRecord(
            equipment="RMU",
            defect_area="Area 1",
            additional_remarks="Remark 1",
            technology="US",
            raw_measurement="12.0",
            us_reading="12.0",
        ),
        CbmDefectRecord(
            equipment="RMU",
            defect_area="Area 1",
            additional_remarks="Remark 1",
            technology="TEV",
            raw_measurement="24.0",
            tev_reading="24.0",
        ),
        CbmDefectRecord(
            equipment="TX",
            defect_area="Area 2",
            additional_remarks="",
            technology="IR",
            raw_measurement="45.0",
            ir_reading="45.0",
        ),
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 2

    rmu_row = next(r for r in rows if r.equipment == "RMU")
    assert rmu_row.ir_reading == "54.2 °C"
    assert rmu_row.us_reading == "12dB"
    assert rmu_row.tev_reading == "24dB"

    tx_row = next(r for r in rows if r.equipment == "TX")
    assert tx_row.ir_reading == "45.0 °C"
    assert tx_row.us_reading == "-"
    assert tx_row.tev_reading == "-"


def test_preflight_validation_missing_templates(tmp_path: Path):
    """Verify Stage 1 _validate_preconditions fails fast when an always-required template is missing."""
    tpl_file = tmp_path / "tpl.docx"
    tpl_file.touch()

    env = MagicMock(spec=ProjectEnvironment)
    env.get_vi_front_page_template.return_value = tmp_path / "missing_front_page.docx"
    env.get_template.return_value = tpl_file

    workflow = QuickReportWorkflow()
    req = QuickReportRequest(mode=QuickReportMode.FOLDER, target_folders=["01-01-2026"])

    with pytest.raises(FileNotFoundError, match="VI front page template missing"):
        workflow._validate_preconditions(env, req)


@patch("src.quick_report.extractor.SubstationTestsheetRepository")
def test_workflow_error_isolation(mock_repo_cls, tmp_path: Path):
    """Verify failure in one station does not abort batch in QuickReportWorkflow."""
    mock_repo = Mock()
    mock_repo_cls.return_value = mock_repo

    pkg1 = MagicMock()
    pkg1.station = "STATION 1"
    pkg1.substation_number = 1
    pkg1.data = MagicMock()
    pkg1.data.fl_erms = "FL1"
    pkg1.data.substation_name_erms = "STATION 1"
    pkg1.data.station_name = "STATION 1"

    pkg2 = MagicMock()
    pkg2.station = "STATION 2"
    pkg2.substation_number = 2
    pkg2.data = MagicMock()
    pkg2.data.fl_erms = "FL2"
    pkg2.data.substation_name_erms = "STATION 2"
    pkg2.data.station_name = "STATION 2"

    mock_repo.discover_packages.return_value = [pkg1, pkg2]

    tpl_file = tmp_path / "tpl.docx"
    tpl_file.touch()

    out2_file = tmp_path / "out2.docx"
    out2_file.write_text("dummy")

    env = MagicMock(spec=ProjectEnvironment)
    env.po_number = "42360565"
    env.state = "PAHANG"
    env.get_vi_front_page_template.return_value = tpl_file
    env.get_cbm_summary_template.return_value = tpl_file
    env.get_vi_summary_template.return_value = tpl_file
    env.get_vi_defect_template.return_value = tpl_file
    env.get_template.return_value = tpl_file

    mock_composer = Mock()
    mock_composer.load.side_effect = [Exception("Mock Error"), out2_file]

    workflow = QuickReportWorkflow(composer=mock_composer)
    req = QuickReportRequest(
        mode=QuickReportMode.FOLDER, target_folders=["01-01-2026"]
    )

    with (
        patch.object(workflow.extractor, "extract_defects", return_value=([], [])),
        patch("src.workflows.quick_report.win32com"),
        patch("src.workflows.quick_report.pythoncom"),
    ):
        result = workflow.execute(env, req)

    assert result.reports_generated == 1
    assert len(result.generated_paths) == 1
    assert result.generated_paths[0] == out2_file
    assert len(result.errors) == 1
    assert "Mock Error" in result.errors[0]
    assert "STATION 1" in result.errors[0]


def test_suffix_calculation():
    """Verify Canonical IR+US+TEV+VI suffix generation in QuickReportTransformer."""
    transformer = QuickReportTransformer()

    assert transformer._calculate_suffix([], []) == ("", [])
    assert transformer._calculate_suffix([CbmDefectRecord(technology="IR")], []) == (
        " (IR)",
        ["IR"],
    )
    assert transformer._calculate_suffix(
        [CbmDefectRecord(technology="IR"), CbmDefectRecord(technology="US")], []
    ) == (" (IR+US)", ["IR", "US"])
    assert transformer._calculate_suffix(
        [CbmDefectRecord(technology="IR"), CbmDefectRecord(technology="TEV")], [ViDefectRecord(equipment="SWG")]
    ) == (" (IR+TEV+VI)", ["IR", "TEV", "VI"])
    assert transformer._calculate_suffix([], [ViDefectRecord(equipment="SWG")]) == (" (VI)", ["VI"])
    assert transformer._calculate_suffix([CbmDefectRecord(technology="TEV")], []) == (
        " (TEV)",
        ["TEV"],
    )


@patch("src.quick_report.extractor.SubstationTestsheetRepository")
def test_extractor_folder_path_resolution(mock_repo_cls, tmp_path: Path):
    """Verify QuickReportExtractor resolves existing direct paths and relative testsheet paths."""
    mock_repo = Mock()
    mock_repo_cls.return_value = mock_repo
    mock_repo.discover_packages.return_value = []

    testsheet_dir = tmp_path / "TESTSHEET"
    testsheet_dir.mkdir()
    rel_folder = testsheet_dir / "01-01-2026"
    rel_folder.mkdir()

    direct_folder = tmp_path / "DIRECT_PATH"
    direct_folder.mkdir()

    env = MagicMock(spec=ProjectEnvironment)
    env.po_number = "42360565"
    env.state = "PAHANG"
    env.get_testsheet_dir.return_value = testsheet_dir

    extractor = QuickReportExtractor(repository=mock_repo)

    # Test direct path (Path(str) exists)
    req_direct = QuickReportRequest(
        mode=QuickReportMode.FOLDER, target_folders=[str(direct_folder)]
    )
    extractor.extract(env, req_direct)
    mock_repo.discover_packages.assert_called_with(direct_folder)

    # Test relative path under testsheet dir
    req_rel = QuickReportRequest(
        mode=QuickReportMode.FOLDER, target_folders=["01-01-2026"]
    )
    extractor.extract(env, req_rel)
    mock_repo.discover_packages.assert_called_with(rel_folder)


def test_extractor_raises_for_missing_folder(tmp_path: Path):
    """Verify QuickReportExtractor raises FileNotFoundError when a requested target folder does not exist."""
    testsheet_dir = tmp_path / "TESTSHEET"
    testsheet_dir.mkdir()

    env = MagicMock(spec=ProjectEnvironment)
    env.get_testsheet_dir.return_value = testsheet_dir

    extractor = QuickReportExtractor()
    req = QuickReportRequest(
        mode=QuickReportMode.FOLDER, target_folders=["nonexistent_folder"]
    )

    with pytest.raises(FileNotFoundError, match="Requested target folder does not exist"):
        extractor.extract(env, req)


def test_transformer_output_dir_resolution(tmp_path: Path):
    """Verify 3-tier Pahang output directory structure in QuickReportTransformer._resolve_output_dir."""
    quick_report_dir = tmp_path / "QUICK REPORT"
    env = MagicMock(spec=ProjectEnvironment)
    env.po_number = "42360565"
    env.state = "PAHANG"
    env.get_quick_report_dir.return_value = quick_report_dir

    transformer = QuickReportTransformer()

    # Case 1: 3-tier structure (station, month, date_str) formatted as XX. MONTH
    pkg_3tier = MagicMock()
    pkg_3tier.station = "KUANTAN"
    pkg_3tier.month = "01. JANUARY"
    pkg_3tier.date_str = "01-01-2026"

    out_3tier = transformer._resolve_output_dir(env, pkg_3tier)
    assert out_3tier == quick_report_dir / "KUANTAN" / "01. JANUARY" / "01-01-2026"

    # Case 2: 1-tier structure (date_str only)
    pkg_1tier = MagicMock()
    pkg_1tier.station = ""
    pkg_1tier.month = ""
    pkg_1tier.date_str = "02-01-2026"

    out_1tier = transformer._resolve_output_dir(env, pkg_1tier)
    assert out_1tier == quick_report_dir / "02-01-2026"

    # Case 3: Root structure (no date_str, station, or month)
    pkg_root = MagicMock()
    pkg_root.station = ""
    pkg_root.month = ""
    pkg_root.date_str = ""

    out_root = transformer._resolve_output_dir(env, pkg_root)
    assert out_root == quick_report_dir


def test_master_qr03_defect_repository_empty(tmp_path: Path):
    """Verify MasterQr03DefectRepository raises FileNotFoundError when ENGR files/directory are missing."""
    repo = MasterQr03DefectRepository(tmp_path / "nonexistent")
    with pytest.raises(FileNotFoundError):
        repo.fetch_cbm_defects("12345")
    with pytest.raises(FileNotFoundError):
        repo.fetch_vi_defects("12345")


def test_master_qr03_defect_repository_with_excel(tmp_path: Path):
    """Verify MasterQr03DefectRepository extracts CBM and VI defects strictly from QR03 CBA and QR03 VI sheets."""
    import pandas as pd

    cba_df = pd.DataFrame(
        [
            {
                "FUNCTIONAL LOCATION": "F/L 12345",
                "EQUIPMENT": "RMU SF6",
                "TECHNOLOGY": "IR",
                "BRAND": "ABB",
                "MODEL": "SafePlus",
                "RATING": "11kV",
                "DEFECT AREA": "Cable Compartment",
                "REMARKS": "Hotspot detected",
                "READING": "65.4",
            },
            {
                "FUNCTIONAL LOCATION": "F/L 12345",
                "EQUIPMENT": "RMU SF6",
                "TECHNOLOGY": "US",
                "BRAND": "ABB",
                "MODEL": "SafePlus",
                "RATING": "11kV",
                "DEFECT AREA": "Cable Compartment",
                "REMARKS": "Arcing detected",
                "READING": "18.0",
            },
        ]
    )
    vi_df = pd.DataFrame(
        [
            {
                "FUNCTIONAL LOCATION": "F/L 12345",
                "EQUIPMENT": "SUBSTATION BUILDING",
                "DEFECT AREA": "Door",
                "REMARKS": "Rusty hinge",
            }
        ]
    )

    engr_path = tmp_path / "ENGR-750-36-CBA-TEST-2026.xlsx"
    with pd.ExcelWriter(engr_path, engine="openpyxl") as writer:
        cba_df.to_excel(writer, sheet_name="QR03 CBA", index=False)
        vi_df.to_excel(writer, sheet_name="QR03 VI", index=False)

    repo = MasterQr03DefectRepository(tmp_path)
    cbm = repo.fetch_cbm_defects("12345")
    vi = repo.fetch_vi_defects("12345")

    assert len(cbm) == 2
    assert cbm[0].equipment == "RMU SF6"
    assert cbm[0].technology == "IR"
    assert cbm[0].raw_measurement == "65.4"

    assert len(vi) == 1
    assert vi[0].equipment == "SUBSTATION BUILDING"
    assert vi[0].defect_area == "Door"


def test_build_substation_condition_pairs():
    """Verify equipment pair building in QuickReportTransformer."""
    from src.testsheet.models import (
        BatteryBankSpec,
        FireExtinguisherSpec,
        LVDBSpec,
        SubstationEquipmentPackage,
        SubstationTestsheetPackage,
        SwitchgearSpec,
        TestsheetData,
        TransformerSpec,
    )

    transformer = QuickReportTransformer()

    # None / Empty pkg fallback
    assert transformer._build_substation_condition_pairs(None) == [
        ("SUBSTATION OVERVIEW", "SIGNBOARD")
    ]

    eq = SubstationEquipmentPackage(
        switchgears=(SwitchgearSpec(switchgear_type="GIS"),),
        transformers=(TransformerSpec(tx_id="Tx 1"),),
        lvdb_specs=(LVDBSpec(name="FP 1", label="FP"),),
        battery_banks=(BatteryBankSpec(name="BATTERY BANK 1"),),
        fire_extinguisher=FireExtinguisherSpec(has_fire_extinguisher=True),
        has_battery_charger=True,
        has_rtu=True,
        has_sf6=True,
        has_efi=True,
    )
    data = TestsheetData(
        substation_number=1,
        substation_name_erms="PE TEST",
        building_type="INDOOR",
        substation_type="MRMU SF6",
        equipment=eq,
    )
    pkg = SubstationTestsheetPackage(
        testsheet_path=Path("dummy.xlsx"),
        unsorted_raw_data_dir=Path("dummy_raw"),
        station="TEST",
        month="08. AUGUST",
        date_str="12-08-2026",
        substation_number=1,
        data=data,
    )

    pairs = transformer._build_substation_condition_pairs(pkg)
    labels = [p[0] for p in pairs]

    assert "SUBSTATION OVERVIEW" in labels
    assert "SWITCHGEAR" in labels
    assert "TRANSFORMER" in labels
    assert "FEEDER PILLAR" in labels
    assert "BATTERY CHARGER" in labels
    assert "RTU" in labels
    assert "FIRE EXTINGUISHER\n(SWITCHGEAR ROOM)" in labels
    assert "FIRE EXTINGUISHER\n(TX ROOM)" in labels
    assert "TRANSFORMER OIL LEVEL INDICATOR" in labels


def test_composer_cbm_defect_pages_with_cbm_defects(tmp_path: Path):
    """Verify composer calls CBM defect pages generation correctly when cbm_defects exist."""
    composer = QuickReportComposer()
    tpl_file = tmp_path / "tpl.docx"
    tpl_file.touch()

    pkg = MagicMock()
    pkg.station = "CAMERON HIGHLAND"
    pkg.substation_number = 1

    pe_info = {"substation": {"name_erms": "CAMERON HIGHLAND"}}
    cbm_defects = (
        CbmDefectRecord(
            equipment="FP (D)",
            technology="IR",
            brand="ABB",
            model="X",
            rating="11kV",
            defect_area="Body",
            additional_remarks="Hotspot",
            ir_reading="50",
            us_reading="",
            tev_reading="",
            raw_measurement="50",
        ),
    )

    from src.quick_report.models import CbmDefectFamilyPlan, CbmDefectGroup

    family_plan = CbmDefectFamilyPlan(
        spec=QUICK_REPORT_FAMILY_SPECS_BY_ID["fp_lvdb"],
        overview_template=tpl_file,
        detail_templates=(("fp_detail", tpl_file),),
        groups=(
            CbmDefectGroup(
                item_key="FP (D)",
                item_suffix="",
                defects=(cbm_defects[0],),
                overview=cbm_defects[0],
            ),
        ),
    )

    plan = QuickReportStationPlan(
        package=pkg,
        pe_info=pe_info,
        cbm_defects=cbm_defects,
        vi_defects=(),
        suffix=" (IR)",
        suffix_parts=("IR",),
        output_dir=tmp_path,
        output_filename="001. CAMERON HIGHLAND (IR).docx",
        final_output_path=tmp_path / "001. CAMERON HIGHLAND (IR).docx",
        condition_pairs=(),
        cond_template_path=None,
        front_page_template=tpl_file,
        cbm_summary_template=tpl_file,
        vi_summary_template=tpl_file,
        vi_defect_template=tpl_file,
        sticker_template=tpl_file,
        cbm_defect_family_plans=(family_plan,),
    )

    with (
        patch(
            "src.quick_report.composer.generate_front_page",
            return_value=tmp_path / "p1.docx",
        ),
        patch(
            "src.quick_report.composer.generate_cbm_tech_summary",
            return_value=tmp_path / "p2a.docx",
        ),
        patch(
            "src.quick_report.composer.generate_cbm_defect_pages",
            return_value=[tmp_path / "p2b.docx"],
        ) as mock_gen_cbm,
        patch(
            "src.quick_report.composer.generate_sticker_page",
            return_value=tmp_path / "p7.docx",
        ),
    ):
        parts = composer._generate_parts(plan, tmp_path)
        assert len(parts) >= 3
        assert mock_gen_cbm.called
        call_args = mock_gen_cbm.call_args[0]
        assert call_args[0] == family_plan
        assert call_args[1] == tmp_path  # temp_dir
        assert call_args[2] == 1  # substation_number
        assert call_args[3] == pe_info  # pe_info


def test_compile_document_word_com_success(tmp_path: Path):
    """Verify _compile_document uses Documents.Add and Recopy & Paste to combine document parts."""
    composer = QuickReportComposer()

    mock_word = MagicMock()
    mock_main_doc = MagicMock()
    mock_part_doc = MagicMock()
    mock_rng = MagicMock()

    mock_word.Documents.Add.return_value = mock_main_doc
    mock_word.Documents.Open.return_value = mock_part_doc
    mock_main_doc.Content = mock_rng
    mock_main_doc.Tables.Count = 0
    mock_rng.Information.return_value = False

    part1 = tmp_path / "part1.docx"
    part2 = tmp_path / "part2.docx"
    output_path = tmp_path / "output.docx"

    part1.touch()
    part2.touch()

    composer._compile_document([part1, part2], output_path, word_app=mock_word)

    # Verification
    mock_word.Documents.Add.assert_called_once()
    assert mock_word.Documents.Open.call_count == 2
    mock_part_doc.Content.Copy.assert_called()
    mock_part_doc.Close.assert_called_with(False)

    mock_rng.InsertBreak.assert_called_once_with(7)
    assert mock_rng.Paste.call_count == 2

    mock_main_doc.SaveAs2.assert_called_once_with(str(output_path.resolve()))
    mock_main_doc.Close.assert_called_once_with(False)


def test_compile_document_word_com_cleanup_on_error(tmp_path: Path):
    """Verify _compile_document executes cleanup (Close) and re-raises exception when COM fails."""
    composer = QuickReportComposer()

    mock_word = MagicMock()
    mock_main_doc = MagicMock()

    mock_word.Documents.Add.return_value = mock_main_doc
    mock_word.Documents.Open.side_effect = RuntimeError("Open failed")

    part1 = tmp_path / "part1.docx"
    output_path = tmp_path / "output.docx"
    part1.touch()

    with pytest.raises(RuntimeError, match="Open failed"):
        composer._compile_document([part1], output_path, word_app=mock_word)

    # Verification of cleanup in finally block
    mock_main_doc.Close.assert_called_once_with(False)


def test_generate_cbm_defect_pages_filename_uniqueness(tmp_path: Path):

    """Verify unique filenames are generated when multiple groups and defects exist."""
    from src.quick_report.cbm_defect_pages import generate_cbm_defect_pages
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.defects import CbmDefectRecord
    from src.quick_report.models import CbmDefectDetailGroup, CbmDefectFamilyPlan, CbmDefectGroup

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]

    overview_t = tmp_path / "swg_overview.docx"
    panel_t = tmp_path / "swg_panel.docx"
    overview_t.touch()
    panel_t.touch()

    d1 = CbmDefectRecord(equipment="RMU SF6", technology="IR")
    d2 = CbmDefectRecord(equipment="RMU SF6", technology="US")
    d3 = CbmDefectRecord(equipment="RMU SF6", technology="IR")

    group1 = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="",
        defects=(d1, d2),
        overview=d1,
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(d1, d2)),
        ),
    )
    group2 = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="",
        defects=(d3,),
        overview=d3,
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(d3,)),
        ),
    )

    family_plan = CbmDefectFamilyPlan(
        spec=spec,
        overview_template=overview_t,
        detail_templates=(("panel_area", panel_t),),
        groups=(group1, group2),
    )

    with patch("src.quick_report.cbm_defect_pages._render_docx_template"):
        paths = generate_cbm_defect_pages(
            family_plan, tmp_path, 1, {"pe": "info"}
        )

    filenames = [p.name for p in paths]
    assert len(filenames) == len(
        set(filenames)
    ), f"Duplicate filenames found: {filenames}"
    assert "001_3 SWG OVERVIEW_grp1.docx" in filenames
    assert "001_3 SWG OVERVIEW_grp2.docx" in filenames
    assert "001_3 SWG RMU SF6_grp1_part1.docx" in filenames
    assert "001_3 SWG RMU SF6_grp1_part2.docx" in filenames
    assert "001_3 SWG RMU SF6_grp2.docx" in filenames


def test_vi_defect_pages_remove_empty_cell_borders_exception_handling(
    tmp_path: Path, caplog
):
    """Verify _remove_empty_cell_borders logs a warning and does not crash when an exception occurs."""
    from src.quick_report.vi_defect_pages import _remove_empty_cell_borders

    invalid_path = tmp_path / "non_existent.docx"
    with caplog.at_level(logging.WARNING):
        _remove_empty_cell_borders(invalid_path, 0)

    assert "Failed to remove empty cell borders" in caplog.text


def test_substation_condition_remove_empty_cell_borders_multi_table(tmp_path: Path):
    """Verify _remove_empty_cell_borders_sub_cond dehighlights unused tables in multi-table layout."""
    from docx import Document

    from src.quick_report.substation_condition import (
        _remove_empty_cell_borders_sub_cond,
    )

    doc_path = tmp_path / "sub_cond_multi.docx"
    doc = Document()
    doc.add_table(rows=2, cols=2)
    t1 = doc.add_table(rows=2, cols=2)
    t2 = doc.add_table(rows=2, cols=2)

    doc.tables[0].cell(0, 0).paragraphs[0].text = "Table 0 Active"
    t1.cell(0, 0).paragraphs[0].text = "Table 1 Unused"
    t2.cell(0, 0).paragraphs[0].text = "Table 2 Unused"
    doc.save(doc_path)

    _remove_empty_cell_borders_sub_cond(doc_path, active_count=1)

    updated = Document(doc_path)
    assert updated.tables[0].cell(0, 0).text == "Table 0 Active"
    assert updated.tables[1].cell(0, 0).text == ""
    assert updated.tables[2].cell(0, 0).text == ""

    tcPr1 = updated.tables[1].cell(0, 0)._tc.get_or_add_tcPr()
    assert (
        tcPr1.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders"
        )
        is not None
    )

    tcPr2 = updated.tables[2].cell(0, 0)._tc.get_or_add_tcPr()
    assert (
        tcPr2.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders"
        )
        is not None
    )


def test_substation_condition_remove_empty_cell_borders_single_table(tmp_path: Path):
    """Verify _remove_empty_cell_borders_sub_cond dehighlights unused slot rows in single-table layout."""
    from docx import Document

    from src.quick_report.substation_condition import (
        _remove_empty_cell_borders_sub_cond,
    )

    doc_path = tmp_path / "sub_cond_single.docx"
    doc = Document()
    t = doc.add_table(rows=6, cols=2)

    t.cell(0, 0).paragraphs[0].text = "Header 0"
    t.cell(1, 0).paragraphs[0].text = "Photo 0"
    t.cell(2, 0).paragraphs[0].text = "Header 1"
    t.cell(3, 0).paragraphs[0].text = "Photo 1"
    t.cell(4, 0).paragraphs[0].text = "Header 2"
    t.cell(5, 0).paragraphs[0].text = "Photo 2"
    doc.save(doc_path)

    _remove_empty_cell_borders_sub_cond(doc_path, active_count=1)

    updated = Document(doc_path)
    assert updated.tables[0].cell(0, 0).text == "Header 0"
    assert updated.tables[0].cell(1, 0).text == "Photo 0"

    assert updated.tables[0].cell(2, 0).text == ""
    assert updated.tables[0].cell(3, 0).text == ""
    assert updated.tables[0].cell(4, 0).text == ""
    assert updated.tables[0].cell(5, 0).text == ""

    tcPr = updated.tables[0].cell(2, 0)._tc.get_or_add_tcPr()
    assert (
        tcPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders"
        )
        is not None
    )


def test_substation_condition_remove_empty_cell_borders_exception_handling(
    tmp_path: Path, caplog
):
    """Verify _remove_empty_cell_borders_sub_cond logs a warning and handles errors gracefully."""
    from src.quick_report.substation_condition import (
        _remove_empty_cell_borders_sub_cond,
    )

    invalid_path = tmp_path / "non_existent.docx"
    with caplog.at_level(logging.WARNING):
        _remove_empty_cell_borders_sub_cond(invalid_path, 0)

    assert "Failed to remove empty cell borders" in caplog.text


def test_quick_report_fl_mode_fl_erms_matching(monkeypatch, tmp_path: Path):
    """Verify QuickReportWorkflow FL mode filters using TestsheetData.fl_erms without AttributeError."""
    from src.testsheet.models import SubstationTestsheetPackage, TestsheetData

    mock_pkg1 = SubstationTestsheetPackage(
        testsheet_path=Path("test1.xlsx"),
        unsorted_raw_data_dir=Path("dir1"),
        station="CCHL",
        month="05. MAY",
        date_str="2026-05-15",
        substation_number=1,
        data=TestsheetData(
            substation_number=1,
            substation_name_erms="SUBSTATION A",
            fl_erms="CCHL/PCE/J00059",
            fl_site="CCHL/PCE/J00059",
        ),
    )

    mock_pkg2 = SubstationTestsheetPackage(
        testsheet_path=Path("test2.xlsx"),
        unsorted_raw_data_dir=Path("dir2"),
        station="CCHL",
        month="05. MAY",
        date_str="2026-05-15",
        substation_number=2,
        data=TestsheetData(
            substation_number=2,
            substation_name_erms="SUBSTATION B",
            fl_erms="CCHL/PCE/J00060",
            fl_site="CCHL/PCE/J00060",
        ),
    )

    monkeypatch.setattr(
        "src.testsheet.repository.SubstationTestsheetRepository.discover_packages",
        lambda self, folder_path: [mock_pkg1, mock_pkg2],
    )

    tpl_file = tmp_path / "tpl.docx"
    tpl_file.touch()

    env = MagicMock()
    env.get_vi_front_page_template.return_value = tpl_file
    env.get_cbm_summary_template.return_value = tpl_file
    env.get_vi_summary_template.return_value = tpl_file
    env.get_vi_defect_template.return_value = tpl_file
    env.get_template.return_value = tpl_file

    req = QuickReportRequest(
        mode=QuickReportMode.FL,
        target_package_names=("CCHL/PCE/J00059",),
    )

    dummy_path = tmp_path / "dummy.docx"
    dummy_path.write_text("content")

    mock_composer = Mock()
    mock_composer.load.return_value = dummy_path

    workflow = QuickReportWorkflow(composer=mock_composer)
    with (
        patch("src.workflows.quick_report.win32com"),
        patch("src.workflows.quick_report.pythoncom"),
        patch.object(workflow.extractor, "extract_defects", return_value=([], [])),
    ):
        result = workflow.execute(env, req)
    assert result.reports_generated == 1
    assert len(result.generated_paths) == 1


def test_workflow_fetches_defects_from_repository(tmp_path: Path):
    """Regression: QuickReport workflow must call MasterQr03DefectRepository to fetch
    cbm_defects and vi_defects and pass them through to transformer/composer."""
    mock_data = Mock()
    mock_data.substation_name_erms = "KEA FARM"
    mock_data.station_name = "KEA FARM"
    mock_data.date_str = "17-06-2026"
    mock_data.gps_coordinate = "4.123,101.456"
    mock_data.substation_type = "PMU"
    mock_data.building_type = "BRICK"
    mock_data.fl_erms = "CCHL/PCE/J00059"
    mock_data.fl_site = "CCHL/PCE/J00059"

    mock_pkg = Mock()
    mock_pkg.data = mock_data
    mock_pkg.station = "CAMERON HIGHLAND"
    mock_pkg.month = "06. JUNE"
    mock_pkg.date_str = "17-06-2026"
    mock_pkg.substation_number = 303

    sample_cbm = [CbmDefectRecord(equipment="SWG", technology="IR", defect_area="Cable lug")]
    sample_vi = [ViDefectRecord(equipment="SWG", defect_area="Rust", additional_remarks="Mild")]

    tpl_file = tmp_path / "tpl.docx"
    tpl_file.touch()

    env = MagicMock()
    env.get_quick_report_dir.return_value = tmp_path
    env.get_vi_front_page_template.return_value = tpl_file
    env.get_cbm_summary_template.return_value = tpl_file
    env.get_vi_summary_template.return_value = tpl_file
    env.get_vi_defect_template.return_value = tpl_file
    env.get_template.return_value = tpl_file

    req = QuickReportRequest(
        mode=QuickReportMode.FL, target_package_names=("CCHL/PCE/J00059",)
    )

    dummy_path = tmp_path / "dummy.docx"
    dummy_path.write_text("content")

    captured_plan = {}

    def spy_load(plan, word_app=None):
        captured_plan["plan"] = plan
        return dummy_path

    mock_composer = Mock()
    mock_composer.load.side_effect = spy_load

    with (
        patch("src.workflows.quick_report.win32com"),
        patch("src.workflows.quick_report.pythoncom"),
        patch("src.quick_report.extractor.SubstationTestsheetRepository") as MockRepo,
        patch(
            "src.quick_report.extractor.MasterQr03DefectRepository"
        ) as MockDefectRepo,
    ):
        MockRepo.return_value.discover_packages.return_value = [mock_pkg]

        defect_repo_instance = MockDefectRepo.return_value
        defect_repo_instance.fetch_cbm_defects.return_value = sample_cbm
        defect_repo_instance.fetch_vi_defects.return_value = sample_vi

        workflow = QuickReportWorkflow(composer=mock_composer)
        workflow.execute(env, req)

    assert "plan" in captured_plan
    plan = captured_plan["plan"]
    assert (
        plan.cbm_defects == tuple(sample_cbm)
    ), "cbm_defects not passed through from repository"
    assert (
        plan.vi_defects == tuple(sample_vi)
    ), "vi_defects not passed through from repository"


def test_extract_defects_returns_empty_when_valid_sheet_has_no_matching_fl(tmp_path: Path):
    """Valid ENGR workbook with QR03 CBA/VI sheets but no rows matching the FL → returns ([], [])."""
    import pandas as pd

    cba_df = pd.DataFrame([
        {
            "FUNCTIONAL LOCATION": "F/L 99999",
            "EQUIPMENT": "RMU SF6",
            "TECHNOLOGY": "IR",
            "DEFECT AREA": "Cable Compartment",
            "READING": "65.4",
        }
    ])
    vi_df = pd.DataFrame([
        {
            "FUNCTIONAL LOCATION": "F/L 99999",
            "EQUIPMENT": "SUBSTATION BUILDING",
            "DEFECT AREA": "Door",
        }
    ])

    engr_path = tmp_path / "ENGR-750-36-CBA-TEST-2026.xlsx"
    with pd.ExcelWriter(engr_path, engine="openpyxl") as writer:
        cba_df.to_excel(writer, sheet_name="QR03 CBA", index=False)
        vi_df.to_excel(writer, sheet_name="QR03 VI", index=False)

    pkg = MagicMock()
    pkg.data.fl_erms = "12345"

    env = MagicMock()
    env.storage.get_engr_folder.return_value = tmp_path

    extractor = QuickReportExtractor()
    cbm, vi = extractor.extract_defects(pkg, env)

    assert cbm == []
    assert vi == []


def test_extract_defects_raises_when_engr_folder_missing(tmp_path: Path):
    """Missing ENGR directory → FileNotFoundError, NOT empty lists."""
    pkg = MagicMock()
    pkg.data.fl_erms = "12345"

    env = MagicMock()
    env.storage.get_engr_folder.return_value = tmp_path / "nonexistent"

    extractor = QuickReportExtractor()
    with pytest.raises(FileNotFoundError, match="Required ENGR directory does not exist"):
        extractor.extract_defects(pkg, env)


def test_extract_defects_raises_when_no_engr_workbooks(tmp_path: Path):
    """ENGR directory exists but contains no .xlsx files → FileNotFoundError."""
    pkg = MagicMock()
    pkg.data.fl_erms = "12345"

    env = MagicMock()
    env.storage.get_engr_folder.return_value = tmp_path

    extractor = QuickReportExtractor()
    with pytest.raises(FileNotFoundError, match="No ENGR Excel workbooks found"):
        extractor.extract_defects(pkg, env)


def test_extract_defects_raises_when_qr03_cba_sheet_missing(tmp_path: Path):
    """ENGR workbook exists but missing 'QR03 CBA' sheet → RuntimeError."""
    import pandas as pd

    dummy_df = pd.DataFrame([{"FL": "12345"}])
    engr_path = tmp_path / "ENGR-750-36-CBA-TEST-2026.xlsx"
    with pd.ExcelWriter(engr_path, engine="openpyxl") as writer:
        dummy_df.to_excel(writer, sheet_name="Sheet1", index=False)

    pkg = MagicMock()
    pkg.data.fl_erms = "12345"

    env = MagicMock()
    env.storage.get_engr_folder.return_value = tmp_path

    extractor = QuickReportExtractor()
    with pytest.raises(RuntimeError, match="Missing required sheet 'QR03 CBA'"):
        extractor.extract_defects(pkg, env)


def test_extract_defects_raises_when_qr03_vi_sheet_missing(tmp_path: Path):
    """ENGR workbook exists with QR03 CBA but missing 'QR03 VI' sheet → RuntimeError."""
    import pandas as pd

    cba_df = pd.DataFrame([
        {
            "FUNCTIONAL LOCATION": "F/L 12345",
            "EQUIPMENT": "RMU SF6",
            "TECHNOLOGY": "IR",
            "DEFECT AREA": "Cable Compartment",
            "READING": "65.4",
        }
    ])
    engr_path = tmp_path / "ENGR-750-36-CBA-TEST-2026.xlsx"
    with pd.ExcelWriter(engr_path, engine="openpyxl") as writer:
        cba_df.to_excel(writer, sheet_name="QR03 CBA", index=False)

    pkg = MagicMock()
    pkg.data.fl_erms = "12345"

    env = MagicMock()
    env.storage.get_engr_folder.return_value = tmp_path

    extractor = QuickReportExtractor()
    with pytest.raises(RuntimeError, match="Missing required sheet 'QR03 VI'"):
        extractor.extract_defects(pkg, env)


def test_workflow_fresh_word_com_session_per_package(tmp_path: Path):
    """Verify batch execution processes multiple packages with fresh isolated Word sessions and cleans up resources cleanly."""
    pkg1 = MagicMock()
    pkg1.station = "STATION A"
    pkg1.substation_number = 1
    pkg1.data = MagicMock()
    pkg1.data.fl_erms = "FL1"
    pkg1.data.substation_name_erms = "STATION A"
    pkg1.data.station_name = "STATION A"

    pkg2 = MagicMock()
    pkg2.station = "STATION B"
    pkg2.substation_number = 2
    pkg2.data = MagicMock()
    pkg2.data.fl_erms = "FL2"
    pkg2.data.substation_name_erms = "STATION B"
    pkg2.data.station_name = "STATION B"

    pkg3 = MagicMock()
    pkg3.station = "STATION C"
    pkg3.substation_number = 3
    pkg3.data = MagicMock()
    pkg3.data.fl_erms = "FL3"
    pkg3.data.substation_name_erms = "STATION C"
    pkg3.data.station_name = "STATION C"

    tpl_file = tmp_path / "tpl.docx"
    tpl_file.touch()

    out1 = tmp_path / "out1.docx"
    out2 = tmp_path / "out2.docx"
    out3 = tmp_path / "out3.docx"
    out1.write_text("dummy1")
    out2.write_text("dummy2")
    out3.write_text("dummy3")

    env = MagicMock(spec=ProjectEnvironment)
    env.get_vi_front_page_template.return_value = tpl_file
    env.get_template.return_value = tpl_file

    mock_word1 = MagicMock()
    mock_word2 = MagicMock()
    mock_word3 = MagicMock()

    dispatched_apps = [mock_word1, mock_word2, mock_word3]
    captured_word_apps = []

    def spy_load(plan, word_app=None):
        captured_word_apps.append(word_app)
        if len(captured_word_apps) == 1:
            return out1
        elif len(captured_word_apps) == 2:
            return out2
        else:
            return out3

    mock_composer = Mock()
    mock_composer.load.side_effect = spy_load

    workflow = QuickReportWorkflow(composer=mock_composer)
    workflow.extractor = Mock()
    workflow.extractor.extract.return_value = [pkg1, pkg2, pkg3]
    workflow.filter_stage = Mock()
    workflow.filter_stage.filter.return_value = [pkg1, pkg2, pkg3]
    workflow.extractor.extract_defects.return_value = ([], [])

    req = QuickReportRequest(mode=QuickReportMode.FOLDER, target_folders=["01-01-2026"])

    mock_win32com = MagicMock()
    mock_win32com.client.Dispatch.return_value = mock_word1

    with (
        patch("src.workflows.quick_report.win32com", mock_win32com),
        patch("src.workflows.quick_report.pythoncom") as mock_pythoncom,
    ):
        result = workflow.execute(env, req)

    mock_win32com.client.Dispatch.assert_called_once_with("Word.Application")
    mock_pythoncom.CoInitialize.assert_called_once()
    mock_pythoncom.CoUninitialize.assert_called_once()
    mock_word1.Quit.assert_called_once()
    assert captured_word_apps == [mock_word1, mock_word1, mock_word1]
    assert result.reports_generated == 3
    assert result.generated_paths == [out1, out2, out3]
    assert len(result.errors) == 0


def test_workflow_fresh_word_com_session_cleanup_on_per_package_error(tmp_path: Path):
    """Verify that when a package fails midway, its isolated Word COM session is still quit cleanly."""
    pkg1 = MagicMock()
    pkg1.station = "STATION FAILING"
    pkg1.substation_number = 1
    pkg1.data = MagicMock()
    pkg1.data.fl_erms = "FL1"
    pkg1.data.substation_name_erms = "STATION FAILING"
    pkg1.data.station_name = "STATION FAILING"

    pkg2 = MagicMock()
    pkg2.station = "STATION OK"
    pkg2.substation_number = 2
    pkg2.data = MagicMock()
    pkg2.data.fl_erms = "FL2"
    pkg2.data.substation_name_erms = "STATION OK"
    pkg2.data.station_name = "STATION OK"

    tpl_file = tmp_path / "tpl.docx"
    tpl_file.touch()

    out2 = tmp_path / "out2.docx"
    out2.write_text("dummy2")

    env = MagicMock(spec=ProjectEnvironment)
    env.get_vi_front_page_template.return_value = tpl_file
    env.get_template.return_value = tpl_file

    mock_word1 = MagicMock()
    mock_word2 = MagicMock()

    mock_composer = Mock()
    mock_composer.load.side_effect = [RuntimeError("Composer crash"), out2]

    workflow = QuickReportWorkflow(composer=mock_composer)
    workflow.extractor = Mock()
    workflow.extractor.extract.return_value = [pkg1, pkg2]
    workflow.filter_stage = Mock()
    workflow.filter_stage.filter.return_value = [pkg1, pkg2]
    workflow.extractor.extract_defects.return_value = ([], [])

    req = QuickReportRequest(mode=QuickReportMode.FOLDER, target_folders=["01-01-2026"])

    mock_win32com = MagicMock()
    mock_win32com.client.Dispatch.return_value = mock_word1

    with (
        patch("src.workflows.quick_report.win32com", mock_win32com),
        patch("src.workflows.quick_report.pythoncom") as mock_pythoncom,
    ):
        result = workflow.execute(env, req)

    mock_win32com.client.Dispatch.assert_called_once_with("Word.Application")
    mock_pythoncom.CoInitialize.assert_called_once()
    mock_pythoncom.CoUninitialize.assert_called_once()
    mock_word1.Quit.assert_called_once()
    assert result.reports_generated == 1
    assert len(result.errors) == 1
    assert "Composer crash" in result.errors[0]


def test_generate_substation_condition_pages_removes_trailing_sectpr(tmp_path: Path):
    """Verify generate_substation_condition_pages merges multi-part condition files without trailing sectPr on last paragraph."""
    from docx import Document
    from docx.oxml.ns import qn

    from src.quick_report.substation_condition import (
        generate_substation_condition_pages,
    )

    template_doc = Document()
    template_doc.add_paragraph("Substation Condition Template")
    tpl_path = tmp_path / "substation_condition_template.docx"
    template_doc.save(tpl_path)

    # 5 pairs forces 2 parts (chunks of 3) -> triggers DocxComposer merge
    pairs = [
        ("SUBSTATION OVERVIEW", "SIGNBOARD"),
        ("SWITCHGEAR 1", "SWITCHGEAR 1 NAMEPLATE"),
        ("TRANSFORMER 1", "TRANSFORMER 1 NAMEPLATE"),
        ("FEEDER PILLAR 1", "FEEDER PILLAR 1 NAMEPLATE"),
        ("BATTERY CHARGER", "BATTERY CHARGER NAMEPLATE"),
    ]

    out_files = generate_substation_condition_pages(
        pe_info={"pe_name": "PE TEST"},
        condition_pairs_or_pkg=pairs,
        template_path=tpl_path,
        output_dir=tmp_path,
        substation_number=1,
    )

    assert len(out_files) == 1
    merged_path = out_files[0]
    assert merged_path.name == "001_5 SUBSTATION CONDITION.docx"

    merged_doc = Document(merged_path)
    assert len(merged_doc.paragraphs) > 0
    last_para = merged_doc.paragraphs[-1]
    pPr = last_para._p.get_or_add_pPr()
    sectPr = pPr.find(qn("w:sectPr"))
    assert sectPr is None
    assert len(last_para._p.findall(qn("w:br"))) == 0

def test_generate_substation_condition_pages_removes_br_page_breaks_cleanly(tmp_path: Path):
    """Verify generate_substation_condition_pages strips w:br from last paragraph without affecting preceding content."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from src.quick_report.substation_condition import (
        generate_substation_condition_pages,
    )

    template_doc = Document()
    p1 = template_doc.add_paragraph("Preceding Paragraph")
    run1 = p1.add_run("Text in paragraph 1")
    
    # Add a paragraph with a break
    p_last = template_doc.add_paragraph()
    r_last = p_last.add_run("Trailing text")
    
    # Inject w:br inside r_last
    br1 = OxmlElement("w:br")
    r_last._r.append(br1)
    
    # Inject direct w:br inside p_last._p
    br2 = OxmlElement("w:br")
    p_last._p.append(br2)

    tpl_path = tmp_path / "substation_condition_template_br.docx"
    template_doc.save(tpl_path)

    pairs = [
        ("SUBSTATION OVERVIEW", "SIGNBOARD"),
        ("SWITCHGEAR 1", "SWITCHGEAR 1 NAMEPLATE"),
        ("TRANSFORMER 1", "TRANSFORMER 1 NAMEPLATE"),
        ("FEEDER PILLAR 1", "FEEDER PILLAR 1 NAMEPLATE"),
    ]

    out_files = generate_substation_condition_pages(
        pe_info={"pe_name": "PE TEST BR"},
        condition_pairs_or_pkg=pairs,
        template_path=tpl_path,
        output_dir=tmp_path,
        substation_number=2,
    )

    assert len(out_files) == 1
    merged_doc = Document(out_files[0])
    
    # Check preceding paragraph is intact
    assert merged_doc.paragraphs[0].text == "Preceding ParagraphText in paragraph 1"

    
    # Check last paragraph has no w:br elements
    last_para = merged_doc.paragraphs[-1]
    assert len(last_para._p.findall(qn("w:br"))) == 0
    for r in last_para.runs:
        assert len(r._r.findall(qn("w:br"))) == 0





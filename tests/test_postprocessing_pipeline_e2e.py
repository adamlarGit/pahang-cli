"""End-to-End Integration Test Suite for the 1-Click Substation Post-Processing Pipeline.

Covers full 6-stage lifecycle across:
- Scenario A: Full end-to-end by_date execution (Pre-flight, Renaming Sync, WhatsApp, Signatures, Diagonals, Merge).
- Scenario B: End-to-end by_fl execution (Subset scoping, WhatsApp skipped, Signatures 'none' placeholder stripping).
- Scenario C: Pre-flight count mismatch fail-fast (Testsheet vs Quick Report mismatch halts before modification).
- Scenario D: Substation converter failure resilience (1 failed station isolated, remaining batch succeeds).
- Scenario E: Pre-flight raw material count mismatch fail-fast.
- Scenario F: CLI Presentation Adapter execution (PostProcessingPipelineAction end-to-end).
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, patch
import openpyxl
from PIL import Image as PILImage
import pytest
from PyPDF2 import PdfReader
from docx import Document

from src.postprocessing.converters import FakeDocumentConverter
from src.project.environment import ProjectEnvironment
from src.project.models import ProjectMetadata
from src.project.storage import LocalWorkspaceStorage
from src.project_workflow_actions import (
    PostProcessingPipelineAction,
    _print_postprocessing_summary,
)
from src.workflows.models import (
    PostProcessingFailure,
    PostProcessingMode,
    PostProcessingRequest,
    PostProcessingSummary,
)
from src.workflows.postprocessing_pipeline import (
    PostProcessingPipelineWorkflow,
    run_postprocessing_pipeline,
)
from src.workflows.postprocessing_preflight import PreFlightValidationError
from src.workflows.service import WorkflowService


def _setup_e2e_workspace(root_path: Path) -> ProjectEnvironment:
    """Scaffold a full realistic project workspace backed by temporary root_path."""
    python_dir = root_path / "PYTHON"
    templates_dir = root_path / "templates"
    sign_dir = root_path / "OTHERS" / "SIGN"

    python_dir.mkdir(parents=True, exist_ok=True)
    templates_dir.mkdir(parents=True, exist_ok=True)
    sign_dir.mkdir(parents=True, exist_ok=True)

    # 1. TOTAL PE.xlsx with DataCycle1 sheet
    total_pe_path = python_dir / "TOTAL PE.xlsx"
    wb_pe = openpyxl.Workbook()
    ws_pe = wb_pe.active
    ws_pe.title = "DataCycle1"
    ws_pe.append(["PE NO", "FL NUMBER", "SUBSTATION NAME", "DATE", "TYPE", "WO", "SCOPE"])
    for num in range(1, 10):
        ws_pe.append([
            num,
            f"CKTN-{num:03d}",
            f"PE STATION_{num}",
            datetime(2026, 5, 10),
            "P-E",
            40010000 + num,
            "FULL",
        ])
    wb_pe.save(total_pe_path)
    wb_pe.close()

    # 2. WhatsApp template
    wa_template_dir = templates_dir / "WHATSAPP"
    wa_template_dir.mkdir(parents=True, exist_ok=True)
    wa_template_file = wa_template_dir / "TEMPLATE WHATSAPP PYTHON.docx"

    repo_template = Path(__file__).parent.parent / "templates" / "WHATSAPP" / "TEMPLATE WHATSAPP PYTHON.docx"
    if repo_template.exists():
        wa_template_file.write_bytes(repo_template.read_bytes())
    else:
        doc = Document()
        doc.add_paragraph("WhatsApp Report Date: {{ date }} Station: {{ station }}")
        doc.add_paragraph("{% for item in items %}{{ item.name }} - {{ item.defect }} - {{ item.msms }}{% endfor %}")
        doc.save(wa_template_file)

    # 3. Signature image fixtures
    ali_dir = sign_dir / "ALI"
    ali_dir.mkdir(parents=True, exist_ok=True)
    ali_img_path = ali_dir / "signature.png"
    img_ali = PILImage.new("RGBA", (60, 30), color=(0, 0, 255, 255))
    img_ali.save(ali_img_path)

    bakar_dir = sign_dir / "BAKAR"
    bakar_dir.mkdir(parents=True, exist_ok=True)
    bakar_img_path = bakar_dir / "signature.png"
    img_bakar = PILImage.new("RGBA", (60, 30), color=(255, 0, 0, 255))
    img_bakar.save(bakar_img_path)

    metadata = ProjectMetadata(
        key="pahang_e2e",
        name="Pahang E2E Integration Test",
        po_number="PO-E2E-2026",
        state="pahang",
        voltage_type="11kV",
        year="2026",
        cycle="Cycle 1",
        technologies=("IR", "US", "TEV"),
        base_path=str(root_path),
    )
    storage = LocalWorkspaceStorage(root_path, templates_dir=templates_dir)
    return ProjectEnvironment(metadata=metadata, storage=storage)


def _create_sample_testsheet_workbook(
    path: Path,
    substation_number: int,
    station_name: str,
    fl_number: str,
) -> None:
    """Create a realistic openpyxl workbook with PCE Testsheet and PCE VI sheets."""
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = openpyxl.Workbook()

    # Sheet 1: PCE Testsheet
    ws1 = wb.active
    ws1.title = "PCE Testsheet"
    ws1["A1"] = "SUBSTATION"
    ws1["B1"] = station_name
    ws1["C5"] = station_name
    ws1["A2"] = "FL"
    ws1["B2"] = fl_number
    ws1["W5"] = fl_number
    ws1["Y1"] = substation_number
    ws1["P4"] = datetime(2026, 5, 10)
    ws1["D10"] = "{{signvendor}}"
    ws1["T10"] = "{{signtnb}}"
    ws1["A5"] = None  # blank cell to be diagonalized

    # Sheet 2: PCE VI
    ws2 = wb.create_sheet(title="PCE VI")
    ws2["A1"] = "VISUAL INSPECTION"
    ws2["C6"] = None  # blank cell to be diagonalized
    ws2["C10"] = "{{signvendor}}"
    ws2["K10"] = "{{signtnb}}"

    wb.save(path)
    wb.close()


def _create_substation_fixtures(
    env: ProjectEnvironment,
    date_folder: str,
    count: int = 2,
    mismatched_testsheet_name: bool = False,
    include_raw_material: bool = True,
    raw_material_count: int | None = None,
) -> tuple[list[Path], list[Path], list[Path]]:
    """Create synchronized or mismatched file fixtures for testing."""
    qr_dir = env.get_quick_report_dir() / date_folder
    ts_dir = env.get_testsheet_dir() / date_folder
    raw_dir = env.get_raw_material_dir() / date_folder

    qr_dir.mkdir(parents=True, exist_ok=True)
    ts_dir.mkdir(parents=True, exist_ok=True)
    if include_raw_material:
        raw_dir.mkdir(parents=True, exist_ok=True)

    qr_files: list[Path] = []
    ts_files: list[Path] = []
    raw_folders: list[Path] = []

    for i in range(1, count + 1):
        stem = f"0{i}. PE STATION_{i} (VI)"
        qr_file = qr_dir / f"{stem}.docx"
        doc = Document()
        doc.add_heading(f"Quick Report for PE STATION_{i}", level=1)
        doc.add_paragraph("Visual inspection and testing report details.")
        doc.save(qr_file)
        qr_files.append(qr_file)

        if mismatched_testsheet_name and i == 1:
            ts_file = ts_dir / f"01. OLD NAME MISMATCH.xlsx"
        else:
            ts_file = ts_dir / f"{stem}.xlsx"

        _create_sample_testsheet_workbook(
            ts_file,
            substation_number=i,
            station_name=f"PE STATION_{i}",
            fl_number=f"CKTN-{i:03d}",
        )
        ts_files.append(ts_file)

    if include_raw_material:
        actual_raw_count = raw_material_count if raw_material_count is not None else count
        for i in range(1, actual_raw_count + 1):
            if mismatched_testsheet_name and i == 1:
                raw_folder = raw_dir / "01. OLD RAW FOLDER"
            else:
                raw_folder = raw_dir / f"0{i}. PE STATION_{i} (VI)"
            raw_folder.mkdir(parents=True, exist_ok=True)
            raw_folders.append(raw_folder)

    return qr_files, ts_files, raw_folders


# ─── Scenario A: Full end-to-end by_date execution ───


def test_e2e_by_date_happy_path(tmp_path: Path) -> None:
    """Scenario A: Full end-to-end BY_DATE execution.

    - Real workspace setup: QUICK REPORT, TESTSHEET, RAW MATERIAL, OTHERS/SIGN, PYTHON/TOTAL PE.
    - Signatures enabled with real image files.
    - WhatsApp daily report requested and generated in PYTHON/WHATSAPP/.
    - Pre-flight passes.
    - Renaming sync synchronizes mismatched filename '01. OLD NAME MISMATCH.xlsx' -> '01. PE STATION_1 (VI).xlsx'.
    - Intermediate testsheet PDFs created in TESTSHEET/<DATE>/processed_testsheet/pdf/.
    - Processed testsheet working copies created in TESTSHEET/<DATE>/processed_testsheet/ with stamped signatures and diagonalized blanks.
    - Final deliverable combined PDFs created in QUICK REPORT/<DATE>/<STEM>.pdf.
    - Source testsheets in TESTSHEET/<DATE>/ remain strictly untouched (immutability).
    """
    env = _setup_e2e_workspace(tmp_path / "workspace_a")
    date_str = "10-08-2026"
    qr_files, ts_files, raw_folders = _create_substation_fixtures(
        env,
        date_folder=date_str,
        count=2,
        mismatched_testsheet_name=True,
    )

    vendor_sign = env.get_sign_dir() / "ALI" / "signature.png"
    tnb_sign = env.get_sign_dir() / "BAKAR" / "signature.png"
    converter = FakeDocumentConverter()
    progress_messages: list[str] = []

    request = PostProcessingRequest(
        mode=PostProcessingMode.BY_DATE,
        target_dates=(date_str,),
        apply_signatures=True,
        vendor_signature_path=vendor_sign,
        tnb_signature_path=tnb_sign,
        generate_whatsapp=True,
        converter=converter,
        progress_sink=progress_messages.append,
    )

    service = WorkflowService()
    summary = service.run_postprocessing_pipeline(env, request)

    # 1. Summary validation
    assert isinstance(summary, PostProcessingSummary)
    assert summary.is_successful is True
    assert len(summary.processed_packages) == 2
    assert len(summary.final_deliverables) == 2
    assert len(summary.failed_packages) == 0
    assert summary.duration_seconds > 0

    # 2. Renaming sync verification
    ts_dir = env.get_testsheet_dir() / date_str
    assert (ts_dir / "01. PE STATION_1 (VI).xlsx").exists()
    assert not (ts_dir / "01. OLD NAME MISMATCH.xlsx").exists()

    raw_dir = env.get_raw_material_dir() / date_str
    assert (raw_dir / "01. PE STATION_1 (VI)").exists()
    assert not (raw_dir / "01. OLD RAW FOLDER").exists()

    # 3. WhatsApp report verification
    wa_dir = env.get_whatsapp_dir()
    assert wa_dir.exists()
    wa_reports = list(wa_dir.glob("*.docx"))
    assert len(wa_reports) >= 1
    generated_wa = wa_reports[0]
    assert generated_wa.stat().st_size > 0

    # 4. Deliverable PDF verification in QUICK REPORT/<DATE>/
    qr_dir = env.get_quick_report_dir() / date_str
    for i in (1, 2):
        stem = f"0{i}. PE STATION_{i} (VI)"
        deliv_pdf = qr_dir / f"{stem}.pdf"
        assert deliv_pdf.exists()
        assert deliv_pdf in summary.final_deliverables

        # Verify PDF contains merged pages
        with open(deliv_pdf, "rb") as f:
            reader = PdfReader(f)
            assert len(reader.pages) >= 2

    # 5. Intermediate testsheet PDF verification in TESTSHEET/<DATE>/processed_testsheet/pdf/
    for i in (1, 2):
        stem = f"0{i}. PE STATION_{i} (VI)"
        ts_pdf = ts_dir / "processed_testsheet" / "pdf" / f"{stem}.pdf"
        assert ts_pdf.exists()

    # 6. Processed working copy testsheet verification
    for i in (1, 2):
        stem = f"0{i}. PE STATION_{i} (VI)"
        proc_xlsx = ts_dir / "processed_testsheet" / f"{stem}.xlsx"
        assert proc_xlsx.exists()

        wb_proc = openpyxl.load_workbook(proc_xlsx)
        ws_proc = wb_proc["PCE Testsheet"]
        # Signatures replaced: placeholder string removed and drawings added
        assert ws_proc["D10"].value != "{{signvendor}}"
        assert ws_proc["T10"].value != "{{signtnb}}"
        assert len(ws_proc._images) > 0  # openpyxl drawing images attached
        wb_proc.close()

    # 7. Testsheet Immutability: Original source files remained strictly untouched
    for i in (1, 2):
        stem = f"0{i}. PE STATION_{i} (VI)"
        orig_xlsx = ts_dir / f"{stem}.xlsx"
        assert orig_xlsx.exists()

        wb_orig = openpyxl.load_workbook(orig_xlsx)
        ws_orig = wb_orig["PCE Testsheet"]
        assert ws_orig["D10"].value == "{{signvendor}}"
        assert ws_orig["T10"].value == "{{signtnb}}"
        assert len(ws_orig._images) == 0
        wb_orig.close()


# ─── Scenario B: End-to-end by_fl execution ───


def test_e2e_by_fl_happy_path_signatures_none(tmp_path: Path) -> None:
    """Scenario B: End-to-end BY_FL execution with signatures disabled (mode='none').

    - Sets up 3 substations in date folder.
    - Targets a single substation by FL code / station name.
    - WhatsApp report is automatically skipped.
    - Signatures disabled: 'none' mode strips placeholders cleanly without inserting images.
    - Blank cells are diagonalized.
    - Merged deliverable PDF created only for the targeted substation.
    - Untargeted substations remain untouched.
    """
    env = _setup_e2e_workspace(tmp_path / "workspace_b")
    date_str = "11-08-2026"
    qr_files, ts_files, _ = _create_substation_fixtures(
        env,
        date_folder=date_str,
        count=3,
    )

    converter = FakeDocumentConverter()
    progress_messages: list[str] = []

    # Request only Station 2 by FL code
    request = PostProcessingRequest(
        mode=PostProcessingMode.BY_FL,
        target_fls=("CKTN-002",),
        apply_signatures=False,
        generate_whatsapp=True,  # Should be skipped in BY_FL mode
        converter=converter,
        progress_sink=progress_messages.append,
    )

    service = WorkflowService()
    summary = service.run_postprocessing_pipeline(env, request)

    # 1. Summary validation
    assert summary.is_successful is True
    assert len(summary.processed_packages) == 1
    assert summary.processed_packages[0].substation_number == 2
    assert len(summary.final_deliverables) == 1

    # 2. WhatsApp report was NOT generated
    wa_dir = env.get_whatsapp_dir()
    if wa_dir.exists():
        assert len(list(wa_dir.glob("*.docx"))) == 0

    # 3. Deliverable PDF created only for Station 2
    qr_dir = env.get_quick_report_dir() / date_str
    deliv_2 = qr_dir / "02. PE STATION_2 (VI).pdf"
    deliv_1 = qr_dir / "01. PE STATION_1 (VI).pdf"
    deliv_3 = qr_dir / "03. PE STATION_3 (VI).pdf"
    assert deliv_2.exists()
    assert not deliv_1.exists()
    assert not deliv_3.exists()

    # 4. Mode='none' placeholder stripping verified on working copy
    ts_dir = env.get_testsheet_dir() / date_str
    proc_xlsx_2 = ts_dir / "processed_testsheet" / "02. PE STATION_2 (VI).xlsx"
    assert proc_xlsx_2.exists()

    wb_proc = openpyxl.load_workbook(proc_xlsx_2)
    ws_proc = wb_proc["PCE Testsheet"]
    assert ws_proc["D10"].value in (None, "")
    assert ws_proc["T10"].value in (None, "")
    assert len(ws_proc._images) == 0  # No images added
    wb_proc.close()

    # 5. Source testsheet remains untouched with placeholders intact
    orig_xlsx_2 = ts_dir / "02. PE STATION_2 (VI).xlsx"
    wb_orig = openpyxl.load_workbook(orig_xlsx_2)
    ws_orig = wb_orig["PCE Testsheet"]
    assert ws_orig["D10"].value == "{{signvendor}}"
    assert ws_orig["T10"].value == "{{signtnb}}"
    wb_orig.close()


# ─── Scenario C: Pre-flight count mismatch fail-fast ───


def test_e2e_preflight_count_mismatch_halts_fail_fast(tmp_path: Path) -> None:
    """Scenario C: Pre-flight file count mismatch halts execution fail-fast.

    - Sets up 2 Quick Reports but 3 Testsheets in date folder.
    - PreFlightValidationError is raised.
    - Zero files are renamed, modified, or converted.
    - No processed_testsheet directory or PDFs created.
    """
    env = _setup_e2e_workspace(tmp_path / "workspace_c")
    date_str = "12-08-2026"

    # Create 2 Quick Reports and 2 Testsheets first
    qr_files, ts_files, _ = _create_substation_fixtures(
        env,
        date_folder=date_str,
        count=2,
    )

    # Add an extra 3rd Testsheet to cause a mismatch (2 QR vs 3 TS)
    extra_ts = env.get_testsheet_dir() / date_str / "03. PE STATION_3 (VI).xlsx"
    _create_sample_testsheet_workbook(
        extra_ts,
        substation_number=3,
        station_name="PE STATION_3",
        fl_number="CKTN-003",
    )

    converter = FakeDocumentConverter()
    request = PostProcessingRequest(
        mode=PostProcessingMode.BY_DATE,
        target_dates=(date_str,),
        converter=converter,
    )

    service = WorkflowService()

    with pytest.raises(PreFlightValidationError) as exc_info:
        service.run_postprocessing_pipeline(env, request)

    err = exc_info.value
    assert err.date_folder == date_str
    assert err.quick_report_count == 2
    assert err.testsheet_count == 3
    assert "Count mismatch detected between QUICK REPORT (2) and TESTSHEET (3)" in str(err)

    # Fail-Fast verification: zero conversion calls, no processed folders created
    assert len(converter.convert_docx_calls) == 0
    assert len(converter.convert_testsheet_calls) == 0
    assert len(converter.merge_pdfs_calls) == 0

    proc_dir = env.get_testsheet_dir() / date_str / "processed_testsheet"
    assert not proc_dir.exists()

    qr_dir = env.get_quick_report_dir() / date_str
    assert len(list(qr_dir.glob("*.pdf"))) == 0


# ─── Scenario D: Substation converter failure resilience ───


def test_e2e_substation_converter_failure_resilience(tmp_path: Path) -> None:
    """Scenario D: Per-substation COM failure resilience.

    - 3 substations in batch.
    - Converter raises an exception for Substation 2 during docx-to-pdf conversion.
    - Substations 1 and 3 complete successfully.
    - Summary reports 2 succeeded, 1 failed with detailed error message.
    """
    env = _setup_e2e_workspace(tmp_path / "workspace_d")
    date_str = "13-08-2026"
    qr_files, ts_files, _ = _create_substation_fixtures(
        env,
        date_folder=date_str,
        count=3,
    )

    class FailingOnStation2Converter(FakeDocumentConverter):
        def convert_docx_to_pdf(self, docx_path: Path, pdf_path: Path, *args, **kwargs) -> Path:
            if "STATION_2" in docx_path.name:
                raise RuntimeError("Simulated Word COM RPC Server Unavailable for STATION_2")
            return super().convert_docx_to_pdf(docx_path, pdf_path, *args, **kwargs)

    converter = FailingOnStation2Converter()
    progress_messages: list[str] = []

    request = PostProcessingRequest(
        mode=PostProcessingMode.BY_DATE,
        target_dates=(date_str,),
        apply_signatures=False,
        generate_whatsapp=False,
        converter=converter,
        progress_sink=progress_messages.append,
    )

    service = WorkflowService()
    summary = service.run_postprocessing_pipeline(env, request)

    # 1. Summary captures partial batch success and isolated failure
    assert summary.is_successful is False
    assert len(summary.processed_packages) == 2
    assert len(summary.final_deliverables) == 2
    assert len(summary.failed_packages) == 1

    failed_item = summary.failed_packages[0]
    assert isinstance(failed_item, PostProcessingFailure)
    assert failed_item.package.substation_number == 2
    assert "Simulated Word COM RPC Server Unavailable for STATION_2" in failed_item.error
    assert len(summary.errors) == 1

    # 2. Verify Deliverable PDFs for Station 1 and Station 3 exist
    qr_dir = env.get_quick_report_dir() / date_str
    assert (qr_dir / "01. PE STATION_1 (VI).pdf").exists()
    assert not (qr_dir / "02. PE STATION_2 (VI).pdf").exists()
    assert (qr_dir / "03. PE STATION_3 (VI).pdf").exists()


# ─── Scenario E: RAW MATERIAL count mismatch fail-fast ───


def test_e2e_preflight_raw_material_mismatch_halts_fail_fast(tmp_path: Path) -> None:
    """Scenario E: Pre-flight fail-fast on RAW MATERIAL count mismatch when directory exists."""
    env = _setup_e2e_workspace(tmp_path / "workspace_e")
    date_str = "14-08-2026"

    # Create 2 QR, 2 TS, but only 1 RAW MATERIAL folder
    _create_substation_fixtures(
        env,
        date_folder=date_str,
        count=2,
        include_raw_material=True,
        raw_material_count=1,
    )

    converter = FakeDocumentConverter()
    request = PostProcessingRequest(
        mode=PostProcessingMode.BY_DATE,
        target_dates=(date_str,),
        converter=converter,
    )

    service = WorkflowService()
    with pytest.raises(PreFlightValidationError) as exc_info:
        service.run_postprocessing_pipeline(env, request)

    err = exc_info.value
    assert err.quick_report_count == 2
    assert err.testsheet_count == 2
    assert err.raw_material_count == 1
    assert "Count mismatch detected between RAW MATERIAL (1) and QUICK REPORT / TESTSHEET (2)" in str(err)


# ─── Scenario F: CLI Presentation Adapter End-to-End ───


def test_e2e_cli_action_adapter_execution(tmp_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """Scenario F: PostProcessingPipelineAction execution via simulated interactive prompts."""
    env = _setup_e2e_workspace(tmp_path / "workspace_f")
    date_str = "15-08-2026"
    qr_files, ts_files, _ = _create_substation_fixtures(
        env,
        date_folder=date_str,
        count=2,
    )

    vendor_sign = env.get_sign_dir() / "ALI" / "signature.png"
    tnb_sign = env.get_sign_dir() / "BAKAR" / "signature.png"
    fake_converter = FakeDocumentConverter()

    action = PostProcessingPipelineAction(
        "Run Full Substation Post-Processing Pipeline (1-Click)",
        converter=fake_converter,
    )
    date_path = env.get_testsheet_dir() / date_str

    with (
        patch("src.cli_selectors.select_one", return_value="by_date"),
        patch("src.cli_selectors.select_pahang_date_folder", return_value=date_path),
        patch("src.cli_selectors.confirm", side_effect=[True, True]),  # Signatures=True, WhatsApp=True
        patch("src.workflows.replace_signatures._select_signature_path", side_effect=[
            (vendor_sign, "ALI"),
            (tnb_sign, "BAKAR"),
        ]),
    ):
        result = action.execute(env)

        assert isinstance(result, PostProcessingSummary)
        assert result.is_successful is True
        assert len(result.processed_packages) == 2
        assert len(result.final_deliverables) == 2

        # Check stdout captured summary box
        captured = capsys.readouterr().out
        assert "1-CLICK POST-PROCESSING PIPELINE SUMMARY" in captured
        assert "Total Queued    : 2" in captured
        assert "Succeeded       : 2" in captured
        assert "FINAL DELIVERABLES:" in captured
        assert "✓ 01. PE STATION_1 (VI).pdf" in captured
        assert "✓ 02. PE STATION_2 (VI).pdf" in captured


def test_e2e_by_date_3tier_nested_pahang_hierarchy(tmp_path: Path) -> None:
    """Verify full end-to-end execution in 3-tier Pahang directory (<STATION>/<MONTH>/<DATE>/)."""
    env = _setup_e2e_workspace(tmp_path / "workspace_nested")
    date_str = "28-08-2026"
    station = "TEMERLOH"
    month = "08. AUGUST"

    ts_date_dir = env.get_testsheet_dir() / station / month / date_str
    qr_date_dir = env.get_quick_report_dir() / station / month / date_str
    raw_date_dir = env.get_raw_material_dir() / station / month / date_str

    ts_date_dir.mkdir(parents=True, exist_ok=True)
    qr_date_dir.mkdir(parents=True, exist_ok=True)
    raw_date_dir.mkdir(parents=True, exist_ok=True)

    for i in range(1, 3):
        ts_path = ts_date_dir / f"0{i}. MISMATCH_NAME_{i}.xlsx"
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "PCE Testsheet"
        ws1["A1"] = f"Testsheet {i}"
        ws2 = wb.create_sheet("PCE VI")
        ws2["A1"] = f"VI {i}"
        wb.save(ts_path)
        wb.close()

        qr_path = qr_date_dir / f"0{i}. PE STATION_{i} (IR+VI).docx"
        doc = Document()
        doc.add_paragraph(f"Quick Report for Station {i}")
        doc.save(qr_path)

        raw_pe_dir = raw_date_dir / f"0{i}. PE STATION_{i}"
        raw_pe_dir.mkdir(parents=True, exist_ok=True)

    fake_converter = FakeDocumentConverter()
    workflow = PostProcessingPipelineWorkflow(converter=fake_converter)

    req = PostProcessingRequest(
        mode=PostProcessingMode.BY_DATE,
        target_dates=(date_str,),
        apply_signatures=False,
        generate_whatsapp=False,
    )
    result = workflow.execute(env, req)

    assert result.is_successful is True
    assert len(result.processed_packages) == 2
    assert len(result.final_deliverables) == 2

    # Renaming sync verified in nested directory
    assert (ts_date_dir / "01. PE STATION_1 (IR+VI).xlsx").exists()
    assert (ts_date_dir / "02. PE STATION_2 (IR+VI).xlsx").exists()

    # Deliverables verified in nested QR directory
    assert (qr_date_dir / "01. PE STATION_1 (IR+VI).pdf").exists()
    assert (qr_date_dir / "02. PE STATION_2 (IR+VI).pdf").exists()


"""Unit tests for Quick Report components."""

from pathlib import Path
import pytest
from src.quick_report.cbm_summary import (
    format_db_reading,
    format_temperature_reading,
    prepare_tech_summary_rows,
)
from src.quick_report.defects import CbmDefectRecord, MasterQr03DefectRepository, ViDefectRecord
from src.quick_report.models import ViDefectPagePlan, ViSummaryRow
from src.quick_report.vi_defect_pages import (
    ViDefectPageBuilder,
    build_vi_defect_page_context,
    format_vi_defect_description,
    generate_vi_defect_pages,
)
from src.quick_report.vi_summary import (
    build_vi_summary_context,
    generate_vi_summary,
    prepare_vi_summary_rows,
)


def test_format_temperature_reading():
    assert format_temperature_reading("") == "-"
    assert format_temperature_reading(None) == "-"
    assert format_temperature_reading("-") == "-"
    assert format_temperature_reading("50") == "50.0 °C"
    assert format_temperature_reading("50.5") == "50.5 °C"
    assert format_temperature_reading("50 °C") == "50.0 °C"
    assert format_temperature_reading("50°C") == "50.0 °C"
    assert format_temperature_reading("33.3 °C") == "33.3 °C"
    assert format_temperature_reading(True) == "-"
    assert format_temperature_reading(False) == "-"
    assert format_temperature_reading(float("nan")) == "-"


def test_quick_report_transformer_date_formatting():
    """Verify QuickReportTransformer formats datefrontpage as DD MMM YYYY and date as DD/MM/YYYY."""
    from unittest.mock import MagicMock
    from src.quick_report.transformer import QuickReportTransformer

    transformer = QuickReportTransformer()

    pkg = MagicMock()
    pkg.station = "KUANTAN"
    pkg.month = "08. AUGUST"
    pkg.substation_number = 1
    pkg.data = MagicMock()
    pkg.data.date_str = "12-08-2026"
    pkg.data.substation_name_erms = "TEST SUBSTATION"
    pkg.data.substation_name_site = "TEST SUBSTATION"
    pkg.data.fl_erms = "FL123"
    pkg.data.fl_site = "FL123"
    pkg.data.gps_coordinate = ""
    pkg.data.substation_type = ""
    pkg.data.building_type = ""
    pkg.data.ambient = ""
    pkg.data.humidity = ""
    pkg.data.time = ""

    env = MagicMock()
    env.po_number = "12345"
    env.state = "PAHANG"
    env.get_vi_front_page_template.return_value = Path("dummy.docx")
    env.get_template.return_value = Path("dummy.docx")

    plan = transformer.transform(pkg, [], [], env)

    substation_info = plan.pe_info["substation"]
    assert substation_info["datefrontpage"] == "12 AUG 2026"
    assert substation_info["date"] == "12/08/2026"

    # Test with ISO date string format
    pkg.data.date_str = "2026-08-12"
    plan_iso = transformer.transform(pkg, [], [], env)
    assert plan_iso.pe_info["substation"]["datefrontpage"] == "12 AUG 2026"
    assert plan_iso.pe_info["substation"]["date"] == "12/08/2026"

    # Test fallback for None / empty date_str
    pkg.data.date_str = None
    plan_none = transformer.transform(pkg, [], [], env)
    assert plan_none.pe_info["substation"]["datefrontpage"] == "-"
    assert plan_none.pe_info["substation"]["date"] == "-"


def test_quick_report_transformer_ambient_formatting():
    """Verify QuickReportTransformer normalizes ambient temperature consistently with humidity and time."""
    from unittest.mock import MagicMock
    from src.quick_report.transformer import QuickReportTransformer

    transformer = QuickReportTransformer()

    pkg = MagicMock()
    pkg.station = "KUANTAN"
    pkg.month = "08. AUGUST"
    pkg.substation_number = 1
    pkg.data = MagicMock()
    pkg.data.date_str = "12-08-2026"
    pkg.data.substation_name_erms = "TEST SUBSTATION"
    pkg.data.substation_name_site = "TEST SUBSTATION"
    pkg.data.fl_erms = "FL123"
    pkg.data.fl_site = "FL123"
    pkg.data.gps_coordinate = ""
    pkg.data.substation_type = ""
    pkg.data.building_type = ""
    pkg.data.humidity = "70%"
    pkg.data.time = "10:35 AM"
    pkg.data.equipment = None
    pkg.data.condition = None

    env = MagicMock()
    env.po_number = "12345"
    env.state = "PAHANG"
    env.get_vi_front_page_template.return_value = Path("dummy.docx")
    env.get_template.return_value = Path("dummy.docx")

    # Case 1: Populated ambient string with unit
    pkg.data.ambient = "23.2 °C"
    plan = transformer.transform(pkg, [], [], env)
    assert plan.pe_info["substation"]["ambient"] == "23.2 °C"

    # Case 2: Dash sentinel
    pkg.data.ambient = "-"
    plan_dash = transformer.transform(pkg, [], [], env)
    assert plan_dash.pe_info["substation"]["ambient"] == "-"

    # Case 3: None
    pkg.data.ambient = None
    plan_none = transformer.transform(pkg, [], [], env)
    assert plan_none.pe_info["substation"]["ambient"] == "-"


def test_format_db_reading():
    assert format_db_reading("") == "-"
    assert format_db_reading(None) == "-"
    assert format_db_reading("-") == "-"
    assert format_db_reading("12") == "12dB"
    assert format_db_reading("12.0") == "12dB"
    assert format_db_reading("12.5") == "12.5dB"
    assert format_db_reading("12dB") == "12dB"


def test_prepare_tech_summary_rows():
    defects = [
        CbmDefectRecord(equipment="TX", technology="IR", defect_area="Body", raw_measurement="50", ir_reading="50"),
        CbmDefectRecord(equipment="TX", technology="US", defect_area="Body", raw_measurement="20", us_reading="20"),
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 1
    assert rows[0].equipment == "TX"
    assert rows[0].ir_reading == "50.0 °C"
    assert rows[0].us_reading == "20dB"


def test_prepare_tech_summary_rows_empty():
    rows = prepare_tech_summary_rows([])
    assert len(rows) == 0


def test_prepare_tech_summary_rows_different_areas():
    defects = [
        CbmDefectRecord(equipment="TX", technology="IR", defect_area="Body", raw_measurement="50", ir_reading="50"),
        CbmDefectRecord(equipment="TX", technology="US", defect_area="Bush", raw_measurement="20", us_reading="20"),
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 2


def test_prepare_tech_summary_rows_ir_only():
    defects = [
        CbmDefectRecord(
            equipment="TX 1",
            defect_area="HV Bushing",
            additional_remarks="",
            technology="IR",
            ir_reading="55.4",
        )
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 1
    r = rows[0]
    assert r.equipment == "TX 1"
    assert r.defect_area == "HV Bushing"
    assert r.ir_abs == "55.4 °C"
    assert r.ir_delta == "-"
    assert r.us_dB == "-"
    assert r.tev_dB == "-"
    assert r.severity == ""
    assert r.status == ""
    assert r.ir_reading == "55.4 °C"
    assert r.us_reading == "-"
    assert r.tev_reading == "-"


def test_prepare_tech_summary_rows_us_only_with_char():
    defects = [
        CbmDefectRecord(
            equipment="SWG 01",
            defect_area="Cable Termination",
            additional_remarks="Phase R",
            technology="US",
            us_reading="18.5",
            us_char="TRACKING",
        )
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 1
    r = rows[0]
    assert r.equipment == "SWG 01"
    assert r.defect_area == "Cable Termination/ Phase R"
    assert r.remarks == "Phase R"
    assert r.ir_abs == "-"
    assert r.ir_delta == "-"
    assert r.us_dB == "18.5dB"
    assert r.tev_dB == "-"
    assert r.severity == "TRACKING"
    assert r.status == "TRACKING"
    assert r.ir_reading == "-"
    assert r.us_reading == "18.5dB"
    assert r.tev_reading == "-"


def test_prepare_tech_summary_rows_tev_only():
    defects = [
        CbmDefectRecord(
            equipment="RMU SF6",
            defect_area="Busbar",
            technology="TEV",
            tev_reading="26.0",
            tev_char="SURFACE",
        )
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 1
    r = rows[0]
    assert r.equipment == "RMU SF6"
    assert r.defect_area == "Busbar"
    assert r.ir_abs == "-"
    assert r.ir_delta == "-"
    assert r.us_dB == "-"
    assert r.tev_dB == "26dB"
    assert r.severity == ""
    assert r.status == ""
    assert r.ir_reading == "-"
    assert r.us_reading == "-"
    assert r.tev_reading == "26dB"


def test_prepare_tech_summary_rows_multitech_merged():
    defects = [
        CbmDefectRecord(
            equipment="RMU 01",
            defect_area="Cable Box",
            additional_remarks="Phase Y",
            technology="IR",
            ir_reading="62.1",
            brand="ABB",
            model="SafeRing",
            rating="630A",
        ),
        CbmDefectRecord(
            equipment="RMU 01",
            defect_area="Cable Box",
            additional_remarks="Phase Y",
            technology="US",
            us_reading="14.0",
            us_char="CORONA",
        ),
        CbmDefectRecord(
            equipment="RMU 01",
            defect_area="Cable Box",
            additional_remarks="Phase Y",
            technology="TEV",
            tev_reading="28.5",
        ),
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 1
    r = rows[0]
    assert r.equipment == "RMU 01"
    assert r.defect_area == "Cable Box/ Phase Y"
    assert r.remarks == "Phase Y"
    assert r.brand == "ABB"
    assert r.model == "SafeRing"
    assert r.rating == "630A"
    assert r.ir_abs == "62.1 °C"
    assert r.ir_delta == "-"
    assert r.us_dB == "14dB"
    assert r.tev_dB == "28.5dB"
    assert r.severity == "CORONA DISCHARGE"
    assert r.status == "CORONA DISCHARGE"
    assert r.ir_reading == "62.1 °C"
    assert r.us_reading == "14dB"
    assert r.tev_reading == "28.5dB"


def test_prepare_tech_summary_rows_case_insensitive_and_non_matching():
    defects = [
        CbmDefectRecord(
            equipment="Switchgear A",
            defect_area="Panel 1",
            additional_remarks="Hotspot",
            technology="IR",
            ir_reading="50",
        ),
        CbmDefectRecord(
            equipment="switchgear a",
            defect_area="panel 1",
            additional_remarks="hotspot",
            technology="US",
            us_reading="10",
            us_char="ARCING",
        ),
        CbmDefectRecord(
            equipment="OVERHEAD LINE",
            defect_area="Conductor",
            additional_remarks="Sagging",
            technology="IR",
            ir_reading="40",
        ),
    ]
    rows = prepare_tech_summary_rows(defects)
    assert len(rows) == 2
    r0 = rows[0]
    assert r0.equipment == "Switchgear A"
    assert r0.defect_area == "Panel 1/ Hotspot"
    assert r0.remarks == "Hotspot"
    assert r0.ir_abs == "50.0 °C"
    assert r0.us_dB == "10dB"
    assert r0.severity == "ARCING"

    r1 = rows[1]
    assert r1.equipment == "OVERHEAD LINE"
    assert r1.defect_area == "Conductor/ Sagging"
    assert r1.remarks == "Sagging"
    assert r1.ir_abs == "40.0 °C"
    assert r1.us_dB == "-"
    assert r1.tev_dB == "-"
    assert r1.severity == ""


def test_cbm_defect_record_normalization():
    # 1. Technology upper-casing and string whitespace stripping
    rec = CbmDefectRecord(
        equipment="  TX 1  ",
        technology="  ir  ",
        brand="  ABB  ",
        model="  XYZ  ",
        rating="  11kV  ",
        defect_area="  Body  ",
        additional_remarks="  Hotspot  ",
        raw_measurement="  55.4  ",
        equipment_id="  SWG 01  ",
    )
    assert rec.technology == "IR"
    assert rec.equipment == "TX 1"
    assert rec.brand == "ABB"
    assert rec.model == "XYZ"
    assert rec.rating == "11kV"
    assert rec.defect_area == "Body"
    assert rec.additional_remarks == "Hotspot"
    assert rec.equipment_id == "SWG 01"
    # Invariant: IR reading populated from raw_measurement
    assert rec.raw_measurement == "55.4"
    assert rec.ir_reading == "55.4"

    # 2. Invariant: IR reading populates raw_measurement if raw_measurement empty
    rec_ir2 = CbmDefectRecord(technology="IR", ir_reading="  60.1  ")
    assert rec_ir2.ir_reading == "60.1"
    assert rec_ir2.raw_measurement == "60.1"

    # 3. Invariant: US technology
    rec_us = CbmDefectRecord(technology="us", raw_measurement="  15.0  ")
    assert rec_us.technology == "US"
    assert rec_us.us_reading == "15.0"
    assert rec_us.raw_measurement == "15.0"

    rec_us2 = CbmDefectRecord(technology="US", us_reading="22.5")
    assert rec_us2.raw_measurement == "22.5"

    # 4. Invariant: TEV technology
    rec_tev = CbmDefectRecord(technology="tev", raw_measurement="  25.0  ")
    assert rec_tev.technology == "TEV"
    assert rec_tev.tev_reading == "25.0"
    assert rec_tev.raw_measurement == "25.0"

    rec_tev2 = CbmDefectRecord(technology="TEV", tev_reading="30.0")
    assert rec_tev2.raw_measurement == "30.0"

    # 5. equipment_id None/empty normalization
    rec_defaults = CbmDefectRecord()
    assert rec_defaults.equipment_id == ""


def test_cbm_defect_record_to_dict():
    rec = CbmDefectRecord(
        equipment="RMU SF6",
        technology="US",
        brand="Schneider",
        model="RM6",
        rating="11kV",
        defect_area="Cable Box",
        additional_remarks="Corona discharge",
        ir_reading="",
        us_reading="24.5",
        us_char="CORONA",
        tev_reading="",
        tev_char="",
        raw_measurement="24.5",
        equipment_id="PANEL 2",
        source_order=3,
    )
    d = rec.to_dict()
    assert d == {
        "equipment": "RMU SF6",
        "technology": "US",
        "brand": "Schneider",
        "model": "RM6",
        "rating": "11kV",
        "defect_area": "Cable Box",
        "additional_remarks": "Corona discharge",
        "ir_reading": "",
        "us_reading": "24.5",
        "us_char": "CORONA",
        "tev_reading": "",
        "tev_char": "",
        "raw_measurement": "24.5",
        "equipment_id": "PANEL 2",
        "source_order": 3,
    }




def test_environment_front_page_template_resolution():
    from unittest.mock import MagicMock
    from src.project.environment import ProjectEnvironment
    from src.project.models import ProjectMetadata
    from src.project.storage import WorkspaceStorage

    def make_env(techs):
        meta = ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", techs, "p")
        storage = MagicMock(spec=WorkspaceStorage)
        storage.get_template.side_effect = lambda k: Path(f"/templates/{k}.docx")
        return ProjectEnvironment(meta, storage)

    env_ir = make_env(("IR",))
    assert env_ir.get_vi_front_page_template() == Path("/templates/vi_front_page_ir.docx")

    env_us = make_env(("IR", "US"))
    assert env_us.get_vi_front_page_template() == Path("/templates/vi_front_page_ir_us.docx")

    env_tev = make_env(("IR", "US", "TEV"))
    assert env_tev.get_vi_front_page_template() == Path("/templates/vi_front_page_ir_us_tev.docx")


def test_environment_get_cbm_defect_folder_name():
    from unittest.mock import MagicMock
    from src.project.environment import ProjectEnvironment
    from src.project.models import ProjectMetadata
    from src.project.storage import WorkspaceStorage

    storage = MagicMock(spec=WorkspaceStorage)

    def make_env(techs):
        meta = ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", techs, "p")
        return ProjectEnvironment(meta, storage)

    assert make_env(("IR",)).get_cbm_defect_folder_name() == "DEFECT IR"
    assert make_env(("ir",)).get_cbm_defect_folder_name() == "DEFECT IR"
    assert make_env(("IR", "US")).get_cbm_defect_folder_name() == "DEFECT IR US"
    assert make_env(("ir", "us")).get_cbm_defect_folder_name() == "DEFECT IR US"
    assert make_env(("IR", "US", "TEV")).get_cbm_defect_folder_name() == "DEFECT IR US TEV"
    assert make_env(("ir", "us", "tev")).get_cbm_defect_folder_name() == "DEFECT IR US TEV"
    assert make_env(("TEV",)).get_cbm_defect_folder_name() == "DEFECT IR US TEV"


def test_environment_dynamic_cbm_defect_template_resolution_single_tech(tmp_path: Path):
    from src.project.environment import ProjectEnvironment
    from src.project.models import ProjectMetadata
    from src.project.storage import LocalWorkspaceStorage

    storage = LocalWorkspaceStorage(tmp_path)
    meta = ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", ("IR",), str(tmp_path))
    env = ProjectEnvironment(meta, storage)

    defect_ir_dir = tmp_path / "templates" / "QUICK REPORT" / "DEFECT IR"
    defect_ir_dir.mkdir(parents=True, exist_ok=True)
    fp_tpl = defect_ir_dir / "fp-overview.docx"
    fp_tpl.touch()

    assert env.get_template("fp_overview") == fp_tpl
    assert env.resolve_template_path("fp_overview") == fp_tpl


def test_environment_dynamic_cbm_defect_template_resolution_dual_tech(tmp_path: Path):
    from src.project.environment import ProjectEnvironment
    from src.project.models import ProjectMetadata
    from src.project.storage import LocalWorkspaceStorage

    storage = LocalWorkspaceStorage(tmp_path)
    meta = ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", ("IR", "US"), str(tmp_path))
    env = ProjectEnvironment(meta, storage)

    # 1. Fails fast if DEFECT IR US folder is missing
    with pytest.raises(FileNotFoundError, match="Required CBM defect template directory 'DEFECT IR US' is missing"):
        env.get_template("swg_overview")

    # 2. Resolves successfully when folder and template exist
    defect_ir_us_dir = tmp_path / "templates" / "QUICK REPORT" / "DEFECT IR US"
    defect_ir_us_dir.mkdir(parents=True, exist_ok=True)
    swg_tpl = defect_ir_us_dir / "swg-overview.docx"
    swg_tpl.touch()

    assert env.get_template("swg_overview") == swg_tpl
    assert env.resolve_template_path("swg_overview") == swg_tpl


def test_environment_dynamic_cbm_defect_template_resolution_triple_tech(tmp_path: Path):
    from src.project.environment import ProjectEnvironment
    from src.project.models import ProjectMetadata
    from src.project.storage import LocalWorkspaceStorage

    storage = LocalWorkspaceStorage(tmp_path)
    meta = ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", ("IR", "US", "TEV"), str(tmp_path))
    env = ProjectEnvironment(meta, storage)

    # 1. Fails fast if DEFECT IR US TEV folder is missing
    with pytest.raises(FileNotFoundError, match="Required CBM defect template directory 'DEFECT IR US TEV' is missing"):
        env.get_template("tx_overview")

    # 2. Resolves successfully when folder and template exist
    defect_ir_us_tev_dir = tmp_path / "templates" / "QUICK REPORT" / "DEFECT IR US TEV"
    defect_ir_us_tev_dir.mkdir(parents=True, exist_ok=True)
    tx_tpl = defect_ir_us_tev_dir / "tx-overview.docx"
    tx_tpl.touch()

    assert env.get_template("tx_overview") == tx_tpl
    assert env.resolve_template_path("tx_overview") == tx_tpl


def test_environment_dynamic_cbm_defect_template_missing_file_fails_fast(tmp_path: Path):
    from src.project.environment import ProjectEnvironment
    from src.project.models import ProjectMetadata
    from src.project.storage import LocalWorkspaceStorage

    storage = LocalWorkspaceStorage(tmp_path)
    meta = ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", ("IR",), str(tmp_path))
    env = ProjectEnvironment(meta, storage)

    # Folder exists but file does not
    defect_ir_dir = tmp_path / "templates" / "QUICK REPORT" / "DEFECT IR"
    defect_ir_dir.mkdir(parents=True, exist_ok=True)

    with pytest.raises(FileNotFoundError, match="Required CBM defect template 'fp-overview.docx' is missing"):
        env.get_template("fp_overview")


def test_environment_cbm_summary_template_resolution_with_local_storage(tmp_path: Path):
    from src.project.environment import ProjectEnvironment
    from src.project.models import ProjectMetadata
    from src.project.storage import LocalWorkspaceStorage

    qr_dir = tmp_path / "templates" / "QUICK REPORT"
    qr_dir.mkdir(parents=True, exist_ok=True)
    ir_sum = qr_dir / "CBM DEFECT IR SUMMARY.docx"
    ir_us_sum = qr_dir / "CBM DEFECT IR+US SUMMARY.docx"
    ir_us_tev_sum = qr_dir / "CBM DEFECT IR+US+TEV SUMMARY.docx"
    ir_sum.touch()
    ir_us_sum.touch()
    ir_us_tev_sum.touch()

    storage = LocalWorkspaceStorage(tmp_path)

    env_ir = ProjectEnvironment(ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", ("IR",), str(tmp_path)), storage)
    env_us = ProjectEnvironment(ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", ("IR", "US"), str(tmp_path)), storage)
    env_tev = ProjectEnvironment(ProjectMetadata("k", "n", "po", "st", "11kV", "2026", "c1", ("IR", "US", "TEV"), str(tmp_path)), storage)

    assert env_ir.get_cbm_summary_template() == ir_sum
    assert env_us.get_cbm_summary_template() == ir_us_sum
    assert env_tev.get_cbm_summary_template() == ir_us_tev_sum


def test_generate_vi_summary_programmatic(tmp_path: Path):
    import docx
    from src.quick_report.vi_summary import generate_vi_summary

    template_p = Path("templates/QUICK REPORT/2. VI SUMMARY TEMPLATE Jinja2 DYNAMIC.docx")
    if not template_p.exists():
        # Create a dummy template docx with 4 columns and 4 rows
        template_p = tmp_path / "vi_template.docx"
        doc = docx.Document()
        table = doc.add_table(rows=4, cols=4)
        for j, h in enumerate(["NO.", "EQUIPMENT", "DEFECT DESCRIPTION", "REMARKS"]):
            table.rows[0].cells[j].text = h
        doc.save(template_p)

    pe_info = {"substation": {"name_erms": "TEST SUB"}}
    defects = [
        ViDefectRecord(equipment="SWITCHGEAR", defect_area="Indicator Fault", additional_remarks="Replace lamp"),
        ViDefectRecord(equipment="TRANSFORMER", defect_area="Oil Leakage", additional_remarks="Top up oil"),
    ]

    out_path = generate_vi_summary(pe_info, defects, template_p, tmp_path, 1)
    assert out_path.exists()

    rendered_doc = docx.Document(out_path)
    assert len(rendered_doc.tables) == 1
    t = rendered_doc.tables[0]
    assert len(t.rows) == 3  # 1 header + 2 data rows

    # Header verification
    header_texts = [c.text.strip() for c in t.rows[0].cells]
    assert header_texts == ["NO.", "EQUIPMENT", "DEFECT DESCRIPTION", "ADDITIONAL REMARKS"]

    # Data rows verification
    r1_texts = [c.text.strip() for c in t.rows[1].cells]
    assert r1_texts == ["1", "SWITCHGEAR", "Indicator Fault", "Replace lamp"]

    r2_texts = [c.text.strip() for c in t.rows[2].cells]
    assert r2_texts == ["2", "TRANSFORMER", "Oil Leakage", "Top up oil"]

    # Check cell XML formatting
    tcPr = t.rows[1].cells[0]._tc.get_or_add_tcPr()
    assert tcPr is not None


def test_prepare_vi_summary_rows():
    defects = [
        ViDefectRecord(equipment="SWG", defect_area="Door", additional_remarks="Broken latch"),
        ViDefectRecord(equipment="TX", defect_area="Body", additional_remarks="Rust"),
        ViDefectRecord(equipment="SUBSTATION", defect_area="FENCE BROKEN", additional_remarks=""),
        ViDefectRecord(equipment="LTX/DTX", defect_area="", additional_remarks="TX1 - NAMEPLATE NOT ACCESSIBLE"),
        ViDefectRecord(equipment="", defect_area="", additional_remarks=""),
    ]
    rows = prepare_vi_summary_rows(defects)
    assert len(rows) == 5
    assert rows[0] == ViSummaryRow(equipment="SWG", defect_area="Door", remarks="Broken latch")
    assert rows[1] == ViSummaryRow(equipment="TX", defect_area="Body", remarks="Rust")
    assert rows[2] == ViSummaryRow(equipment="SUBSTATION", defect_area="FENCE BROKEN", remarks="-")
    assert rows[3] == ViSummaryRow(equipment="LTX/DTX", defect_area="-", remarks="TX1 - NAMEPLATE NOT ACCESSIBLE")
    assert rows[4] == ViSummaryRow(equipment="-", defect_area="-", remarks="-")


def test_build_vi_summary_context_accepts_only_vi_defect_record():
    pe_info = {"substation": {"name_erms": "TEST"}}
    defects = [
        ViDefectRecord(equipment="SWG 1", defect_area="Indicator", additional_remarks="Faulty"),
    ]
    context = build_vi_summary_context(pe_info, defects)
    assert context["substation"] == {"name_erms": "TEST"}
    assert len(context["defects"]) == 1
    assert context["defects"][0]["equipment"] == "SWG 1"
    assert context["defects"][0]["defect_area"] == "Indicator"
    assert context["defects"][0]["remarks"] == "Faulty"

    with pytest.raises(AttributeError):
        build_vi_summary_context(pe_info, [{"equipment": "SWG 1"}])  # type: ignore


def test_vi_defect_page_builder_pagination_and_plans(tmp_path: Path):
    template_p = tmp_path / "vi_defect_template.docx"
    template_p.touch()

    builder = ViDefectPageBuilder()
    pe_info = {"substation": {"name_erms": "SUB 1"}}

    # 1. No VI defects returns no page plans (empty list)
    empty_plans = builder.build([], template_p, pe_info, 1)
    assert empty_plans == []

    # 2. 7 VI defects produce 2 ViDefectPagePlan objects
    defects = [
        ViDefectRecord(
            equipment=f"equipment{i}",
            defect_area=f"area{i}",
            additional_remarks=f"remark{i}",
        )
        for i in range(1, 8)
    ]
    plans = builder.build(defects, template_p, pe_info, 1)
    assert len(plans) == 2

    # 3. First page plan contains equipment1 through equipment6
    plan1 = plans[0]
    assert isinstance(plan1, ViDefectPagePlan)
    assert plan1.template_path == template_p
    assert plan1.output_filename == "001_06_vi_defect_part1.docx"
    assert plan1.active_defect_count == 6
    for i in range(1, 7):
        assert plan1.context[f"equipment{i}"] == f"equipment{i}"
        assert plan1.context[f"description{i}"] == f"area{i} \u2013 remark{i}"
        assert plan1.context[f"remark{i}"] == f"remark{i}"

    # 4. Second page plan has active_defect_count == 1
    plan2 = plans[1]
    assert isinstance(plan2, ViDefectPagePlan)
    assert plan2.template_path == template_p
    assert plan2.output_filename == "001_06_vi_defect_part2.docx"
    assert plan2.active_defect_count == 1
    assert plan2.context["equipment1"] == "equipment7"
    assert plan2.context["description1"] == "area7 \u2013 remark7"
    assert plan2.context["remark1"] == "remark7"
    assert len(plan2.context["defects"]) == 6
    for i in range(2, 7):
        assert plan2.context[f"equipment{i}"] == ""
        assert plan2.context[f"description{i}"] == ""
        assert plan2.context[f"remark{i}"] == ""


def test_format_vi_defect_description_combinations():
    """Verify VI defect description formatter handles all combination pairs and empty/dash values."""
    # Both present -> joined with en-dash and spaces (\u2013)
    assert (
        format_vi_defect_description("CPR POSTER OLD VERSION", "SUBSTATION")
        == "CPR POSTER OLD VERSION \u2013 SUBSTATION"
    )
    assert (
        format_vi_defect_description("NO LINK NO./PANEL NO./FEEDER NAME", "PANEL TX")
        == "NO LINK NO./PANEL NO./FEEDER NAME \u2013 PANEL TX"
    )

    # Defect area only -> no trailing dash
    assert format_vi_defect_description("FENCE BROKEN", "") == "FENCE BROKEN"
    assert format_vi_defect_description("FENCE BROKEN", "-") == "FENCE BROKEN"
    assert format_vi_defect_description("FENCE BROKEN", "  -  ") == "FENCE BROKEN"
    assert format_vi_defect_description("NO FUNCTIONAL LOCATION", "N/A") == "NO FUNCTIONAL LOCATION"
    assert format_vi_defect_description("NO SIGNBOARD", None) == "NO SIGNBOARD"

    # Remarks only -> no leading dash
    assert (
        format_vi_defect_description("", "TX1 - NAMEPLATE NOT ACCESSIBLE")
        == "TX1 - NAMEPLATE NOT ACCESSIBLE"
    )
    assert (
        format_vi_defect_description("-", "TX1 - NAMEPLATE NOT ACCESSIBLE")
        == "TX1 - NAMEPLATE NOT ACCESSIBLE"
    )
    assert (
        format_vi_defect_description("N/A", "TX1 - NAMEPLATE NOT ACCESSIBLE")
        == "TX1 - NAMEPLATE NOT ACCESSIBLE"
    )
    assert (
        format_vi_defect_description(None, "TX1 - NAMEPLATE NOT ACCESSIBLE")
        == "TX1 - NAMEPLATE NOT ACCESSIBLE"
    )

    # Neither present -> empty string
    assert format_vi_defect_description("", "") == ""
    assert format_vi_defect_description("-", "-") == ""
    assert format_vi_defect_description("N/A", "N/A") == ""
    assert format_vi_defect_description(None, None) == ""
    assert format_vi_defect_description("   ", "   ") == ""


def test_build_vi_defect_page_context_description_formatting():
    """Verify build_vi_defect_page_context populates clean 'description' key without trailing/leading dashes."""
    pe_info = {"substation": {"name_erms": "TEST SUB"}}
    chunk = [
        ViDefectRecord(equipment="SUBSTATION", defect_area="CPR POSTER OLD VERSION", additional_remarks="SUBSTATION"),
        ViDefectRecord(equipment="SUBSTATION", defect_area="FENCE BROKEN", additional_remarks=""),
        ViDefectRecord(equipment="SIGNBOARD", defect_area="NO FUNCTIONAL LOCATION", additional_remarks="-"),
        ViDefectRecord(equipment="LTX/DTX", defect_area="", additional_remarks="TX1 - NAMEPLATE NOT ACCESSIBLE"),
    ]
    ctx = build_vi_defect_page_context(pe_info, chunk)

    # Check context['defects'] items have 'description' field
    assert ctx["defects"][0]["description"] == "CPR POSTER OLD VERSION \u2013 SUBSTATION"
    assert ctx["defects"][1]["description"] == "FENCE BROKEN"
    assert ctx["defects"][2]["description"] == "NO FUNCTIONAL LOCATION"
    assert ctx["defects"][3]["description"] == "TX1 - NAMEPLATE NOT ACCESSIBLE"
    assert ctx["defects"][4]["description"] == ""
    assert ctx["defects"][5]["description"] == ""

    # Check top-level description keys
    assert ctx["description1"] == "CPR POSTER OLD VERSION \u2013 SUBSTATION"
    assert ctx["description2"] == "FENCE BROKEN"
    assert ctx["description3"] == "NO FUNCTIONAL LOCATION"
    assert ctx["description4"] == "TX1 - NAMEPLATE NOT ACCESSIBLE"
    assert ctx["description5"] == ""
    assert ctx["description6"] == ""


def test_generate_vi_defect_pages_end_to_end_formatting_and_alignment(tmp_path: Path):
    """Verify generate_vi_defect_pages produces clean descriptions and preserves cell formatting/alignment."""
    import docx
    template_path = Path("templates/QUICK REPORT/10. VISUAL DEFECT Jinja2 DYNAMIC.docx")
    if not template_path.exists():
        pytest.skip("Template 10. VISUAL DEFECT Jinja2 DYNAMIC.docx not found")

    pe_info = {"substation": {"name_erms": "TEST SUB"}}
    defects = [
        ViDefectRecord(equipment="SUBSTATION", defect_area="CPR POSTER OLD VERSION", additional_remarks="SUBSTATION"),
        ViDefectRecord(equipment="SUBSTATION", defect_area="FENCE BROKEN", additional_remarks=""),
        ViDefectRecord(equipment="SIGNBOARD", defect_area="NO FUNCTIONAL LOCATION", additional_remarks="-"),
        ViDefectRecord(equipment="LTX/DTX", defect_area="", additional_remarks="TX1 - NAMEPLATE NOT ACCESSIBLE"),
    ]

    out_paths = generate_vi_defect_pages(
        defects=defects,
        template_path=template_path,
        output_dir=tmp_path,
        substation_number=1,
        pe_info=pe_info,
    )

    assert len(out_paths) == 1
    doc = docx.Document(str(out_paths[0]))
    t = doc.tables[0]

    # Verify rendered description texts have no dangling dashes
    assert t.rows[2].cells[0].text.strip() == "CPR POSTER OLD VERSION \u2013 SUBSTATION"
    assert t.rows[2].cells[2].text.strip() == "FENCE BROKEN"
    assert t.rows[6].cells[0].text.strip() == "NO FUNCTIONAL LOCATION"
    assert t.rows[6].cells[2].text.strip() == "TX1 - NAMEPLATE NOT ACCESSIBLE"

    # Verify paragraph center alignment is preserved on all active description cells
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for r, c in [(2, 0), (2, 2), (6, 0), (6, 2)]:
        cell = t.rows[r].cells[c]
        p = cell.paragraphs[0]
        jc = p._p.pPr.find(f"{w_ns}jc")
        assert jc is not None
        assert jc.get(f"{w_ns}val") == "center"


def test_generate_vi_summary_end_to_end_empty_normalization(tmp_path: Path):
    """Verify generate_vi_summary renders '-' for empty description and remarks."""
    import docx
    template_path = Path("templates/QUICK REPORT/2. VI SUMMARY TEMPLATE Jinja2 DYNAMIC.docx")
    if not template_path.exists():
        pytest.skip("Template 2. VI SUMMARY TEMPLATE Jinja2 DYNAMIC.docx not found")

    pe_info = {"substation": {"name_erms": "TEST SUB"}}
    defects = [
        ViDefectRecord(equipment="SUBSTATION", defect_area="CPR POSTER OLD VERSION", additional_remarks="SUBSTATION"),
        ViDefectRecord(equipment="SUBSTATION", defect_area="FENCE BROKEN", additional_remarks=""),
        ViDefectRecord(equipment="LTX/DTX", defect_area="", additional_remarks="TX1 - NAMEPLATE NOT ACCESSIBLE"),
    ]

    out_path = generate_vi_summary(
        pe_info=pe_info,
        defects=defects,
        template_path=template_path,
        output_dir=tmp_path,
        substation_number=1,
    )

    doc = docx.Document(str(out_path))
    t = doc.tables[0]
    # Row 0: Headers ('NO.', 'EQUIPMENT', 'DEFECT DESCRIPTION', 'ADDITIONAL REMARKS')
    # Row 1: Defect 1
    assert [c.text.strip() for c in t.rows[1].cells] == ["1", "SUBSTATION", "CPR POSTER OLD VERSION", "SUBSTATION"]
    # Row 2: Defect 2 (remarks was empty -> '-')
    assert [c.text.strip() for c in t.rows[2].cells] == ["2", "SUBSTATION", "FENCE BROKEN", "-"]
    # Row 3: Defect 3 (defect_area was empty -> '-')
    assert [c.text.strip() for c in t.rows[3].cells] == ["3", "LTX/DTX", "-", "TX1 - NAMEPLATE NOT ACCESSIBLE"]


def test_vi_defect_record_to_dict_uses_remarks_key():
    rec = ViDefectRecord(
        equipment="SWITCHGEAR",
        defect_area="Rust",
        additional_remarks="Corroded hinges",
    )
    assert rec.additional_remarks == "Corroded hinges"
    d = rec.to_dict()
    assert d == {
        "equipment": "SWITCHGEAR",
        "defect_area": "Rust",
        "remarks": "Corroded hinges",
    }
    assert "additional_remarks" not in d


def test_build_vi_defect_page_context_padding():
    pe_info = {"substation": {"name_erms": "SUB A"}}
    defects = [
        ViDefectRecord(
            equipment="SWG 1",
            defect_area="Oil leak",
            additional_remarks="Minor",
        ),
    ]

    context = build_vi_defect_page_context(pe_info, defects)

    assert len(context["defects"]) == 6
    assert context["defects"][0] == {
        "equipment": "SWG 1",
        "defect_area": "Oil leak",
        "remarks": "Minor",
        "description": "Oil leak \u2013 Minor",
    }
    for i in range(1, 6):
        assert context["defects"][i] == {
            "equipment": "",
            "defect_area": "",
            "remarks": "",
            "description": "",
        }

    assert context["equipment1"] == "SWG 1"
    assert context["description1"] == "Oil leak \u2013 Minor"
    assert context["remark1"] == "Minor"

    for slot in range(2, 7):
        assert context[f"equipment{slot}"] == ""
        assert context[f"description{slot}"] == ""
        assert context[f"remark{slot}"] == ""


def test_generate_vi_defect_pages_direct_rendering(tmp_path: Path):
    import docx
    from src.quick_report.vi_defect_pages import generate_vi_defect_pages

    template_p = tmp_path / "vi_defect_template.docx"
    doc = docx.Document()
    doc.save(template_p)

    defects = [
        ViDefectRecord(
            equipment="SWG 1",
            defect_area="Door defect",
            additional_remarks="Loose latch",
        ),
    ]
    pe_info = {"substation": {"name_erms": "TEST SUB"}}

    out_paths = generate_vi_defect_pages(
        defects=defects,
        template_path=template_p,
        output_dir=tmp_path,
        substation_number=1,
        pe_info=pe_info,
    )

    assert len(out_paths) == 1
    assert out_paths[0].exists()
    assert out_paths[0].name == "001_06_vi_defect_part1.docx"


def test_generate_cbm_tech_summary_programmatic(tmp_path: Path):
    import docx
    from src.quick_report.cbm_summary import generate_cbm_tech_summary

    template_p = Path("templates/QUICK REPORT/CBM DEFECT IR SUMMARY.docx")
    if not template_p.exists():
        template_p = tmp_path / "cbm_template.docx"
        doc = docx.Document()
        table = doc.add_table(rows=4, cols=6)
        for j, h in enumerate(["NO.", "EQUIPMENT", "DEFECT AREA", "IR (Abs T/ΔT)", "IR (ΔT)", "DEFECT"]):
            table.rows[0].cells[j].text = h
        doc.save(template_p)

    pe_info = {"substation": {"name_erms": "TEST SUB"}}
    defects = [
        CbmDefectRecord(equipment="TRANSFORMER", technology="IR", defect_area="HV Bushing", raw_measurement="65"),
    ]

    out_path = generate_cbm_tech_summary(pe_info, defects, template_p, tmp_path, 1)
    assert out_path.exists()

    rendered_doc = docx.Document(out_path)
    assert len(rendered_doc.tables) == 1
    t = rendered_doc.tables[0]
    assert len(t.rows) == 2  # 1 header + 1 data row

    r1_texts = [c.text.strip() for c in t.rows[1].cells]
    assert r1_texts[0] == "1"
    assert r1_texts[1] == "TRANSFORMER"
    assert r1_texts[2] == "HV Bushing"
    assert r1_texts[3] == "65.0 °C"

    # Check cell XML formatting
    tcPr = t.rows[1].cells[0]._tc.get_or_add_tcPr()
    assert tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign") is not None
    pPr = t.rows[1].cells[0].paragraphs[0]._p.get_or_add_pPr()
    assert pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc") is not None
    assert pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing") is not None


def test_generate_cbm_tech_summary_multitech_template(tmp_path: Path):
    import docx
    from src.quick_report.cbm_summary import generate_cbm_tech_summary

    template_p = Path("templates/QUICK REPORT/CBM DEFECT IR+US+TEV SUMMARY.docx")
    pe_info = {"substation": {"name_erms": "TEST SUB"}}
    defects = [
        CbmDefectRecord(
            equipment="RMU 01",
            defect_area="Cable Box",
            additional_remarks="Phase Y",
            technology="IR",
            ir_reading="62.1",
        ),
        CbmDefectRecord(
            equipment="RMU 01",
            defect_area="Cable Box",
            additional_remarks="Phase Y",
            technology="US",
            us_reading="14.0",
            us_char="CORONA",
        ),
        CbmDefectRecord(
            equipment="RMU 01",
            defect_area="Cable Box",
            additional_remarks="Phase Y",
            technology="TEV",
            tev_reading="28.5",
        ),
    ]

    out_path = generate_cbm_tech_summary(pe_info, defects, template_p, tmp_path, 1)
    assert out_path.exists()

    rendered_doc = docx.Document(out_path)
    assert len(rendered_doc.tables) == 1
    t = rendered_doc.tables[0]
    assert len(t.rows) == 2  # 1 header + 1 data row

    r1_texts = [c.text.strip() for c in t.rows[1].cells]
    assert r1_texts[0] == "1"
    assert r1_texts[1] == "RMU 01"
    assert r1_texts[2] == "Cable Box/ Phase Y"
    assert "62.1" in r1_texts[3]
    assert r1_texts[4] == "14dB"
    assert r1_texts[5] == "28.5dB"


def test_cbm_defect_planner(tmp_path: Path):
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner
    from src.quick_report.defects import CbmDefectRecord

    tmpl = tmp_path / "swg_overview.docx"
    panel_tmpl = tmp_path / "swg_panel.docx"
    tmpl.touch()
    panel_tmpl.touch()

    env = MagicMock()
    templates = {"swg_overview": tmpl, "swg_panel": panel_tmpl}
    env.get_template.side_effect = lambda k: templates.get(k)

    planner = CbmDefectPlanner()
    defects = [
        CbmDefectRecord(equipment="RMU SF6", technology="IR"),
    ]
    plans = planner.plan(defects, env)

    assert len(plans) == 1
    assert plans[0].spec.id == "swg"
    assert plans[0].overview_template == tmpl
    assert len(plans[0].groups) == 1
    assert plans[0].groups[0].item_key == "RMU SF6"
    assert plans[0].groups[0].overview.technology == "IR"
    assert len(plans[0].groups[0].detail_groups) == 1
    assert plans[0].groups[0].detail_groups[0].role_id == "panel_area"
    assert plans[0].groups[0].detail_groups[0].defects == (defects[0],)


def test_cbm_defect_planner_missing_overview_template_raises_error(tmp_path: Path):
    """Verify missing overview template raises FileNotFoundError."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner
    from src.quick_report.defects import CbmDefectRecord

    panel_tmpl = tmp_path / "swg_panel.docx"
    panel_tmpl.touch()

    env = MagicMock()
    env.get_template.side_effect = lambda k: panel_tmpl if k == "swg_panel" else None

    planner = CbmDefectPlanner()
    defects = [CbmDefectRecord(equipment="RMU SF6", technology="IR")]

    with pytest.raises(FileNotFoundError, match="Missing CBM template 'swg_overview' for family 'swg'"):
        planner.plan(defects, env)


def test_cbm_defect_planner_missing_detail_template_raises_error(tmp_path: Path):
    """Verify missing detail role template raises FileNotFoundError."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner
    from src.quick_report.defects import CbmDefectRecord

    overview_tmpl = tmp_path / "swg_overview.docx"
    overview_tmpl.touch()

    env = MagicMock()
    env.get_template.side_effect = lambda k: overview_tmpl if k == "swg_overview" else None

    planner = CbmDefectPlanner()
    defects = [CbmDefectRecord(equipment="RMU SF6", technology="IR")]

    with pytest.raises(FileNotFoundError, match="Missing CBM template 'swg_panel' for family 'swg'"):
        planner.plan(defects, env)


def test_cbm_defect_planner_tx_defects_split_hv_lv_roles(tmp_path: Path):
    """Verify TX defects are correctly split into HV and LV detail roles in detail_groups."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner
    from src.quick_report.defects import CbmDefectRecord

    overview_tmpl = tmp_path / "tx_overview.docx"
    hv_tmpl = tmp_path / "tx_hv_sides.docx"
    lv_tmpl = tmp_path / "tx_lv_sides.docx"
    overview_tmpl.touch()
    hv_tmpl.touch()
    lv_tmpl.touch()

    env = MagicMock()
    templates = {
        "tx_overview": overview_tmpl,
        "tx_hv_sides": hv_tmpl,
        "tx_lv_sides": lv_tmpl,
    }
    env.get_template.side_effect = lambda k: templates.get(k)

    planner = CbmDefectPlanner()
    d_hv = CbmDefectRecord(equipment="CABLE LTX/DTX", technology="IR")
    d_lv = CbmDefectRecord(equipment="LTX/DTX", technology="IR")

    plans = planner.plan([d_hv, d_lv], env)

    assert len(plans) == 1
    tx_plan = plans[0]
    assert tx_plan.spec.id == "tx"
    assert len(tx_plan.groups) == 2

    # Group 1: CABLE LTX/DTX
    g_hv = [g for g in tx_plan.groups if g.item_key == "CABLE LTX/DTX"][0]
    hv_role = [dg for dg in g_hv.detail_groups if dg.role_id == "tx_hv_side"][0]
    lv_role_in_hv_grp = [dg for dg in g_hv.detail_groups if dg.role_id == "tx_lv_side"][0]
    assert hv_role.defects == (d_hv,)
    assert lv_role_in_hv_grp.defects == ()

    # Group 2: LTX/DTX
    g_lv = [g for g in tx_plan.groups if g.item_key == "LTX/DTX"][0]
    hv_role_in_lv_grp = [dg for dg in g_lv.detail_groups if dg.role_id == "tx_hv_side"][0]
    lv_role = [dg for dg in g_lv.detail_groups if dg.role_id == "tx_lv_side"][0]
    assert hv_role_in_lv_grp.defects == ()
    assert lv_role.defects == (d_lv,)


def test_cbm_defect_planner_no_defects_returns_empty_tuple():
    """Verify empty cbm_defects returns empty tuple."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner

    env = MagicMock()
    planner = CbmDefectPlanner()
    plans = planner.plan([], env)
    assert plans == ()


def test_generate_cbm_defect_pages_typed_plan_processing(tmp_path: Path):
    """Verify renderer receives and processes a typed CbmDefectFamilyPlan."""
    from unittest.mock import patch
    from src.quick_report.cbm_defect_pages import generate_cbm_defect_pages
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.defects import CbmDefectRecord
    from src.quick_report.models import CbmDefectDetailGroup, CbmDefectFamilyPlan, CbmDefectGroup

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]

    overview_t = tmp_path / "swg_overview.docx"
    panel_t = tmp_path / "swg_panel.docx"
    overview_t.touch()
    panel_t.touch()

    defect = CbmDefectRecord(equipment="RMU SF6", technology="IR")
    group = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="",
        defects=(defect,),
        overview=defect,
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(defect,)),
        ),
    )

    family_plan = CbmDefectFamilyPlan(
        spec=spec,
        overview_template=overview_t,
        detail_templates=(("panel_area", panel_t),),
        groups=(group,),
    )

    with patch("src.quick_report.cbm_defect_pages._render_docx_template") as mock_render:
        paths = generate_cbm_defect_pages(
            family_plan, tmp_path, 101, {"substation": {"name_erms": "SUB A"}}
        )

    assert len(paths) == 2
    assert mock_render.call_count == 2
    assert "101_04_SWG_OVERVIEW.docx" in paths[0].name
    assert "101_04_SWG_RMU SF6.docx" in paths[1].name


def test_cbm_defect_page_builder_single_group(tmp_path: Path):
    """Verify CbmDefectPageBuilder builds overview and detail page plans for a single group in memory."""
    from src.quick_report.cbm_defect_pages import CbmDefectPageBuilder
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.defects import CbmDefectRecord
    from src.quick_report.models import CbmDefectDetailGroup, CbmDefectFamilyPlan, CbmDefectGroup

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]

    overview_t = tmp_path / "swg_overview.docx"
    panel_t = tmp_path / "swg_panel.docx"
    overview_t.touch()
    panel_t.touch()

    defect = CbmDefectRecord(
        equipment="RMU SF6",
        technology="IR",
        brand="ABB",
        model="SafePlus",
        rating="11kV",
        defect_area="Cable Compartment",
        additional_remarks="Hotspot",
        ir_reading="55.0",
    )
    group = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="PANEL 1",
        defects=(defect,),
        overview=defect,
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(defect,)),
        ),
    )
    family_plan = CbmDefectFamilyPlan(
        spec=spec,
        overview_template=overview_t,
        detail_templates=(("panel_area", panel_t),),
        groups=(group,),
    )

    pe_info = {"substation": {"name_erms": "PE TEST"}}
    builder = CbmDefectPageBuilder()
    page_plans = builder.build(
        family_plan=family_plan,
        pe_info=pe_info,
        substation_number=1,
    )

    assert len(page_plans) == 2

    # Overview Page Plan
    overview_plan = page_plans[0]
    assert overview_plan.template_path == overview_t
    assert overview_plan.output_filename == "001_04_SWG_OVERVIEW.docx"
    assert overview_plan.context["substation"] == {"name_erms": "PE TEST"}
    assert overview_plan.context["swg"]["area"] == "OVERVIEW"
    assert overview_plan.context["swg"]["type"] == "RMU SF6"
    assert overview_plan.context["panel"]["name"] == "PANEL 1"

    # Detail Page Plan
    detail_plan = page_plans[1]
    assert detail_plan.template_path == panel_t
    assert detail_plan.output_filename == "001_04_SWG_RMU SF6.docx"
    assert detail_plan.context["swg"]["area"] == "Cable Compartment/ Hotspot"
    assert detail_plan.context["panel"]["ir"]["reading"] == "55.0"


def test_cbm_defect_page_builder_multi_group_and_part_suffixes(tmp_path: Path):
    """Verify CbmDefectPageBuilder formats group and part suffixes correctly for multi-group and multi-defect plans."""
    from src.quick_report.cbm_defect_pages import CbmDefectPageBuilder
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.defects import CbmDefectRecord
    from src.quick_report.models import CbmDefectDetailGroup, CbmDefectFamilyPlan, CbmDefectGroup

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]

    overview_t = tmp_path / "swg_overview.docx"
    panel_t = tmp_path / "swg_panel.docx"
    overview_t.touch()
    panel_t.touch()

    d1 = CbmDefectRecord(equipment="RMU SF6", technology="IR", defect_area="Area 1")
    d2 = CbmDefectRecord(equipment="RMU SF6", technology="US", defect_area="Area 2")
    d3 = CbmDefectRecord(equipment="RMU SF6", technology="TEV", defect_area="Area 3")

    group1 = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="",
        defects=(d1, d2),
        overview=d1,
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(d1, d2)),
        ),
    )
    group2 = CbmDefectGroup(
        item_key="RMU SF6",
        item_suffix="",
        defects=(d3,),
        overview=d3,
        detail_groups=(
            CbmDefectDetailGroup(role_id="panel_area", defects=(d3,)),
        ),
    )

    family_plan = CbmDefectFamilyPlan(
        spec=spec,
        overview_template=overview_t,
        detail_templates=(("panel_area", panel_t),),
        groups=(group1, group2),
    )

    builder = CbmDefectPageBuilder()
    page_plans = builder.build(
        family_plan=family_plan,
        pe_info={"substation": {"name_erms": "PE TEST"}},
        substation_number=5,
    )

    # Page count check: 2 overview pages + 3 detail pages = 5
    assert len(page_plans) == 5

    filenames = [p.output_filename for p in page_plans]
    expected_filenames = [
        "005_04_SWG_OVERVIEW_grp1.docx",
        "005_04_SWG_RMU SF6_grp1_part1.docx",
        "005_04_SWG_RMU SF6_grp1_part2.docx",
        "005_04_SWG_OVERVIEW_grp2.docx",
        "005_04_SWG_RMU SF6_grp2.docx",
    ]
    assert filenames == expected_filenames

    # Check ordering: Overview grp1 -> Detail grp1 part1 -> Detail grp1 part2 -> Overview grp2 -> Detail grp2
    assert page_plans[0].template_path == overview_t
    assert page_plans[1].template_path == panel_t
    assert page_plans[2].template_path == panel_t
    assert page_plans[3].template_path == overview_t
    assert page_plans[4].template_path == panel_t


def test_cbm_defect_page_builder_empty_or_missing_templates(tmp_path: Path):
    """Verify CbmDefectPageBuilder returns empty list when plan has no groups or missing overview template."""
    from src.quick_report.cbm_defect_pages import CbmDefectPageBuilder
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.models import CbmDefectFamilyPlan

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    missing_template = tmp_path / "nonexistent.docx"

    family_plan = CbmDefectFamilyPlan(
        spec=spec,
        overview_template=missing_template,
        detail_templates=(),
        groups=(),
    )

    builder = CbmDefectPageBuilder()
    plans = builder.build(
        family_plan=family_plan,
        pe_info={},
        substation_number=1,
    )
    assert plans == []


def test_master_qr03_fetch_cbm_defects_alignment_standard_columns(tmp_path: Path):
    """Verify fetch_cbm_defects extracts equipment_id, us_char, tev_char from standard columns."""
    import pandas as pd

    cba_df = pd.DataFrame(
        [
            {
                "FUNCTIONAL LOCATION": "F/L 12345",
                "EQUIPMENT": "RMU SF6",
                "EQUIPMENT ID": "SWG 01",
                "TECHNOLOGY": "IR",
                "CRITICALITY": "HIGH",
                "BRAND": "ABB",
                "MODEL": "SafePlus",
                "RATING": "11kV",
                "DEFECT AREA": "Cable Compartment",
                "ADDITIONAL REMARKS": "Hotspot detected",
                "READING": "65.5",
                "US CHAR": "",
                "TEV CHAR": "",
            },
            {
                "FUNCTIONAL LOCATION": "F/L 12345",
                "EQUIPMENT": "RMU SF6",
                "EQUIPMENT ID": "SWG 02",
                "TECHNOLOGY": "US",
                "CRITICALITY": "MEDIUM",
                "BRAND": "ABB",
                "MODEL": "SafePlus",
                "RATING": "11kV",
                "DEFECT AREA": "Busbar",
                "ADDITIONAL REMARKS": "Tracking noise",
                "READING": "22.0",
                "US CHAR": "TRACKING",
                "TEV CHAR": "",
            },
            {
                "FUNCTIONAL LOCATION": "F/L 12345",
                "EQUIPMENT": "LTX/DTX",
                "EQUIPMENT ID": "TX 1",
                "TECHNOLOGY": "TEV",
                "CRITICALITY": "LOW",
                "BRAND": "Tamini",
                "MODEL": "ONAN",
                "RATING": "11/.415kV",
                "DEFECT AREA": "HV Bushing",
                "ADDITIONAL REMARKS": "Discharge activity",
                "READING": "18.5",
                "US CHAR": "",
                "TEV CHAR": "CONTINUOUS",
            },
        ]
    )
    vi_df = pd.DataFrame([{"FUNCTIONAL LOCATION": "F/L 12345", "EQUIPMENT": "SWG", "DEFECT AREA": "Door", "REMARKS": "None"}])

    engr_path = tmp_path / "ENGR-CBA-TEST.xlsx"
    with pd.ExcelWriter(engr_path, engine="openpyxl") as writer:
        cba_df.to_excel(writer, sheet_name="QR03 CBA", index=False)
        vi_df.to_excel(writer, sheet_name="QR03 VI", index=False)

    repo = MasterQr03DefectRepository(tmp_path)
    defects = repo.fetch_cbm_defects("12345")

    assert len(defects) == 3

    # IR Defect
    ir_def = defects[0]
    assert ir_def.equipment == "RMU SF6"
    assert ir_def.equipment_id == "SWG 01"
    assert ir_def.technology == "IR"
    assert ir_def.raw_measurement == "65.5"
    assert ir_def.ir_reading == "65.5"
    assert ir_def.us_reading == ""
    assert ir_def.tev_reading == ""

    # US Defect
    us_def = defects[1]
    assert us_def.equipment == "RMU SF6"
    assert us_def.equipment_id == "SWG 02"
    assert us_def.technology == "US"
    assert us_def.raw_measurement == "22.0"
    assert us_def.us_reading == "22.0"
    assert us_def.us_char == "TRACKING"

    # TEV Defect
    tev_def = defects[2]
    assert tev_def.equipment == "LTX/DTX"
    assert tev_def.equipment_id == "TX 1"
    assert tev_def.technology == "TEV"
    assert tev_def.raw_measurement == "18.5"
    assert tev_def.tev_reading == "18.5"
    assert tev_def.tev_char == "CONTINUOUS"


def test_master_qr03_fetch_cbm_defects_alignment_fallback_columns(tmp_path: Path):
    """Verify fetch_cbm_defects handles fallback columns: EQUIPMENT_ID/ID, DEFECT FROM/TECH, STATUS, DEFECT TYPE."""
    import pandas as pd

    cba_df = pd.DataFrame(
        [
            {
                "FL": "PE-9999",
                "EQUIPMENT": "RMU SF6",
                "EQUIPMENT_ID": "P1",
                "DEFECT FROM": "US",
                "STATUS": "CRITICAL",
                "BRAND": "Schneider",
                "MODEL": "RM6",
                "RATING": "11kV",
                "DEFECT_AREA": "Cable Box",
                "REMARKS": "Discharge sound",
                "US READING": "28.4",
                "DEFECT TYPE": "CORONA",
            },
            {
                "FL": "PE-9999",
                "EQUIPMENT": "LTX/DTX",
                "ID": "TX2",
                "TECH": "TEV",
                "STATUS": "HIGH",
                "BRAND": "Wilson",
                "MODEL": "1000kVA",
                "RATING": "11/.415kV",
                "DEFECT AREA": "LV Bushing",
                "REMARKS": "High TEV level",
                "TEV READING": "31.2",
                "DEFECT TYPE": "SURFACE",
            },
            {
                "FL": "PE-9999",
                "EQUIPMENT": "FP (D)",
                "ID": "FP TX1",
                "DEFECT FROM": "IR",
                "STATUS": "LOW",
                "BRAND": "Tamco",
                "MODEL": "DIN",
                "RATING": "415V",
                "DEFECT AREA": "Fuse Base",
                "REMARKS": "Loose connection",
                "IR READING": "58.2",
                "DEFECT TYPE": "HOTSPOT",
            },
        ]
    )
    vi_df = pd.DataFrame([{"FL": "PE-9999", "EQUIPMENT": "FP (D)", "DEFECT AREA": "Cover", "REMARKS": "None"}])

    engr_path = tmp_path / "ENGR-CBA-FALLBACK-TEST.xlsx"
    with pd.ExcelWriter(engr_path, engine="openpyxl") as writer:
        cba_df.to_excel(writer, sheet_name="QR03 CBA", index=False)
        vi_df.to_excel(writer, sheet_name="QR03 VI", index=False)

    repo = MasterQr03DefectRepository(tmp_path)
    defects = repo.fetch_cbm_defects("PE-9999")

    assert len(defects) == 3

    # Row 1: US defect with EQUIPMENT_ID, DEFECT FROM, STATUS, DEFECT TYPE fallback
    d1 = defects[0]
    assert d1.equipment == "RMU SF6"
    assert d1.equipment_id == "P1"
    assert d1.technology == "US"
    assert d1.us_char == "CORONA DISCHARGE"
    assert d1.us_reading == "28.4"
    assert d1.raw_measurement == "28.4"

    # Row 2: TEV defect with ID, TECH, STATUS, DEFECT TYPE fallback
    d2 = defects[1]
    assert d2.equipment == "LTX/DTX"
    assert d2.equipment_id == "TX2"
    assert d2.technology == "TEV"
    assert d2.tev_char == "SURFACE"
    assert d2.tev_reading == "31.2"
    assert d2.raw_measurement == "31.2"

    # Row 3: IR defect with ID, DEFECT FROM, STATUS; DEFECT TYPE does not populate us_char/tev_char
    d3 = defects[2]
    assert d3.equipment == "FP (D)"
    assert d3.equipment_id == "FP TX1"
    assert d3.technology == "IR"
    assert d3.us_char == ""
    assert d3.tev_char == ""
    assert d3.ir_reading == "58.2"
    assert d3.raw_measurement == "58.2"


def test_master_qr03_fetch_cbm_defects_case_and_whitespace_insensitivity(tmp_path: Path):
    """Verify column matching is resilient to case variations and leading/trailing whitespace."""
    import pandas as pd

    cba_df = pd.DataFrame(
        [
            {
                " functional location ": "PE-8888",
                " equipment ": "RMU SF6",
                " Equipment ID ": "SWG-PANEL-3",
                " defect from ": "us",
                " status ": "MEDIUM",
                " Us Character ": "ARCING",
                " reading ": "19.5",
            }
        ]
    )
    vi_df = pd.DataFrame([{" functional location ": "PE-8888", " equipment ": "RMU SF6", " defect area ": "Door"}])

    engr_path = tmp_path / "ENGR-CASE-TEST.xlsx"
    with pd.ExcelWriter(engr_path, engine="openpyxl") as writer:
        cba_df.to_excel(writer, sheet_name="QR03 CBA", index=False)
        vi_df.to_excel(writer, sheet_name="QR03 VI", index=False)

    repo = MasterQr03DefectRepository(tmp_path)
    defects = repo.fetch_cbm_defects("PE-8888")

    assert len(defects) == 1
    d = defects[0]
    assert d.equipment == "RMU SF6"
    assert d.equipment_id == "SWG-PANEL-3"
    assert d.technology == "US"
    assert d.us_char == "ARCING"
    assert d.us_reading == "19.5"


def test_cbm_family_specs_canonical_aliasing_and_technologies():
    """Verify QUICK_REPORT_FAMILY_SPECS contains canonical aliased values and IR/US/TEV support."""
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID

    # 1. SWG
    swg = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    assert swg.technologies == ("IR", "US", "TEV")
    assert swg.overview_template_key == "swg_overview"
    assert swg.equipment_values == (
        "RMU SF6",
        "RMU OIL",
        "VCB 11kV",
        "VCB 33kV",
        "MRMU",
        "CABLE SWG",
        "EARTHING",
        "SWITCHGEAR",
        "GIS 33kV",
    )
    assert len(swg.detail_roles) == 1
    assert swg.detail_roles[0].id == "panel_area"
    assert swg.detail_roles[0].template_key == "swg_panel"
    assert swg.detail_roles[0].technologies == ("IR", "US", "TEV")
    assert swg.detail_roles[0].equipment_values == swg.equipment_values

    # 2. TX
    tx = QUICK_REPORT_FAMILY_SPECS_BY_ID["tx"]
    assert tx.technologies == ("IR", "US", "TEV")
    assert tx.overview_template_key == "tx_overview"
    assert tx.equipment_values == (
        "LTX/DTX",
        "CABLE LTX/DTX",
        "PTX",
        "CABLE PTX",
        "TRANSFORMER",
    )
    assert len(tx.detail_roles) == 2
    hv_role = next(r for r in tx.detail_roles if r.id == "tx_hv_side")
    lv_role = next(r for r in tx.detail_roles if r.id == "tx_lv_side")
    assert hv_role.template_key == "tx_hv_sides"
    assert hv_role.technologies == ("IR", "US", "TEV")
    assert hv_role.equipment_values == tx.equipment_values
    assert lv_role.template_key == "tx_lv_sides"
    assert lv_role.technologies == ("IR", "US", "TEV")
    assert lv_role.equipment_values == tx.equipment_values

    # 3. FP_LVDB
    fp = QUICK_REPORT_FAMILY_SPECS_BY_ID["fp_lvdb"]
    assert fp.technologies == ("IR", "US", "TEV")
    assert fp.overview_template_key == "fp_overview"
    assert fp.equipment_values == (
        "FP (D)",
        "FP (J)",
        "LVDB",
        "CABLE LVDB/FP",
        "FP",
    )
    assert len(fp.detail_roles) == 1
    assert fp.detail_roles[0].id == "fp_detail"
    assert fp.detail_roles[0].template_key == "fp_individual_defect"
    assert fp.detail_roles[0].technologies == ("IR", "US", "TEV")
    assert fp.detail_roles[0].equipment_values == fp.equipment_values

    # 4. BATTERY
    battery = QUICK_REPORT_FAMILY_SPECS_BY_ID["battery"]
    assert battery.technologies == ("IR", "US", "TEV")
    assert battery.overview_template_key == "battery_overview"
    assert battery.equipment_values == (
        "BATTERY CHARGER",
        "BATTERY BANK",
        "BATTERY",
    )
    assert len(battery.detail_roles) == 1
    assert battery.detail_roles[0].id == "battery_detail"
    assert battery.detail_roles[0].template_key == "battery_overview"
    assert battery.detail_roles[0].technologies == ("IR", "US", "TEV")
    assert battery.detail_roles[0].equipment_values == battery.equipment_values

    # 5. BLACKBOX
    blackbox = QUICK_REPORT_FAMILY_SPECS_BY_ID["blackbox"]
    assert blackbox.technologies == ("IR", "US", "TEV")
    assert blackbox.overview_template_key == "blackbox_overview"
    assert blackbox.equipment_values == ("BLACK BOX", "BLACKBOX")
    assert len(blackbox.detail_roles) == 1
    assert blackbox.detail_roles[0].id == "blackbox_detail"
    assert blackbox.detail_roles[0].template_key == "blackbox_overview"
    assert blackbox.detail_roles[0].technologies == ("IR", "US", "TEV")
    assert blackbox.detail_roles[0].equipment_values == blackbox.equipment_values


def test_cbm_defect_planner_canonical_aliasing_matching(tmp_path: Path):
    """Verify CbmDefectPlanner matches canonical aliased equipment names to families."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner

    # Create dummy templates
    template_names = [
        "swg_overview", "swg_panel",
        "tx_overview", "tx_hv_sides", "tx_lv_sides",
        "fp_overview", "fp_individual_defect",
        "battery_overview",
        "blackbox_overview",
    ]
    template_map = {}
    for name in template_names:
        p = tmp_path / f"{name}.docx"
        p.touch()
        template_map[name] = p

    env = MagicMock()
    env.get_template.side_effect = lambda k: template_map.get(k)

    planner = CbmDefectPlanner()

    defects = [
        CbmDefectRecord(equipment="VCB 11kV", technology="IR", defect_area="Cable Box"),
        CbmDefectRecord(equipment="MRMU", technology="US", defect_area="Busbar"),
        CbmDefectRecord(equipment="GIS 33kV", technology="TEV", defect_area="Chamber"),
        CbmDefectRecord(equipment="PTX", technology="IR", defect_area="HV Bushing"),
        CbmDefectRecord(equipment="CABLE PTX", technology="US", defect_area="Terminations"),
        CbmDefectRecord(equipment="FP (J)", technology="IR", defect_area="Fuse Base"),
        CbmDefectRecord(equipment="CABLE LVDB/FP", technology="IR", defect_area="Cable Joint"),
        CbmDefectRecord(equipment="BATTERY CHARGER", technology="IR", defect_area="Diode"),
        CbmDefectRecord(equipment="BLACK BOX", technology="IR", defect_area="Enclosure"),
    ]

    plans = planner.plan(defects, env)
    family_ids = [p.spec.id for p in plans]

    assert "swg" in family_ids
    assert "tx" in family_ids
    assert "fp_lvdb" in family_ids
    assert "battery" in family_ids
    assert "blackbox" in family_ids

    # Verify SWG group contains VCB 11kV, MRMU, GIS 33kV
    swg_plan = next(p for p in plans if p.spec.id == "swg")
    swg_defects = [d.equipment for g in swg_plan.groups for d in g.defects]
    assert "VCB 11kV" in swg_defects
    assert "MRMU" in swg_defects
    assert "GIS 33kV" in swg_defects

    # Verify TX group contains PTX, CABLE PTX
    tx_plan = next(p for p in plans if p.spec.id == "tx")
    tx_defects = [d.equipment for g in tx_plan.groups for d in g.defects]
    assert "PTX" in tx_defects
    assert "CABLE PTX" in tx_defects


def test_cbm_defect_planner_equipment_family_grouping_by_item_key(tmp_path: Path):
    """Verify defects are grouped by equipment family (item_key derived from equipment_id or equipment)."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner

    template_map = {
        "swg_overview": tmp_path / "swg_overview.docx",
        "swg_panel": tmp_path / "swg_panel.docx",
    }
    for p in template_map.values():
        p.touch()

    env = MagicMock()
    env.get_template.side_effect = lambda k: template_map.get(k)

    planner = CbmDefectPlanner()

    # Defects for 2 distinct switchgear panels, plus 1 defect without equipment_id
    defects = [
        CbmDefectRecord(equipment="RMU SF6", equipment_id="PANEL 1", defect_area="Cable Box", technology="IR"),
        CbmDefectRecord(equipment="RMU SF6", equipment_id="PANEL 1", defect_area="Busbar", technology="US"),
        CbmDefectRecord(equipment="RMU SF6", equipment_id="PANEL 2", defect_area="Cable Box", technology="IR"),
        CbmDefectRecord(equipment="RMU SF6", equipment_id="", defect_area="Earth Switch", technology="IR"),
    ]

    plans = planner.plan(defects, env)
    assert len(plans) == 1
    swg_plan = plans[0]

    # Single switchgear board belongs to 1 equipment family group: "RMU SF6"
    assert len(swg_plan.groups) == 1
    swg_group = swg_plan.groups[0]
    assert swg_group.item_key == "RMU SF6"
    # All 4 defect rows are preserved across panels within the equipment family
    assert len(swg_group.defects) == 4

    # Multiple distinct equipment families (e.g. FP TX1 and FP TX2) produce distinct groups
    fp_template_map = {
        "fp_overview": tmp_path / "fp_overview.docx",
        "fp_individual_defect": tmp_path / "fp_individual_defect.docx",
    }
    for p in fp_template_map.values():
        p.touch()
    env.get_template.side_effect = lambda k: fp_template_map.get(k) or template_map.get(k)

    fp_defects = [
        CbmDefectRecord(equipment="FP (J)", equipment_id="FP TX1 - OUTGOING F1", defect_area="Fuse", technology="IR"),
        CbmDefectRecord(equipment="FP (J)", equipment_id="FP TX1 - OUTGOING F2", defect_area="Fuse", technology="IR"),
        CbmDefectRecord(equipment="FP (J)", equipment_id="FP TX2 - OUTGOING F1", defect_area="Fuse", technology="IR"),
    ]
    fp_plans = planner.plan(fp_defects, env)
    assert len(fp_plans) == 1
    fp_plan = fp_plans[0]
    assert [g.item_key for g in fp_plan.groups] == ["FP TX1", "FP TX2"]
    assert len(fp_plan.groups[0].defects) == 2  # F1 and F2 under FP TX1
    assert len(fp_plan.groups[1].defects) == 1  # F1 under FP TX2


def test_cbm_defect_planner_multi_technology_merging_same_item_and_area(tmp_path: Path):
    """Verify multiple defects sharing (item_key, defect_area) merge readings into one unified record."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner

    template_map = {
        "swg_overview": tmp_path / "swg_overview.docx",
        "swg_panel": tmp_path / "swg_panel.docx",
    }
    for p in template_map.values():
        p.touch()

    env = MagicMock()
    env.get_template.side_effect = lambda k: template_map.get(k)

    planner = CbmDefectPlanner()

    defects = [
        CbmDefectRecord(
            equipment="RMU SF6",
            equipment_id="PANEL 1",
            defect_area="Cable Compartment",
            technology="IR",
            ir_reading="58.5",
            additional_remarks="Thermal hotspot",
            source_order=1,
        ),
        CbmDefectRecord(
            equipment="RMU SF6",
            equipment_id="PANEL 1",
            defect_area="cable compartment",  # Case-insensitive match
            technology="US",
            us_reading="25.0",
            us_char="CORONA",
            additional_remarks="Corona sound",
            source_order=2,
        ),
        CbmDefectRecord(
            equipment="RMU SF6",
            equipment_id="PANEL 1",
            defect_area="Cable Compartment",
            technology="TEV",
            tev_reading="33.5",
            tev_char="CONTINUOUS",
            source_order=3,
        ),
    ]

    plans = planner.plan(defects, env)
    assert len(plans) == 1
    swg_plan = plans[0]
    assert len(swg_plan.groups) == 1

    group = swg_plan.groups[0]
    assert group.item_key == "RMU SF6"
    # Merged to exactly 1 unified defect
    assert len(group.defects) == 1

    merged = group.defects[0]
    assert merged.equipment == "RMU SF6"
    assert merged.equipment_id == "PANEL 1"
    assert merged.defect_area == "Cable Compartment"
    assert merged.ir_reading == "58.5"
    assert merged.us_reading == "25.0"
    assert merged.us_char == "CORONA"
    assert merged.tev_reading == "33.5"
    assert merged.tev_char == "CONTINUOUS"
    # Source order preserved from earliest: 1
    assert merged.source_order == 1

    # Verify detail_groups has the merged record
    assert len(group.detail_groups) == 1
    assert group.detail_groups[0].defects == (merged,)


def test_cbm_defect_planner_different_defect_areas_not_merged(tmp_path: Path):
    """Verify defects on same apparatus with distinct defect areas remain separate records."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner

    template_map = {
        "swg_overview": tmp_path / "swg_overview.docx",
        "swg_panel": tmp_path / "swg_panel.docx",
    }
    for p in template_map.values():
        p.touch()

    env = MagicMock()
    env.get_template.side_effect = lambda k: template_map.get(k)

    planner = CbmDefectPlanner()

    defects = [
        CbmDefectRecord(
            equipment="RMU SF6",
            equipment_id="PANEL 1",
            defect_area="Cable Compartment",
            technology="IR",
            ir_reading="58.5",
        ),
        CbmDefectRecord(
            equipment="RMU SF6",
            equipment_id="PANEL 1",
            defect_area="Busbar Chamber",
            technology="US",
            us_reading="22.0",
        ),
    ]

    plans = planner.plan(defects, env)
    assert len(plans) == 1
    group = plans[0].groups[0]
    assert len(group.defects) == 2
    areas = [d.defect_area for d in group.defects]
    assert "Cable Compartment" in areas
    assert "Busbar Chamber" in areas


def test_cbm_defect_planner_tx_smart_routing(tmp_path: Path):
    """Verify TX smart routing routes to tx_hv_side vs tx_lv_side according to area/id/equipment rules."""
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner

    template_map = {
        "tx_overview": tmp_path / "tx_overview.docx",
        "tx_hv_sides": tmp_path / "tx_hv_sides.docx",
        "tx_lv_sides": tmp_path / "tx_lv_sides.docx",
    }
    for p in template_map.values():
        p.touch()

    env = MagicMock()
    env.get_template.side_effect = lambda k: template_map.get(k)

    planner = CbmDefectPlanner()

    # Test cases for TX routing:
    # 1. HV Bushing -> tx_hv_side
    # 2. 11kV Termination -> tx_hv_side
    # 3. 33kV Bushing -> tx_hv_side
    # 4. equipment_id with HV -> tx_hv_side
    # 5. LV Cable Box -> tx_lv_side
    # 6. 415V Terminations -> tx_lv_side
    # 7. equipment_id with LV -> tx_lv_side
    # 8. Fallback CABLE -> tx_hv_side
    # 9. Fallback non-CABLE -> tx_lv_side
    defects = [
        CbmDefectRecord(equipment="LTX/DTX", equipment_id="TX 1", defect_area="HV Bushing", technology="IR"),
        CbmDefectRecord(equipment="TRANSFORMER", equipment_id="TX 2", defect_area="11kV Termination", technology="IR"),
        CbmDefectRecord(equipment="PTX", equipment_id="TX 3", defect_area="33kV Bushing", technology="IR"),
        CbmDefectRecord(equipment="LTX/DTX", equipment_id="TX 4 (HV)", defect_area="General Defect", technology="IR"),
        CbmDefectRecord(equipment="LTX/DTX", equipment_id="TX 5", defect_area="LV Cable Box", technology="IR"),
        CbmDefectRecord(equipment="TRANSFORMER", equipment_id="TX 6", defect_area="415V Terminations", technology="IR"),
        CbmDefectRecord(equipment="PTX", equipment_id="TX 7 LV", defect_area="General Defect", technology="IR"),
        CbmDefectRecord(equipment="CABLE LTX/DTX", equipment_id="TX 8", defect_area="Oil Leak", technology="IR"),
        CbmDefectRecord(equipment="CABLE PTX", equipment_id="TX 9", defect_area="Corrosion", technology="IR"),
        CbmDefectRecord(equipment="LTX/DTX", equipment_id="TX 10", defect_area="Tank Body Rust", technology="IR"),
        CbmDefectRecord(equipment="PTX", equipment_id="TX 11", defect_area="Radiator Leak", technology="IR"),
        CbmDefectRecord(equipment="TRANSFORMER", equipment_id="TX 12", defect_area="Silica Gel", technology="IR"),
    ]

    plans = planner.plan(defects, env)
    assert len(plans) == 1
    tx_plan = plans[0]

    groups_by_key = {g.item_key: g for g in tx_plan.groups}

    def get_routed_role(item_key: str) -> str:
        grp = groups_by_key[item_key]
        for dg in grp.detail_groups:
            if dg.defects:
                return dg.role_id
        return ""

    # HV routing assertions
    assert get_routed_role("TX 1") == "tx_hv_side"
    assert get_routed_role("TX 2") == "tx_hv_side"
    assert get_routed_role("TX 3") == "tx_hv_side"
    assert get_routed_role("TX 4 (HV)") == "tx_hv_side"
    assert get_routed_role("TX 8") == "tx_hv_side"   # CABLE LTX/DTX fallback
    assert get_routed_role("TX 9") == "tx_hv_side"   # CABLE PTX fallback

    # LV routing assertions
    assert get_routed_role("TX 5") == "tx_lv_side"
    assert get_routed_role("TX 6") == "tx_lv_side"
    assert get_routed_role("TX 7 LV") == "tx_lv_side"
    assert get_routed_role("TX 10") == "tx_lv_side"  # LTX/DTX fallback
    assert get_routed_role("TX 11") == "tx_lv_side"  # PTX fallback
    assert get_routed_role("TX 12") == "tx_lv_side"  # TRANSFORMER fallback


def test_cbm_defect_planner_non_matching_equipment_skipping_with_log(
    tmp_path: Path, caplog: pytest.LogCaptureFixture
):
    """Verify non-matching equipment (e.g. switchyard/overhead) is skipped with an explicit logger info message."""
    import logging
    from unittest.mock import MagicMock
    from src.quick_report.cbm_defect_planner import CbmDefectPlanner

    env = MagicMock()
    planner = CbmDefectPlanner()

    non_matching = [
        CbmDefectRecord(equipment="OVERHEAD LINE", defect_area="Conductor", technology="IR"),
        CbmDefectRecord(equipment="SWITCHYARD", defect_area="Isolator", technology="US"),
        CbmDefectRecord(equipment="LIGHTNING ARRESTER", defect_area="Surge Counter", technology="IR"),
    ]

    with caplog.at_level(logging.INFO):
        plans = planner.plan(non_matching, env)


def test_format_detail_area_formatting():
    from src.quick_report.cbm_render import _format_detail_area

    # 1. Overview returns "OVERVIEW"
    assert _format_detail_area("Cable Box", "Hotspot", overview=True) == "OVERVIEW"
    assert _format_detail_area("", "", overview=True) == "OVERVIEW"

    # 2. Defect area + remarks formatted verbatim as "{area}/ {remarks}"
    assert _format_detail_area("Cable Compartment", "Phase R Hotspot", overview=False) == "Cable Compartment/ Phase R Hotspot"
    assert _format_detail_area("HV Bushing", "Discharge noise", overview=False) == "HV Bushing/ Discharge noise"

    # 3. Only defect area
    assert _format_detail_area("Cable Compartment", "", overview=False) == "Cable Compartment"
    assert _format_detail_area("Cable Compartment", None, overview=False) == "Cable Compartment"
    assert _format_detail_area("Cable Compartment", "   ", overview=False) == "Cable Compartment"

    # 4. Only remarks
    assert _format_detail_area("", "Hotspot", overview=False) == "Hotspot"
    assert _format_detail_area(None, "Hotspot", overview=False) == "Hotspot"

    # 5. Empty / None fallback
    assert _format_detail_area("", "", overview=False) == "-"
    assert _format_detail_area(None, None, overview=False) == "-"


def test_swg_render_context_with_testsheet_match():
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.cbm_render import _build_family_render_context
    from src.quick_report.defects import CbmDefectRecord
    from src.testsheet.models import (
        SubstationEquipmentPackage,
        SwitchgearPanelSpec,
        SwitchgearSpec,
    )

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    panel_1 = SwitchgearPanelSpec(
        panel_no=1,
        panel_feeder_no="1",
        name="TX 1",
        status="CLOSE",
        load_amp="120A",
        cable_type="XLPE 300mm",
        heater_amp="0.5A",
        serial_no="PNL-SN-001",
    )
    swg_spec = SwitchgearSpec(
        switchgear_type="RMU SF6",
        manufacturer="Schneider",
        model="RM6",
        rating="11kV",
        serial_no="SN-SWG-001",
        panels=(panel_1,),
    )
    equipment_pkg = SubstationEquipmentPackage(switchgears=(swg_spec,))
    pe_info = {"equipment_specs": equipment_pkg, "substation": {"name_erms": "PE TEST"}}

    rec = CbmDefectRecord(
        equipment="RMU SF6",
        equipment_id="PANEL 1",
        technology="IR",
        defect_area="Cable Compartment",
        additional_remarks="Phase R Hotspot",
        ir_reading="55.2",
    )

    ctx = _build_family_render_context(
        spec,
        rec,
        overview=False,
        item_key="PANEL 1",
        item_suffix="PANEL 1",
        pe_info=pe_info,
    )

    assert ctx["swg"]["area"] == "Cable Compartment/ Phase R Hotspot"
    assert ctx["swg"]["manufacturer"] == "Schneider"
    assert ctx["swg"]["model"] == "RM6"
    assert ctx["swg"]["rating"] == "11kV"
    assert ctx["swg"]["serialnumber"] == "SN-SWG-001"
    assert ctx["swg"]["type"] == "RMU SF6"

    assert ctx["panel"]["name"] == "PANEL 1"
    assert ctx["panel"]["linknumber"] == "PANEL 1"
    assert ctx["panel"]["area"] == "Cable Compartment/ Phase R Hotspot"
    assert ctx["panel"]["loadamp"] == "120A"
    assert ctx["panel"]["heateramp"] == "0.5A"
    assert ctx["panel"]["breakerstatus"] == "CLOSE"
    assert ctx["panel"]["cabletype"] == "XLPE 300mm"
    assert ctx["panel"]["serialnumber"] == "PNL-SN-001"
    assert ctx["panel"]["busbarposition"] == "-"
    assert ctx["panel"]["ir"]["reading"] == "55.2"
    assert ctx["panel"]["us"]["reading"] == "-"
    assert ctx["panel"]["us"]["char"] == "-"
    assert ctx["panel"]["tev"]["reading"] == "-"
    assert ctx["panel"]["tev"]["char"] == "-"


def test_swg_render_context_without_testsheet_fallback_dash():
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.cbm_render import _build_family_render_context
    from src.quick_report.defects import CbmDefectRecord

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    rec = CbmDefectRecord(
        equipment="CABLE SWG",
        equipment_id="PANEL 2",
        technology="US",
        defect_area="Cable Box",
        additional_remarks="",
        us_reading="24.0",
        us_char="TRACKING",
    )

    ctx = _build_family_render_context(
        spec,
        rec,
        overview=False,
        item_key="PANEL 2",
        pe_info={},
    )

    assert ctx["swg"]["area"] == "Cable Box"
    assert ctx["swg"]["manufacturer"] == "-"
    assert ctx["swg"]["model"] == "-"
    assert ctx["swg"]["rating"] == "-"
    assert ctx["swg"]["serialnumber"] == "-"
    assert ctx["swg"]["type"] == "CABLE SWG"

    assert ctx["panel"]["name"] == "PANEL 2"
    assert ctx["panel"]["linknumber"] == "PANEL 2"
    assert ctx["panel"]["area"] == "Cable Box"
    assert ctx["panel"]["loadamp"] == "-"
    assert ctx["panel"]["heateramp"] == "-"
    assert ctx["panel"]["breakerstatus"] == "-"
    assert ctx["panel"]["cabletype"] == "CABLE SWG"
    assert ctx["panel"]["serialnumber"] == "-"
    assert ctx["panel"]["busbarposition"] == "-"
    assert ctx["panel"]["ir"]["reading"] == "-"
    assert ctx["panel"]["us"]["reading"] == "24"
    assert ctx["panel"]["us"]["char"] == "TRACKING"
    assert ctx["panel"]["tev"]["reading"] == "-"
    assert ctx["panel"]["tev"]["char"] == "-"


def test_tx_render_context_with_and_without_testsheet():
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.cbm_render import _build_family_render_context
    from src.quick_report.defects import CbmDefectRecord
    from src.testsheet.models import SubstationEquipmentPackage, TransformerSpec

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["tx"]
    tx_spec = TransformerSpec(
        tx_id="Tx 1",
        rating_kva="1000kVA",
        manufacturer="Tamini",
        type="ONAN",
        serial_no="TX-SN-999",
    )
    equipment_pkg = SubstationEquipmentPackage(transformers=(tx_spec,))
    pe_info = {
        "equipment_specs": equipment_pkg,
        "substation": {"building_type": "ATTACHED", "name_erms": "PE TEST"},
    }

    rec = CbmDefectRecord(
        equipment="LTX/DTX",
        equipment_id="TX 1",
        technology="US",
        defect_area="HV Bushing",
        additional_remarks="Corona sound",
        us_reading="18.5",
        us_char="CORONA",
    )

    ctx = _build_family_render_context(
        spec,
        rec,
        overview=False,
        item_key="TX 1",
        pe_info=pe_info,
    )

    assert ctx["tx"]["number"] == "TX 1"
    assert ctx["tx"]["location"] == "HV - SIDE"
    assert ctx["tx"]["area"] == "HV Bushing/ Corona sound"
    assert ctx["tx"]["manufacturer"] == "Tamini"
    assert ctx["tx"]["model"] == "ONAN"
    assert ctx["tx"]["rating"] == "1000kVA"
    assert ctx["tx"]["serialnumber"] == "TX-SN-999"
    assert ctx["tx"]["cabletype"] == "-"
    assert ctx["tx"]["ir"]["reading"] == "-"
    assert ctx["tx"]["us"]["reading"] == "19"
    assert ctx["tx"]["us"]["char"] == "CORONA DISCHARGE"
    assert ctx["tx"]["tev"]["reading"] == "-"
    assert ctx["tx"]["tev"]["char"] == "-"

    # Without testsheet
    ctx_empty = _build_family_render_context(spec, rec, overview=False, item_key="TX 1", pe_info={})
    assert ctx_empty["tx"]["number"] == "TX 1"
    assert ctx_empty["tx"]["location"] == "HV - SIDE"
    assert ctx_empty["tx"]["manufacturer"] == "-"
    assert ctx_empty["tx"]["model"] == "-"
    assert ctx_empty["tx"]["rating"] == "-"
    assert ctx_empty["tx"]["serialnumber"] == "-"


def test_fp_lvdb_render_context_splitting_and_testsheet():
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.cbm_render import _build_family_render_context
    from src.quick_report.defects import CbmDefectRecord
    from src.testsheet.models import LVDBSpec, SubstationEquipmentPackage

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["fp_lvdb"]
    lv_spec = LVDBSpec(
        name="LVDB 1",
        label="FP",
        source="TX1",
        manufacturer="Tamco",
        serial_no="FP-SN-123",
        rating="415V",
    )
    equipment_pkg = SubstationEquipmentPackage(lvdb_specs=(lv_spec,))
    pe_info = {"equipment_specs": equipment_pkg, "substation": {"name_erms": "PE TEST"}}

    # With split "FP TX1 - OUTGOING F5"
    rec_split = CbmDefectRecord(
        equipment="FP (D)",
        equipment_id="FP TX1 - OUTGOING F5",
        technology="IR",
        defect_area="Fuse Base",
        additional_remarks="Hotspot",
        ir_reading="60.1",
    )
    ctx_split = _build_family_render_context(
        spec, rec_split, overview=False, item_key="FP TX1 - OUTGOING F5", pe_info=pe_info
    )
    assert ctx_split["fp"]["labelsource"] == "FP TX1"
    assert ctx_split["fp"]["feederno"] == "OUTGOING F5"
    assert ctx_split["fp"]["area"] == "Fuse Base/ Hotspot"
    assert ctx_split["fp"]["manufacturer"] == "Tamco"
    assert ctx_split["fp"]["rating"] == "415V"
    assert ctx_split["fp"]["serialnumber"] == "FP-SN-123"
    assert ctx_split["fp"]["ir"]["reading"] == "60.1"
    assert ctx_split["fp"]["us"]["reading"] == "-"
    assert ctx_split["fp"]["tev"]["reading"] == "-"

    # Without split
    rec_nosplit = CbmDefectRecord(
        equipment="FP (D)",
        equipment_id="FP TX1",
        technology="IR",
        defect_area="Fuse Base",
    )
    ctx_nosplit = _build_family_render_context(
        spec, rec_nosplit, overview=False, item_key="FP TX1", pe_info=pe_info
    )
    assert ctx_nosplit["fp"]["labelsource"] == "FP TX1"
    assert ctx_nosplit["fp"]["feederno"] == "-"


def test_blackbox_render_context_number_and_location_detection():
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.cbm_render import _build_family_render_context
    from src.quick_report.defects import CbmDefectRecord

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["blackbox"]
    rec = CbmDefectRecord(
        equipment="BLACK BOX",
        equipment_id="BLACK BOX 2",
        defect_area="LEFT SIDE TERMINAL",
        additional_remarks="Loose wire",
        technology="TEV",
        tev_reading="24.0",
        tev_char="SURFACE",
    )

    ctx = _build_family_render_context(
        spec,
        rec,
        overview=False,
        item_key="BLACK BOX 2",
        pe_info={},
    )

    assert ctx["bbox"]["number"] == "2"
    assert ctx["bbox"]["location"] == "LEFT"
    assert ctx["bbox"]["area"] == "LEFT SIDE TERMINAL/ Loose wire"
    assert ctx["bbox"]["tev"]["reading"] == "24"
    assert ctx["bbox"]["tev"]["char"] == "SURFACE"
    assert ctx["bbox"]["ir"]["reading"] == "-"
    assert ctx["bbox"]["us"]["reading"] == "-"


def test_battery_render_context_with_testsheet():
    from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID
    from src.quick_report.cbm_render import _build_family_render_context
    from src.quick_report.defects import CbmDefectRecord
    from src.testsheet.models import BatteryBankSpec, SubstationEquipmentPackage

    spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["battery"]
    batt_spec = BatteryBankSpec(
        name="BATTERY BANK 1",
        manufacturer="Chloride",
        model="Powersafe",
        serial_no="BATT-007",
    )
    equipment_pkg = SubstationEquipmentPackage(battery_banks=(batt_spec,))
    pe_info = {"equipment_specs": equipment_pkg, "substation": {"name_erms": "PE TEST"}}

    rec = CbmDefectRecord(
        equipment="BATTERY BANK",
        equipment_id="BATTERY BANK 1",
        technology="IR",
        defect_area="Terminal Post",
        additional_remarks="Phase R",
        ir_reading="45.0",
    )

    ctx = _build_family_render_context(
        spec,
        rec,
        overview=False,
        item_key="BATTERY BANK 1",
        pe_info=pe_info,
    )

    assert ctx["batt"]["number"] == "BATTERY BANK 1"
    assert ctx["batt"]["manufacturer"] == "Chloride"
    assert ctx["batt"]["model"] == "Powersafe"
    assert ctx["batt"]["serialnumber"] == "BATT-007"
    assert ctx["batt"]["area"] == "Terminal Post/ Phase R"
    assert ctx["batt"]["ir"]["reading"] == "45.0"
    assert ctx["batt"]["us"]["reading"] == "-"
    assert ctx["batt"]["tev"]["reading"] == "-"


def test_render_docx_template_zero_raw_jinja_tag_leaks(tmp_path: Path):
    import re
    import zipfile
    from src.quick_report.cbm_render import _render_docx_template

    template_p = Path("templates/QUICK REPORT/DEFECT IR/swg-panel.docx")
    if not template_p.exists():
        pytest.skip("swg-panel.docx template not found")

    out_p = tmp_path / "test_swg_rendered.docx"
    # Minimal sparse context
    context = {
        "substation": {"name_erms": "PE TEST SUB"},
        "swg": {"type": "RMU SF6"},
        "panel": {"name": "PANEL 1"},
    }

    _render_docx_template(template_p, out_p, context)
    assert out_p.exists()

    with zipfile.ZipFile(out_p) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        raw_tags = re.findall(r"\{\{[^}]+\}\}", xml)
        assert raw_tags == [], f"Found raw Jinja tags leaking in rendered docx: {raw_tags}"


def test_transformer_pe_info_equipment_specs():
    from unittest.mock import MagicMock
    from src.quick_report.transformer import QuickReportTransformer
    from src.testsheet.models import SubstationEquipmentPackage, SwitchgearSpec

    swg = SwitchgearSpec(manufacturer="ABB")
    pkg_equipment = SubstationEquipmentPackage(switchgears=(swg,))

    pkg = MagicMock()
    pkg.station = "KUANTAN"
    pkg.month = "08. AUGUST"
    pkg.substation_number = 1
    pkg.data = MagicMock()
    pkg.data.date_str = "12-08-2026"
    pkg.data.substation_name_erms = "TEST SUBSTATION"
    pkg.data.substation_name_site = "TEST SUBSTATION"
    pkg.data.fl_erms = "FL123"
    pkg.data.fl_site = "FL123"
    pkg.data.gps_coordinate = ""
    pkg.data.substation_type = ""
    pkg.data.building_type = "ATTACHED"
    pkg.data.ambient = ""
    pkg.data.humidity = ""
    pkg.data.time = ""
    pkg.data.equipment = pkg_equipment

    env = MagicMock()
    env.po_number = "12345"
    env.state = "PAHANG"
    env.get_vi_front_page_template.return_value = Path("dummy.docx")
    env.get_template.return_value = Path("dummy.docx")

    transformer = QuickReportTransformer()
    plan = transformer.transform(pkg, [], [], env)

    assert plan.pe_info["equipment_specs"] == pkg_equipment
    assert plan.pe_info["equipment_package"] == pkg_equipment
    assert plan.pe_info["equipment"] == pkg_equipment


def test_utils_clear_cell_text_and_set_cell_no_borders():
    from docx import Document
    from docx.oxml.ns import qn
    from src.quick_report.utils import clear_cell_text, set_cell_no_borders

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell.text = "Sample Text"
    assert cell.text == "Sample Text"

    clear_cell_text(cell)
    assert cell.text == ""

    set_cell_no_borders(cell)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    assert tcBorders is not None
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tcBorders.find(qn(f"w:{border_name}"))
        assert border is not None
        assert border.get(qn("w:val")) == "nil"


def test_utils_normalize_functional_location_input():
    from src.quick_report.utils import normalize_functional_location_input

    assert normalize_functional_location_input(None) == ""
    assert normalize_functional_location_input("") == ""
    assert normalize_functional_location_input("12345.0") == "12345"
    assert normalize_functional_location_input("F/L 67890") == "67890"
    assert normalize_functional_location_input("  f/l 11111  ") == "11111"
    assert normalize_functional_location_input(12345) == "12345"





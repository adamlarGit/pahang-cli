"""Unit tests for CBM severity cell shading (Ticket 101) and testsheet measurement ingestion (Ticket 102)."""

from __future__ import annotations

from pathlib import Path
import docx
from docx.oxml.ns import qn
import openpyxl
import pytest

from src.quick_report.cbm_render import (
    _build_swg_render_context,
    _build_tx_render_context,
    _render_docx_template,
)
from src.quick_report.defects import CbmDefectRecord
from src.quick_report.utils import clear_cell_text, set_cell_shading
from src.testsheet.extractor import TestsheetExtractor
from src.testsheet.models import (
    SubstationEquipmentPackage,
    SwitchgearPanelSpec,
    SwitchgearSpec,
    TestsheetData,
    TransformerSpec,
)


def test_set_cell_shading_oxml():
    """Test set_cell_shading correctly writes w:shd XML element on table cell."""
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    cell = table.rows[0].cells[0]

    set_cell_shading(cell, "#EE0000")
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")) == "EE0000"

    set_cell_shading(cell, "00B050")
    assert shd.get(qn("w:fill")) == "00B050"


def test_render_docx_template_severity_shading_ir_defect(tmp_path: Path):
    """Test _render_docx_template applies #EE0000 to IR and #00B050 to US/TEV on IR defect."""
    template_path = Path("templates/QUICK REPORT/DEFECT IR US TEV/swg-panel.docx")
    if not template_path.exists():
        pytest.skip("Template swg-panel.docx not found.")

    out_path = tmp_path / "test_swg_ir_defect.docx"
    rec = CbmDefectRecord(
        equipment="VCB PANEL 1",
        technology="IR",
        defect_area="Cable Box",
        ir_reading="55.4",
    )
    context = _build_swg_render_context(rec, overview=False)
    _render_docx_template(template_path, out_path, context)

    doc = docx.Document(out_path)
    t = doc.tables[0]

    # IR severity cell (Row 18, Col 3) -> should be #EE0000 and text cleared
    cell_ir = t.rows[18].cells[3]
    shd_ir = cell_ir._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_ir is not None
    assert shd_ir.get(qn("w:fill")) == "EE0000"
    assert cell_ir.text.strip() == ""

    # US severity cell (Row 30, Col 4) -> should be #00B050 and text cleared
    cell_us = t.rows[30].cells[4]
    shd_us = cell_us._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_us is not None
    assert shd_us.get(qn("w:fill")) == "00B050"
    assert cell_us.text.strip() == ""

    # TEV severity cell (Row 30, Col 18) -> should be #00B050 and text cleared
    cell_tev = t.rows[30].cells[18]
    shd_tev = cell_tev._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_tev is not None
    assert shd_tev.get(qn("w:fill")) == "00B050"
    assert cell_tev.text.strip() == ""


def test_render_docx_template_severity_shading_us_defect(tmp_path: Path):
    """Test _render_docx_template applies #EE0000 to US and #00B050 to IR/TEV on US defect."""
    template_path = Path("templates/QUICK REPORT/DEFECT IR US TEV/swg-panel.docx")
    if not template_path.exists():
        pytest.skip("Template swg-panel.docx not found.")

    out_path = tmp_path / "test_swg_us_defect.docx"
    rec = CbmDefectRecord(
        equipment="VCB PANEL 1",
        technology="US",
        defect_area="Spout",
        us_reading="18.5",
        us_char="TRACKING",
    )
    context = _build_swg_render_context(rec, overview=False)
    _render_docx_template(template_path, out_path, context)

    doc = docx.Document(out_path)
    t = doc.tables[0]

    # IR severity cell (Row 18, Col 3) -> should be #00B050
    cell_ir = t.rows[18].cells[3]
    shd_ir = cell_ir._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_ir is not None
    assert shd_ir.get(qn("w:fill")) == "00B050"
    assert cell_ir.text.strip() == ""

    # US severity cell (Row 30, Col 4) -> should be #EE0000
    cell_us = t.rows[30].cells[4]
    shd_us = cell_us._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_us is not None
    assert shd_us.get(qn("w:fill")) == "EE0000"
    assert cell_us.text.strip() == ""

    # TEV severity cell (Row 30, Col 18) -> should be #00B050
    cell_tev = t.rows[30].cells[18]
    shd_tev = cell_tev._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_tev is not None
    assert shd_tev.get(qn("w:fill")) == "00B050"
    assert cell_tev.text.strip() == ""


def test_render_docx_template_overview_page_severity_dash(tmp_path: Path):
    """Test overview page renders plain '-' for severity without background fill."""
    template_path = Path("templates/QUICK REPORT/DEFECT IR US TEV/swg-overview.docx")
    if not template_path.exists():
        pytest.skip("Template swg-overview.docx not found.")

    out_path = tmp_path / "test_swg_overview.docx"
    rec = CbmDefectRecord(
        equipment="VCB",
        technology="IR",
        defect_area="Overview",
    )
    context = _build_swg_render_context(rec, overview=True)
    _render_docx_template(template_path, out_path, context, overview=True)

    doc = docx.Document(out_path)
    t = doc.tables[0]
    # Severity row (Row 18) should have '-' or unshaded normal text, not red/green fill
    sev_row = t.rows[18]
    for cell in sev_row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            assert fill not in ("EE0000", "00B050")


def test_testsheet_extractor_us_tev_measurements(tmp_path: Path):
    """Test TestsheetExtractor ingests P6 tev.bg, panel Q/S/T/U/V, and tx K/L values."""
    file_path = tmp_path / "001. TEST_MEASUREMENTS.xlsx"
    wb = openpyxl.Workbook()

    # Sheet 1: PCE Testsheet
    ws_pce = wb.active
    ws_pce.title = "PCE Testsheet"
    ws_pce["C5"] = "PE TEST"
    ws_pce["W5"] = "FL-TEST-01"
    ws_pce["P4"] = "2026-06-15"
    ws_pce["P6"] = "12"  # TEV Background dB

    # Switchgear Panel 1 (Row 10)
    ws_pce["B10"] = "F01"
    ws_pce["C10"] = "INCOMING 1"
    ws_pce["I10"] = "SN-SWG-01"
    ws_pce["Q10"] = "14.2"  # Ultrasound dB
    ws_pce["S10"] = "CORONA"  # Ultrasound char
    ws_pce["T10"] = "28.5"  # TEV dB
    ws_pce["U10"] = "150"  # TEV PPC
    ws_pce["V10"] = "CONTINUOUS"  # TEV char

    # Switchgear Panel 2 (Row 14)
    ws_pce["B14"] = "F02"
    ws_pce["C14"] = "TX 1 FEEDER"
    ws_pce["I14"] = "SN-SWG-02"
    ws_pce["Q14"] = "8.1"
    ws_pce["S14"] = "NORMAL"
    ws_pce["T14"] = "5.0"
    ws_pce["U14"] = "0"
    ws_pce["V14"] = "NORMAL"

    # Transformer 1 (Row 33)
    ws_pce["A33"] = "TX1"
    ws_pce["K33"] = "11.7"  # TX1 US dB
    ws_pce["L33"] = "TRACKING"  # TX1 US char

    # Transformer 2 (Row 38)
    ws_pce["A38"] = "TX2"
    ws_pce["K38"] = "6.3"  # TX2 US dB
    ws_pce["L38"] = "NORMAL"  # TX2 US char

    # Sheet 2: PCE VI
    ws_vi = wb.create_sheet("PCE VI")
    ws_vi["N1"] = "INDOOR"
    ws_vi["C7"] = "PE TEST SITE"
    ws_vi["C17"] = "2"  # 2 Transformers

    # TX 1 on PCE VI (Row 18)
    ws_vi["D18"] = "HERMETICALLY SEALED"
    ws_vi["F18"] = "1000KVA"
    ws_vi["I18"] = "2020"
    ws_vi["L18"] = "TAMCO"
    ws_vi["O18"] = "SN-TX-01"

    # TX 2 on PCE VI (Row 19)
    ws_vi["D19"] = "HERMETICALLY SEALED"
    ws_vi["F19"] = "500KVA"
    ws_vi["I19"] = "2018"
    ws_vi["L19"] = "ABB"
    ws_vi["O19"] = "SN-TX-02"

    wb.save(file_path)

    extractor = TestsheetExtractor()
    data = extractor.extract_testsheet_data(file_path)

    # 1. Verify TEV Background
    assert data.tev_background == "12"

    # 2. Verify Switchgear Panels
    panels = data.equipment.switchgear.panels
    assert len(panels) == 2
    assert panels[0].name == "INCOMING 1"
    assert panels[0].us_reading == "14.2"
    assert panels[0].us_char == "CORONA DISCHARGE"
    assert panels[0].tev_reading == "28.5"
    assert panels[0].tev_ppc == "150"
    assert panels[0].tev_char == "CONTINUOUS"

    assert panels[1].name == "TX 1 FEEDER"
    assert panels[1].us_reading == "8.1"
    assert panels[1].us_char == "NORMAL"
    assert panels[1].tev_reading == "5.0"
    assert panels[1].tev_ppc == "0"
    assert panels[1].tev_char == "NORMAL"

    # 3. Verify Transformers
    transformers = data.equipment.transformers
    assert len(transformers) == 2
    assert transformers[0].tx_id == "Tx 1"
    assert transformers[0].us_reading == "11.7"
    assert transformers[0].us_char == "TRACKING"

    assert transformers[1].tx_id == "Tx 2"
    assert transformers[1].us_reading == "6.3"
    assert transformers[1].us_char == "NORMAL"


def test_build_swg_and_tx_render_contexts_measurement_flow():
    """Test measurement metadata correctly flows into SWG and TX render contexts."""
    panel = SwitchgearPanelSpec(
        panel_no=1,
        panel_feeder_no="F01",
        name="INCOMING 1",
        us_reading="16.4",
        us_char="CORONA",
        tev_reading="31.2",
        tev_ppc="200",
        tev_char="INTERMITTENT",
    )
    swg = SwitchgearSpec(
        switchgear_type="VCB",
        manufacturer="TAMCO",
        panels=(panel,),
    )
    tx = TransformerSpec(
        tx_id="Tx 1",
        manufacturer="TAMCO",
        type="HERMETICALLY SEALED",
        us_reading="13.5",
        us_char="PARTIAL DISCHARGE",
    )
    equipment_pkg = SubstationEquipmentPackage(
        switchgears=(swg,),
        transformers=(tx,),
    )
    ts_data = TestsheetData(
        substation_number=1,
        substation_name_erms="PE TEST",
        tev_background="8",
        equipment=equipment_pkg,
    )
    pe_info = {
        "substation": {
            "name_erms": "PE TEST",
            "tev_bg": "8",
        },
        "testsheet_data": ts_data,
        "equipment": equipment_pkg,
    }

    # 1. SWG Context verification
    swg_rec = CbmDefectRecord(
        equipment="INCOMING 1",
        technology="US",
        defect_area="Cable Box",
        us_reading="16.4",
        us_char="CORONA",
    )
    swg_ctx = _build_swg_render_context(swg_rec, overview=False, pe_info=pe_info)

    assert swg_ctx["us"]["reading"] == "16"
    assert swg_ctx["us"]["char"] == "CORONA DISCHARGE"
    assert swg_ctx["tev"]["reading"] == "31"
    assert swg_ctx["tev"]["ppc"] == "200"
    assert swg_ctx["tev"]["bg"] == "8"
    assert swg_ctx["panel"]["tev"]["bg"] == "8"

    # 2. TX Context verification
    tx_rec = CbmDefectRecord(
        equipment="TX 1",
        technology="IR",
        defect_area="HT Bushing",
        ir_reading="58.2",
    )
    tx_ctx = _build_tx_render_context(tx_rec, overview=False, pe_info=pe_info)

    assert tx_ctx["tx"]["us"]["reading"] == "14"
    assert tx_ctx["tx"]["us"]["char"] == "PARTIAL DISCHARGE"
    assert tx_ctx["us"]["reading"] == "14"
    assert tx_ctx["us"]["char"] == "PARTIAL DISCHARGE"
    assert tx_ctx["tev"]["bg"] == "8"

"""Tests for WhatsApp Report Generation deep module and workflow orchestrator."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from unittest.mock import patch
import openpyxl
import pytest

from src.project.environment import ProjectEnvironment
from src.project.models import ProjectMetadata
from src.project.storage import LocalWorkspaceStorage
from src.whatsapp import (
    QUALIFYING_PDF_PATTERN,
    WhatsAppReportItem,
    WhatsAppReportResources,
    WhatsAppReportSummary,
    build_quick_report_batch_confirmation_lines,
    generate_whatsapp_report,
    get_quick_report_batch_option_title,
    is_selectable_quick_report_batch,
    list_qualifying_batch_pdfs,
)
from src.whatsapp_report_workflow import (
    run_generate_whatsapp_report,
    select_quick_report_batch,
)
from src.workflows.models import WhatsAppReportRequest, WhatsAppReportResult
from src.workflows.service import WorkflowService


def test_qualifying_pdf_pattern_matching() -> None:
    """Verify PDF regex matching for PE number, name stem, and defect suffix."""
    m1 = QUALIFYING_PDF_PATTERN.match("001. SSU CHEROH (VI).pdf")
    assert m1 is not None
    assert m1.group(1) == "001"
    assert m1.group(2) == "SSU CHEROH"
    assert m1.group(3) == "VI"

    m2 = QUALIFYING_PDF_PATTERN.match("002 PE PEKAN.pdf")
    assert m2 is not None
    assert m2.group(1) == "002"
    assert m2.group(2) == "PE PEKAN"
    assert m2.group(3) is None

    m3 = QUALIFYING_PDF_PATTERN.match("10. MARAN (IR+US+VI).pdf")
    assert m3 is not None
    assert m3.group(1) == "10"
    assert m3.group(2) == "MARAN"
    assert m3.group(3) == "IR+US+VI"

    assert QUALIFYING_PDF_PATTERN.match("random_file.docx") is None
    assert QUALIFYING_PDF_PATTERN.match("summary.pdf") is None


def test_list_qualifying_batch_pdfs(tmp_path: Path) -> None:
    """Verify sorting and filtering of batch PDFs by numerical PE prefix."""
    batch_dir = tmp_path / "batch"
    batch_dir.mkdir()

    (batch_dir / "010 PE TEN.pdf").write_bytes(b"%PDF-1.4")
    (batch_dir / "001 PE ONE (VI).pdf").write_bytes(b"%PDF-1.4")
    (batch_dir / "002 PE TWO.pdf").write_bytes(b"%PDF-1.4")
    (batch_dir / "notes.txt").write_text("hello")

    pdfs = list_qualifying_batch_pdfs(batch_dir)
    assert len(pdfs) == 3
    assert pdfs[0].name == "001 PE ONE (VI).pdf"
    assert pdfs[1].name == "002 PE TWO.pdf"
    assert pdfs[2].name == "010 PE TEN.pdf"

    assert is_selectable_quick_report_batch(batch_dir) is True
    assert get_quick_report_batch_option_title(batch_dir) == "batch (3 PDFs)"


def test_build_quick_report_batch_confirmation_lines(tmp_path: Path) -> None:
    """Verify confirmation summary text generation."""
    root_dir = tmp_path / "QUICK REPORT"
    batch_dir = root_dir / "MARAN" / "01-05-2026"
    batch_dir.mkdir(parents=True)

    (batch_dir / "001 SSU CHEROH (VI).pdf").write_bytes(b"%PDF-1.4")
    (batch_dir / "005 SSU LUIT.pdf").write_bytes(b"%PDF-1.4")

    lines = build_quick_report_batch_confirmation_lines(root_dir, batch_dir)
    assert len(lines) == 5
    assert "MARAN / 01-05-2026" in lines[1]
    assert "Qualifying PDFs: 2" in lines[2]
    assert "First PE: 1" in lines[3]
    assert "Last PE: 5" in lines[4]


def create_dummy_total_pe_file(total_pe_path: Path) -> None:
    """Helper to build a mock TOTAL PE.xlsx workbook."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "DataCycle1"
    ws.append(["PE NO", "FL NUMBER", "SUBSTATION NAME", "DATE", "TYPE", "WO", "SCOPE"])
    ws.append([1, "CMRN/SUB1", "SSU CHEROH", datetime(2026, 5, 1), "SSU", 40012345, "FULL"])
    ws.append([2, "CMRN/SUB2", "PE PEKAN", datetime(2026, 5, 1), "P-E", 40012346, "FULL"])
    wb.save(total_pe_path)
    wb.close()


def test_generate_whatsapp_report(tmp_path: Path) -> None:
    """Verify end-to-end rendering of WhatsApp report docx document."""
    root_path = tmp_path / "workspace"
    python_dir = root_path / "PYTHON"
    qr_dir = root_path / "QUICK REPORT" / "01-05-2026"
    qr_dir.mkdir(parents=True)
    python_dir.mkdir(parents=True)

    real_template = Path(__file__).parent.parent / "templates" / "WHATSAPP" / "TEMPLATE WHATSAPP PYTHON.docx"
    template_path = tmp_path / "TEMPLATE WHATSAPP PYTHON.docx"
    if real_template.exists():
        template_path.write_bytes(real_template.read_bytes())
    else:
        from docx import Document
        d = Document()
        d.add_paragraph("WhatsApp Report Date: {{ date }} Station: {{ station }}")
        d.add_paragraph("{% for item in items %}{{ item.name }} - {{ item.defect }} - {{ item.msms }}{% endfor %}")
        d.save(template_path)

    total_pe_path = python_dir / "TOTAL PE.xlsx"
    create_dummy_total_pe_file(total_pe_path)

    (qr_dir / "001 SSU CHEROH (VI).pdf").write_bytes(b"%PDF-1.4")
    (qr_dir / "002 PE PEKAN.pdf").write_bytes(b"%PDF-1.4")

    metadata = ProjectMetadata(
        key="pahang_2026",
        name="Pahang 2026",
        po_number="PO123456",
        state="Pahang",
        voltage_type="11kV",
        year="2026",
        cycle="1",
        technologies=("IR", "VI"),
        base_path=str(root_path),
    )
    storage = LocalWorkspaceStorage(root_path, templates_dir=template_path.parent)
    env = ProjectEnvironment(metadata=metadata, storage=storage)

    resources = WhatsAppReportResources(
        quick_report_dir=root_path / "QUICK REPORT",
        save_dir=python_dir / "WHATSAPP",
        template_path=template_path,
        total_pe_path=total_pe_path,
        station_mapping={"CMRN": "CAMERON HIGHLAND"},
    )

    summary = generate_whatsapp_report(resources, qr_dir)
    assert summary.substations_count == 2
    assert summary.output_path.exists()
    assert summary.output_path.name == "01. CAMERON HIGHLAND 01-05-2026.docx"


def test_workflow_service_run_whatsapp(tmp_path: Path) -> None:
    """Verify WorkflowService.run_whatsapp execution."""
    root_path = tmp_path / "workspace"
    python_dir = root_path / "PYTHON"
    qr_dir = root_path / "QUICK REPORT" / "01-05-2026"
    qr_dir.mkdir(parents=True)
    python_dir.mkdir(parents=True)

    real_template = Path(__file__).parent.parent / "templates" / "WHATSAPP" / "TEMPLATE WHATSAPP PYTHON.docx"
    template_path = tmp_path / "TEMPLATE WHATSAPP PYTHON.docx"
    if real_template.exists():
        template_path.write_bytes(real_template.read_bytes())
    else:
        from docx import Document
        d = Document()
        d.add_paragraph("Report {{ date }} {{ station }}")
        d.save(template_path)

    total_pe_path = python_dir / "TOTAL PE.xlsx"
    create_dummy_total_pe_file(total_pe_path)

    (qr_dir / "001 SSU CHEROH (VI).pdf").write_bytes(b"%PDF-1.4")

    metadata = ProjectMetadata(
        key="pahang_2026",
        name="Pahang 2026",
        po_number="PO123456",
        state="Pahang",
        voltage_type="11kV",
        year="2026",
        cycle="1",
        technologies=("IR", "VI"),
        base_path=str(root_path),
    )
    storage = LocalWorkspaceStorage(root_path, templates_dir=template_path.parent)
    env = ProjectEnvironment(metadata=metadata, storage=storage)

    service = WorkflowService()
    req = WhatsAppReportRequest(report_dir=qr_dir)

    original_resources = env.get_whatsapp_report_resources
    def mock_resources():
        res = original_resources()
        return WhatsAppReportResources(
            quick_report_dir=res.quick_report_dir,
            save_dir=res.save_dir,
            template_path=template_path,
            total_pe_path=res.total_pe_path,
            station_mapping={"CMRN": "CAMERON HIGHLAND"},
        )
    env.get_whatsapp_report_resources = mock_resources

    result = service.run_whatsapp(env, req)
    assert isinstance(result, WhatsAppReportResult)
    assert result.substations_count == 1
    assert result.output_path is not None
    assert result.output_path.exists()

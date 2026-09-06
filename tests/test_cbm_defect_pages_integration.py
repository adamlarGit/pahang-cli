"""Unit & Integration tests for Part 4 CBM Defect Detail Pages & Composer Integration (Ticket 105)."""

from __future__ import annotations

import base64
import gzip
from pathlib import Path
import struct
import docx
from docx.oxml.ns import qn
import pytest

from src.quick_report.cbm_defect_pages import generate_cbm_defect_pages
from src.quick_report.cbm_family import QUICK_REPORT_FAMILY_SPECS_BY_ID, QuickReportFamilySpec
from src.quick_report.defects import CbmDefectRecord
from src.quick_report.models import (
    CbmDefectDetailGroup,
    CbmDefectFamilyPlan,
    CbmDefectGroup,
)
from src.testsheet.models import (
    SubstationEquipmentPackage,
    SwitchgearPanelSpec,
    SwitchgearSpec,
    TestsheetData,
    TransformerSpec,
)


def _build_synthetic_tev_file(path: Path) -> None:
    """Write a synthetic valid FlatBuffers binary buffer eventData.js."""
    events_data = bytearray()
    for peak, phase, cycle in [(25.5, 45, 1), (42.0, 180, 2)]:
        events_data.extend(struct.pack("<fiHHHHff", peak, 100, phase, cycle, 10, 20, 0.5, 1.2))

    vector_bytes = bytearray(struct.pack("<I", 2)) + events_data
    vtable = struct.pack("<HH H H", 8, 8, 4, 0)
    root_table = struct.pack("<i I", 8, 4)

    buf = bytearray()
    buf.extend(struct.pack("<I", 24))
    buf.extend(b"UE01")
    buf.extend(b"\x00" * 8)
    buf.extend(vtable)
    buf.extend(root_table)
    buf.extend(vector_bytes)

    compressed = gzip.compress(bytes(buf))
    b64_str = base64.b64encode(compressed).decode("ascii")
    path.write_text(f'var eventData="{b64_str}";\n', encoding="utf-8")


def _build_synthetic_us_file(path: Path) -> None:
    """Write a synthetic ultrasonic_phase_plot.js file."""
    path.write_text('var ultra_events = {"data": [[14.2, 50, 1], [8.0, 200, 2]]};\n', encoding="utf-8")


def test_generate_cbm_defect_pages_swg_with_prpd_and_severity(tmp_path: Path):
    """Test full CBM defect page generation for SWG with PRPD graphs and severity shading."""
    swg_spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    overview_tpl = Path("templates/QUICK REPORT/DEFECT IR US TEV/swg-overview.docx")
    detail_tpl = Path("templates/QUICK REPORT/DEFECT IR US TEV/swg-panel.docx")

    if not overview_tpl.exists() or not detail_tpl.exists():
        pytest.skip("SWG templates not found.")

    # 1. Setup mock raw survey folder
    raw_survey = tmp_path / "RAW_SURVEY" / "RAW DATA" / "US+TEV" / "20260810T104017_001-PE-TEST"
    feeder1_dir = raw_survey / "SWG" / "FEEDER_1"
    tev_run = feeder1_dir / "20260810T105322_TEV"
    us_run = feeder1_dir / "20260810T104448_Ultrasonic"
    tev_run.mkdir(parents=True)
    us_run.mkdir(parents=True)

    _build_synthetic_tev_file(tev_run / "eventData.js")
    _build_synthetic_us_file(us_run / "ultrasonic_phase_plot.js")

    # 2. Setup equipment package and testsheet data
    panel1 = SwitchgearPanelSpec(
        panel_no=1,
        panel_feeder_no="F01",
        name="INCOMING 1",
        heater_amp="0.5A",
        us_reading="14.2",
        us_char="CORONA",
        tev_reading="25.5",
        tev_ppc="100",
        tev_char="CONTINUOUS",
    )
    swg_equipment = SwitchgearSpec(
        switchgear_type="VCB",
        manufacturer="TAMCO",
        panels=(panel1,),
    )
    equipment_pkg = SubstationEquipmentPackage(switchgears=(swg_equipment,))
    ts_data = TestsheetData(
        substation_number=1,
        substation_name_erms="PE TEST",
        tev_background="8",
        equipment=equipment_pkg,
    )

    pe_info = {
        "substation": {"name_erms": "PE TEST", "tev_bg": "8"},
        "tev_bg": "8",
        "raw_data_dir": tmp_path / "RAW_SURVEY",
        "testsheet_data": ts_data,
        "equipment": equipment_pkg,
    }

    # 3. Setup CBM Defect Family Plan
    defect_rec = CbmDefectRecord(
        equipment="INCOMING 1",
        technology="US",
        defect_area="Cable Box",
        additional_remarks="Discharge sound",
        us_reading="14.2",
        us_char="CORONA",
    )
    detail_group = CbmDefectDetailGroup(
        role_id="panel_area",
        defects=(defect_rec,),
    )
    group = CbmDefectGroup(
        item_key="INCOMING 1",
        item_suffix="",
        defects=(defect_rec,),
        overview=defect_rec,
        detail_groups=(detail_group,),
    )
    family_plan = CbmDefectFamilyPlan(
        spec=swg_spec,
        overview_template=overview_tpl,
        detail_templates=(("panel_area", detail_tpl),),
        groups=(group,),
    )

    # 4. Generate pages
    out_dir = tmp_path / "out_docx"
    generated_files = generate_cbm_defect_pages(
        plan=family_plan,
        output_dir=out_dir,
        substation_number=1,
        pe_info=pe_info,
    )

    assert len(generated_files) == 2  # 1 overview + 1 detail page

    # Verify Overview Page
    overview_doc = docx.Document(generated_files[0])
    assert len(overview_doc.tables) >= 1

    # Verify Detail Page
    detail_doc = docx.Document(generated_files[1])
    assert len(detail_doc.tables) >= 1
    t = detail_doc.tables[0]

    # Verify severity shading on detail page
    # IR severity cell (Row 18, Col 3) -> Non-defective -> Green #00B050
    shd_ir = t.rows[18].cells[3]._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_ir is not None and shd_ir.get(qn("w:fill")) == "00B050"
    assert t.rows[18].cells[3].text.strip() == ""

    # US severity cell (Row 30, Col 4) -> Defective -> Red #EE0000
    shd_us = t.rows[30].cells[4]._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_us is not None and shd_us.get(qn("w:fill")) == "EE0000"
    assert t.rows[30].cells[4].text.strip() == ""

    # TEV severity cell (Row 30, Col 18) -> Non-defective -> Green #00B050
    shd_tev = t.rows[30].cells[18]._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_tev is not None and shd_tev.get(qn("w:fill")) == "00B050"
    assert t.rows[30].cells[18].text.strip() == ""

    # Verify PRPD images were embedded into detail doc
    assert len(detail_doc.inline_shapes) >= 2

    # Verify rendered domain values in tables: heateramp, busbarposition, us.char
    full_table_text = " ".join(cell.text for row in t.rows for cell in row.cells)
    assert "ON:0.5A/OFF:0.0A" in full_table_text
    assert "MAIN" in full_table_text
    assert "CORONA DISCHARGE" in full_table_text



def test_generate_cbm_defect_pages_missing_survey_dir_graceful_fallback(tmp_path: Path):
    """Test CBM defect page generation without raw survey directory leaves PRPD slots clean."""
    swg_spec = QUICK_REPORT_FAMILY_SPECS_BY_ID["swg"]
    overview_tpl = Path("templates/QUICK REPORT/DEFECT IR US TEV/swg-overview.docx")
    detail_tpl = Path("templates/QUICK REPORT/DEFECT IR US TEV/swg-panel.docx")

    if not overview_tpl.exists() or not detail_tpl.exists():
        pytest.skip("SWG templates not found.")

    defect_rec = CbmDefectRecord(
        equipment="FEEDER 1",
        technology="IR",
        defect_area="Cable Box",
        ir_reading="55.0",
    )
    detail_group = CbmDefectDetailGroup(role_id="panel_area", defects=(defect_rec,))
    group = CbmDefectGroup(
        item_key="FEEDER 1",
        item_suffix="",
        defects=(defect_rec,),
        overview=defect_rec,
        detail_groups=(detail_group,),
    )
    family_plan = CbmDefectFamilyPlan(
        spec=swg_spec,
        overview_template=overview_tpl,
        detail_templates=(("panel_area", detail_tpl),),
        groups=(group,),
    )

    out_dir = tmp_path / "out_docx_no_survey"
    generated_files = generate_cbm_defect_pages(
        plan=family_plan,
        output_dir=out_dir,
        substation_number=2,
        pe_info={"substation": {"name_erms": "PE TEST NO SURVEY"}},
    )

    assert len(generated_files) == 2
    detail_doc = docx.Document(generated_files[1])
    assert len(detail_doc.tables) >= 1
    t = detail_doc.tables[0]

    # Verify IR is Red #EE0000, US is Green #00B050, TEV is Green #00B050
    shd_ir = t.rows[18].cells[3]._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_ir is not None and shd_ir.get(qn("w:fill")) == "EE0000"

    shd_us = t.rows[30].cells[4]._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd_us is not None and shd_us.get(qn("w:fill")) == "00B050"
